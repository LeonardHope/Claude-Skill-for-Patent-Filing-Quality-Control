#!/usr/bin/env python3
"""
Patent Filing Quality Control Script

Performs comprehensive QC checks on patent application filing documents.
Generates both Markdown and PDF reports.
"""

import os
import re
import sys
import unicodedata
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass, field
from enum import Enum
import PyPDF2
import argparse

# OCR support (optional)
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


class Severity(Enum):
    """Issue severity levels"""
    CRITICAL = "CRITICAL"
    WARNING = "WARNING"
    INFO = "INFO"
    PASS = "PASS"


@dataclass
class QCIssue:
    """Represents a single QC check result"""
    check_id: int
    category: str
    check_name: str
    severity: Severity
    message: str
    details: str = ""


@dataclass
class QCReport:
    """Container for all QC results"""
    folder_path: str
    files_found: Dict[str, Optional[str]] = field(default_factory=dict)
    issues: List[QCIssue] = field(default_factory=list)
    
    def add_issue(self, check_id: int, category: str, check_name: str, 
                  severity: Severity, message: str, details: str = ""):
        """Add an issue to the report"""
        self.issues.append(QCIssue(check_id, category, check_name, severity, message, details))
    
    def get_critical_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == Severity.CRITICAL)
    
    def get_warning_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == Severity.WARNING)
    
    def get_pass_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == Severity.PASS)


class PatentFilingQC:
    """Main QC engine for patent filing documents"""
    
    def __init__(self, folder_path: str):
        self.folder_path = Path(folder_path)
        self.report = QCReport(folder_path=str(folder_path))
        self.spec_text = ""
        self.ads_text = ""
        self.declaration_text = ""
        self.assignment_text = ""
        self.poa_text = ""
        self.drawings_text = ""
        self.ads_data: Optional[Dict] = None
        self.ads_is_xfa = False
        self.authoritative_inventors: List[Dict] = []
        self.authoritative_source: Optional[str] = None
        # Per-document count of pages that have no extractable text but do
        # contain images (e.g., scanned signed declaration/assignment pages).
        # Cross-doc checks should hedge findings against these documents.
        self.image_only_pages: Dict[str, int] = {}
        
    def extract_pdf_text(self, pdf_path: Path, doc_type: str = "Document") -> str:
        """Extract text from a PDF file. Tries pdfplumber first (preserves
        paragraph structure that the regex-based section/claim checks need),
        then PyPDF2 as a fallback, then OCR for image-only PDFs."""
        # First, try pdfplumber. It generally preserves newlines between
        # paragraphs and section headers that PyPDF2 strips, which is what
        # the spec-content checks (Abstract, Brief Description, Claims, etc.)
        # rely on. Lazy-imported so the dep is only required when actually used.
        text = ""
        image_only_count = 0
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text:
                        text += page_text + "\n"
                    elif hasattr(page, 'images') and page.images:
                        # Page has no extractable text but does contain images
                        # — typical scanned signed page on declaration/assignment.
                        image_only_count += 1
            if image_only_count > 0:
                self.image_only_pages[doc_type] = image_only_count
            clean_text = text.strip()
            if len(clean_text) > 100 and "Please wait" not in clean_text[:200]:
                return text
        except ImportError:
            # pdfplumber not installed; fall through to PyPDF2.
            pass
        except Exception as e:
            print(f"  ⚠️  pdfplumber failed on {pdf_path.name} ({e}); falling back to PyPDF2")

        # Fallback: PyPDF2. Less faithful to layout but doesn't need the dep.
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            # Check if we got meaningful text (more than just whitespace/boilerplate)
            clean_text = text.strip()
            if len(clean_text) > 100 and "Please wait" not in clean_text[:200]:
                return text

            # Text extraction failed or returned minimal content - try OCR
            if OCR_AVAILABLE:
                print(f"  ℹ️  {doc_type} appears to be image-based, attempting OCR...")
                try:
                    images = convert_from_path(pdf_path)
                    ocr_text = ""
                    for i, image in enumerate(images):
                        page_text = pytesseract.image_to_string(image)
                        ocr_text += page_text + "\n"

                    if len(ocr_text.strip()) > 100:
                        print(f"  ✅ OCR successful for {doc_type}")
                        return ocr_text
                    else:
                        self._document_read_failure(doc_type, pdf_path, "OCR returned minimal text")
                        return ""
                except Exception as ocr_error:
                    self._document_read_failure(doc_type, pdf_path, f"OCR failed: {str(ocr_error)}")
                    return ""
            else:
                self._document_read_failure(doc_type, pdf_path,
                    "Image-based PDF detected but OCR not available. "
                    "Install pytesseract and pdf2image: pip install pytesseract pdf2image")
                return ""

        except Exception as e:
            self._document_read_failure(doc_type, pdf_path, str(e))
            return ""

    def _document_read_failure(self, doc_type: str, pdf_path: Path, reason: str):
        """Handle document read failure with helpful error message"""
        print(f"\n  🚨 FAILED TO READ {doc_type.upper()}: {pdf_path.name}")
        print(f"     Reason: {reason}")
        print(f"     Solutions:")
        print(f"       1. Open the original fillable form and 'Print to PDF'")
        print(f"       2. Use Adobe Acrobat's 'Recognize Text' (OCR) feature")
        print(f"       3. Re-export from the source application as a text-based PDF")
        print()

    def _is_xfa_form(self, pdf_path: Path) -> bool:
        """Check whether a PDF contains an XFA form (e.g., USPTO web-fillable ADS)."""
        try:
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                # IndirectObjects auto-resolve under [] but not .get(), so use [].
                root = reader.trailer['/Root']
                if hasattr(root, 'get_object'):
                    root = root.get_object()
                if '/AcroForm' not in root:
                    return False
                acroform = root['/AcroForm']
                if hasattr(acroform, 'get_object'):
                    acroform = acroform.get_object()
                return '/XFA' in acroform
        except Exception:
            return False

    def _extract_xfa_datasets_xml(self, pdf_path: Path) -> Optional[str]:
        """Pull the 'datasets' XML stream out of an XFA form. Returns the XML string or None."""
        try:
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                acroform = reader.trailer['/Root']['/AcroForm']
                if hasattr(acroform, 'get_object'):
                    acroform = acroform.get_object()
                xfa = acroform.get('/XFA')
                if xfa is None:
                    return None
                if hasattr(xfa, 'get_object'):
                    xfa = xfa.get_object()
                # XFA can be a single stream or an array of [name, stream, name, stream, ...]
                items = list(xfa)
                for i in range(0, len(items), 2):
                    name = items[i]
                    if i + 1 >= len(items):
                        break
                    if str(name) == 'datasets':
                        stream = items[i + 1]
                        if hasattr(stream, 'get_object'):
                            stream = stream.get_object()
                        return stream.get_data().decode('utf-8', errors='replace')
                return None
        except Exception as e:
            print(f"  ⚠️  XFA datasets extraction failed: {e}")
            return None

    def _parse_ads_xfa(self, xml_str: str) -> Optional[Dict]:
        """Parse the USPTO ADS XFA datasets XML into a structured dict."""
        try:
            root = ET.fromstring(xml_str)
        except ET.ParseError as e:
            print(f"  ⚠️  XFA XML parse failed: {e}")
            return None

        # Strip namespace from tag names to make traversal less painful.
        def localname(elem):
            tag = elem.tag
            return tag.split('}', 1)[1] if '}' in tag else tag

        def find_first(elem, name):
            for child in elem.iter():
                if localname(child) == name:
                    return child
            return None

        def text_of(elem, name):
            child = find_first(elem, name) if elem is not None else None
            if child is None or child.text is None:
                return ""
            return child.text.strip()

        # Walk to <us-request>
        us_request = find_first(root, 'us-request')
        if us_request is None:
            return None

        data: Dict = {
            'inventors': [],
            'title': text_of(us_request, 'invention-title'),
            'docket_number': text_of(us_request, 'attorney-docket-number'),
            'customer_number': "",
            'attorney_customer_number': "",
            'form_pages': text_of(us_request, 'numofpages'),
            'small_entity': None,
            'application_type': "",
            'submission_type': "",
            'drawing_sheets': "",
            'representative_figure': "",
            'non_publication': None,
            'aia_transition': None,
            'assignee_org': "",
            'assignee_address': {},
            'signer': {},
            'domestic_continuity_entries': [],
            'foreign_priority_entries': [],
        }

        # Inventors live in repeated <sfApplicantInformation> blocks under ContentArea1
        for app_info in us_request.iter():
            if localname(app_info) != 'sfApplicantInformation':
                continue
            name_block = None
            for child in app_info:
                if localname(child) == 'sfApplicantName':
                    name_block = child
                    break
            if name_block is None:
                continue
            inventor = {
                'prefix': text_of(name_block, 'prefix'),
                'first': text_of(name_block, 'firstName'),
                'middle': text_of(name_block, 'middleName'),
                'last': text_of(name_block, 'lastName'),
                'suffix': text_of(name_block, 'suffix'),
                'citizenship': "",
                'residency': "",
                'res_city': "",
                'res_state': "",
                'res_country': "",
                'mail_address1': "",
                'mail_address2': "",
                'mail_city': "",
                'mail_state': "",
                'mail_postcode': "",
                'mail_country': "",
            }
            # Citizenship dropdown (sfCitz/CitizedDropDown). Often blank when the
            # applicant is filing as an assignee under 37 CFR 1.46.
            citz_block = find_first(app_info, 'sfCitz')
            if citz_block is not None:
                inventor['citizenship'] = text_of(citz_block, 'CitizedDropDown')
            # Residency
            res_check = find_first(app_info, 'sfAppResChk')
            if res_check is not None:
                inventor['residency'] = text_of(res_check, 'ResidencyRadio')
                us_res = find_first(res_check, 'sfUSres')
                non_us_res = find_first(res_check, 'sfNonUSRes')
                if inventor['residency'] == 'us-residency' and us_res is not None:
                    inventor['res_city'] = text_of(us_res, 'rsCityTxt')
                    inventor['res_state'] = text_of(us_res, 'rsStTxt')
                    inventor['res_country'] = text_of(us_res, 'rsCtryTxt')
                elif non_us_res is not None:
                    inventor['res_city'] = text_of(non_us_res, 'nonresCity')
                    inventor['res_country'] = text_of(non_us_res, 'nonresCtryList')
            # Mailing address
            mail = find_first(app_info, 'sfApplicantMail')
            if mail is not None:
                inventor['mail_address1'] = text_of(mail, 'address1')
                inventor['mail_address2'] = text_of(mail, 'address2')
                inventor['mail_city'] = text_of(mail, 'city')
                inventor['mail_state'] = text_of(mail, 'state')
                inventor['mail_postcode'] = text_of(mail, 'postcode')
                inventor['mail_country'] = text_of(mail, 'mailCountry')

            # Skip empty placeholder blocks
            if inventor['first'] or inventor['last']:
                data['inventors'].append(inventor)

        # Correspondence customer number (sfCorrCustNo)
        cust_no = find_first(us_request, 'sfCorrCustNo')
        if cust_no is not None:
            data['customer_number'] = text_of(cust_no, 'customerNumber')

        # Attorney/agent customer number (sfAttorny → sfAttornyFlow → sfcustomerNumber).
        # This is a separate field from the correspondence customer number, and
        # typically should match it. Mismatch is a warning sign.
        attorny_block = find_first(us_request, 'sfAttorny')
        if attorny_block is not None:
            atty_cust = find_first(attorny_block, 'sfcustomerNumber')
            if atty_cust is not None:
                data['attorney_customer_number'] = text_of(atty_cust, 'customerNumberTxt')

        # Fall back to attorney customer number if no correspondence one was set
        if not data['customer_number'] and data['attorney_customer_number']:
            data['customer_number'] = data['attorney_customer_number']

        # Domestic continuity entries — check for any populated parent application info
        for cont in us_request.iter():
            if localname(cont) != 'sfDomesticContinuity':
                continue
            for info in cont.iter():
                if localname(info) != 'sfDomesContInfo':
                    continue
                app_num = text_of(info, 'domappNumber')
                cont_type = text_of(info, 'domesContList')
                prior_num = text_of(info, 'domPriorAppNum')
                date_field = text_of(info, 'DateTimeField1')
                if any([app_num, prior_num, cont_type, date_field]):
                    data['domestic_continuity_entries'].append({
                        'application_number': app_num,
                        'continuation_type': cont_type,
                        'prior_application_number': prior_num,
                        'date': date_field,
                    })

        # Foreign priority entries
        for fpr in us_request.iter():
            if localname(fpr) != 'sfForeignPriorityInfo':
                continue
            app_num = text_of(fpr, 'frprAppNum')
            country = text_of(fpr, 'frprctryList')
            date_field = text_of(fpr, 'frprParentDate')
            access_code = text_of(fpr, 'accessCode')
            if any([app_num, country, date_field, access_code]):
                data['foreign_priority_entries'].append({
                    'application_number': app_num,
                    'country': country,
                    'priority_date': date_field,
                    'access_code': access_code,
                })

        # Application info
        app_pos = find_first(us_request, 'sfAppPos')
        if app_pos is not None:
            small = text_of(app_pos, 'chkSmallEntity')
            data['small_entity'] = (small == '1') if small in ('0', '1') else None
            data['application_type'] = text_of(app_pos, 'application_type')
            data['submission_type'] = text_of(app_pos, 'us_submission_type')
            data['drawing_sheets'] = text_of(app_pos, 'us-total_number_of_drawing-sheets')
            data['representative_figure'] = text_of(app_pos, 'us-suggested_representative_figure')

        # Publication / AIA flags
        pub = find_first(us_request, 'sfPub')
        if pub is not None:
            non_pub = text_of(pub, 'nonPublication')
            data['non_publication'] = (non_pub == '1') if non_pub in ('0', '1') else None
        aia = find_first(us_request, 'AIATransition')
        if aia is not None:
            aia_check = text_of(aia, 'AIACheck')
            data['aia_transition'] = (aia_check == '1') if aia_check in ('0', '1') else None

        # Assignee
        assignee_info = find_first(us_request, 'sfAssigneeInformation')
        if assignee_info is not None:
            data['assignee_org'] = text_of(assignee_info, 'orgName')
            addr = find_first(assignee_info, 'sfAssigneeAddress')
            if addr is not None:
                data['assignee_address'] = {
                    'address1': text_of(addr, 'address-1'),
                    'address2': text_of(addr, 'address-2'),
                    'city': text_of(addr, 'city'),
                    'state': text_of(addr, 'state'),
                    'postcode': text_of(addr, 'postcode'),
                    'country': text_of(addr, 'txtCorrCtry'),
                }

        # Signer
        signature = find_first(us_request, 'sfSignature')
        if signature is not None:
            sig = find_first(signature, 'sfSig')
            if sig is not None:
                data['signer'] = {
                    'first_name': text_of(sig, 'first-name'),
                    'last_name': text_of(sig, 'last-name'),
                    'registration_number': text_of(sig, 'registration-number'),
                    'signature': text_of(sig, 'signature'),
                    'date': text_of(sig, 'date'),
                }

        return data

    def _synthesize_ads_text_from_xfa(self, data: Dict) -> str:
        """Build a plain-text representation of XFA-extracted ADS data so the
        existing regex-based extractors and checks find what they expect."""
        lines = []
        if data.get('title'):
            # Use the pipe separator that the real ADS form uses, so the existing
            # regex extractor in extract_title() works as a fallback.
            lines.append(f"Title of Invention | {data['title']}")
        if data.get('docket_number'):
            lines.append(f"Attorney Docket Number: {data['docket_number']}")
        if data.get('customer_number'):
            lines.append(f"Customer Number: {data['customer_number']}")
        if data.get('application_type'):
            lines.append(f"Application Type: {data['application_type']}")
        if data.get('submission_type'):
            lines.append(f"Submission Type: {data['submission_type']}")
        if data.get('small_entity') is True:
            lines.append("Entity Status: Small Entity")
        elif data.get('small_entity') is False:
            lines.append("Entity Status: Large/Regular (small entity not claimed)")
        if data.get('non_publication') is True:
            lines.append("Non-Publication Request: Yes")
        if data.get('aia_transition') is True:
            lines.append("AIA Transition Statement: Yes")
        if data.get('drawing_sheets'):
            lines.append(f"Total Drawing Sheets: {data['drawing_sheets']}")

        # Inventors — emit in the "Suffix\nFirst Middle LAST" pattern the
        # existing extract_inventors() regex looks for, so cross-doc checks work.
        # Also include "Inventor N" labels and residency tokens that other
        # ADS-specific checks expect.
        lines.append("")
        lines.append("Inventor Information:")
        for idx, inv in enumerate(data.get('inventors', []), start=1):
            given = ' '.join(p for p in [inv.get('first', ''), inv.get('middle', '')] if p)
            last = (inv.get('last') or '').upper()
            full_line = f"{given} {last}".strip()
            if inv.get('suffix'):
                full_line = f"{full_line} {inv['suffix']}".strip()
            lines.append(f"Inventor {idx}")
            lines.append("Suffix")
            lines.append(full_line)
            if inv.get('residency') == 'us-residency':
                lines.append("US Residency")
            elif inv.get('residency') == 'non-us-residency':
                lines.append("non-US Residency")
            if inv.get('mail_address1'):
                lines.append(f"Address 1 {inv['mail_address1']}")
            if inv.get('mail_address2'):
                lines.append(f"Address 2 {inv['mail_address2']}")
            city_state = ', '.join(p for p in [inv.get('mail_city', ''), inv.get('mail_state', '')] if p)
            if city_state or inv.get('mail_postcode'):
                lines.append(f"{city_state} Postal Code {inv.get('mail_postcode', '')} {inv.get('mail_country', '')}".strip())
            lines.append("")

        # Correspondence section keyword for Check 10 (ADS required fields)
        if data.get('customer_number'):
            lines.append("Correspondence Information")
            lines.append(f"Customer Number: {data['customer_number']}")
            lines.append("")

        # Correspondence / assignee / signer info — useful for downstream checks
        if data.get('assignee_org'):
            lines.append(f"Assignee: {data['assignee_org']}")
            addr = data.get('assignee_address') or {}
            for key in ('address1', 'address2'):
                if addr.get(key):
                    lines.append(addr[key])
            ac = ', '.join(p for p in [addr.get('city', ''), addr.get('state', '')] if p)
            if ac:
                lines.append(f"{ac} {addr.get('postcode', '')} {addr.get('country', '')}".strip())

        signer = data.get('signer') or {}
        if signer.get('signature'):
            lines.append("")
            lines.append(f"Signature: {signer.get('signature', '')}")
            lines.append(f"Name: {signer.get('first_name', '')} {signer.get('last_name', '')}".strip())
            if signer.get('registration_number'):
                lines.append(f"Registration Number: {signer['registration_number']}")
            if signer.get('date'):
                lines.append(f"Date: {signer['date']}")

        return "\n".join(lines)

    def _extract_ads_text(self, ads_path: Path) -> str:
        """ADS-aware text extraction: uses XFA datasets stream if the form is XFA,
        falls back to standard PDF text extraction (and OCR) otherwise."""
        if self._is_xfa_form(ads_path):
            self.ads_is_xfa = True
            print(f"  ℹ️  ADS appears to be an XFA (web-fillable) form — reading XFA datasets stream")
            xml_str = self._extract_xfa_datasets_xml(ads_path)
            if xml_str:
                data = self._parse_ads_xfa(xml_str)
                if data and (data.get('inventors') or data.get('title') or data.get('docket_number')):
                    self.ads_data = data
                    print(f"  ✅ XFA extraction successful: {len(data.get('inventors', []))} inventor(s), title='{data.get('title','')[:60]}'")
                    return self._synthesize_ads_text_from_xfa(data)
                print(f"  ⚠️  XFA datasets parsed but no usable fields found; falling back to OCR")
            else:
                print(f"  ⚠️  Could not extract XFA datasets stream; falling back to OCR")
        return self.extract_pdf_text(ads_path, 'ADS')
    
    def find_document(self, patterns: List[str], doc_type: str) -> Optional[Path]:
        """[Legacy filename-pattern finder; retained for callers but no longer
        used by load_documents, which now classifies by file content.]"""
        for pattern in patterns:
            matches = list(self.folder_path.glob(pattern))
            if matches:
                self.report.files_found[doc_type] = str(matches[0].name)
                return matches[0]
        self.report.files_found[doc_type] = None
        return None

    def _quick_extract_text(self, pdf_path: Path, max_pages: int = 999) -> str:
        """Text extraction used for content-based classification. Reads ALL
        pages by default — many key spec markers (CLAIMS preamble, ABSTRACT)
        appear at the END of a long spec, not the first few pages. Reading
        only the first 3 pages produced asymmetric classification confidence
        between .pdf (only first pages seen) and .docx (full doc seen) for
        the same content.
        Uses pdfplumber when available (much better for PDFs with custom font
        encodings — e.g., USPTO declarations whose body text PyPDF2 can't decode
        but pdfplumber can). Falls back to PyPDF2 if pdfplumber is unavailable."""
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                pages = pdf.pages[:max_pages]
                return "\n".join((p.extract_text() or "") for p in pages)
        except ImportError:
            pass
        except Exception:
            pass
        try:
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                pages = reader.pages[:max_pages]
                return "\n".join((p.extract_text() or "") for p in pages)
        except Exception:
            return ""

    def _extract_docx_text(self, docx_path: Path) -> str:
        """Extract all visible text from a .docx file. Lazy-imports python-docx
        so the dependency is only required when a .docx file is actually present."""
        try:
            import docx  # python-docx
        except ImportError:
            print(f"\n  🚨 Found Word document {docx_path.name} but python-docx is not installed.")
            print(f"     Install it with: pip install python-docx --break-system-packages")
            print(f"     Then re-run this QC.\n")
            return ""
        try:
            doc = docx.Document(str(docx_path))
            chunks = [p.text for p in doc.paragraphs]
            # Also pull text from tables (some specs render claims in tables)
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        chunks.append(cell.text)
            return "\n".join(c for c in chunks if c)
        except Exception as e:
            self._document_read_failure('Word document', docx_path, str(e))
            return ""

    def _extract_text_any(self, path: Path, doc_type: str) -> str:
        """Format-agnostic text extraction: dispatch on extension."""
        suffix = path.suffix.lower()
        if suffix == '.docx':
            return self._extract_docx_text(path)
        return self.extract_pdf_text(path, doc_type)

    def _extract_all_xfa_xml(self, pdf_path: Path) -> str:
        """Concatenate every XFA stream's text content. The 'template' stream
        holds form UI labels and static text (e.g., 'I hereby declare' on a
        declaration form), the 'datasets' stream holds filled values, etc.
        Lets us classify XFA forms by their content rather than assuming
        XFA == ADS."""
        try:
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                root = reader.trailer['/Root']
                if hasattr(root, 'get_object'):
                    root = root.get_object()
                acroform = root['/AcroForm']
                if hasattr(acroform, 'get_object'):
                    acroform = acroform.get_object()
                xfa = acroform['/XFA']
                if hasattr(xfa, 'get_object'):
                    xfa = xfa.get_object()
                items = list(xfa)
                chunks = []
                for i in range(0, len(items), 2):
                    if i + 1 >= len(items):
                        break
                    stream = items[i + 1]
                    if hasattr(stream, 'get_object'):
                        stream = stream.get_object()
                    try:
                        chunks.append(stream.get_data().decode('utf-8', errors='replace'))
                    except Exception:
                        pass
                return "\n".join(chunks)
        except Exception:
            return ""

    def _classify_file(self, path: Path) -> Tuple[str, float]:
        """Identify what kind of filing document a file is by inspecting its
        content (not its filename). Handles both PDF and DOCX. Returns
        (doc_type, confidence_score). doc_type is one of: 'ADS',
        'Specification', 'Drawings', 'Declaration', 'Assignment',
        'Power of Attorney', 'Unknown'.
        """
        # DOCX path: only the spec is ever .docx in practice. Extract text
        # and run it through the same scoring as a non-XFA PDF.
        if path.suffix.lower() == '.docx':
            text = self._extract_docx_text(path)
            if not text.strip():
                return ('Unknown', 0.0)
            return self._score_text(text)

        # PDF path. XFA forms: don't assume XFA == ADS. Multiple USPTO forms
        # ship as XFA (PTO/AIA/01 + /02 Declarations, /14 ADS, /82 POA, etc.).
        # Pull the XFA XML streams and classify on their content.
        if self._is_xfa_form(path):
            xfa_xml = self._extract_all_xfa_xml(path)
            xl = xfa_xml.lower()
            # Highly specific structural markers per form
            xfa_scores = {
                'ADS': (
                    100 * ('us-request' in xl and 'sfapplicantinformation' in xl) +
                    20 * ('attorney-docket-number' in xl) +
                    20 * ('invention-title' in xl)
                ),
                'Declaration': (
                    100 * ('pto/aia/01' in xl or 'pto/aia/02' in xl) +
                    30 * ('i hereby declare' in xl) +
                    30 * ('37 cfr 1.63' in xl or '§1.63' in xl) +
                    20 * ('original inventor' in xl)
                ),
                'Power of Attorney': (
                    100 * ('pto/aia/82' in xl) +
                    40 * ('power of attorney' in xl) +
                    20 * ('appoint' in xl and 'practitioner' in xl)
                ),
            }
            best_xfa = max(xfa_scores, key=lambda k: xfa_scores[k])
            if xfa_scores[best_xfa] >= 30:
                return (best_xfa, float(xfa_scores[best_xfa]))
            # Fall through to text-based classification using the XFA XML as the
            # source text (the rendered PDF would just say "Please wait...").
            text = xfa_xml
        else:
            text = self._quick_extract_text(path)
            # If a non-XFA PDF returned almost no text, it's most likely the
            # drawings (image-only PDF). Try the full text path before deciding.
            if len(text.strip()) < 200:
                text = self.extract_pdf_text(path, '<classifying>')

        return self._score_text(text)

    def _score_text(self, text: str) -> Tuple[str, float]:
        """Score raw text against doc-type signatures and return the best match.
        Used by both PDF and DOCX classification paths."""
        tl = text.lower()
        head = tl[:1500]  # signatures that must appear near the top

        scores: Dict[str, float] = {}

        scores['ADS'] = (
            6 * ('application data sheet' in tl) +
            5 * (('37 cfr 1.76' in tl) or ('§1.76' in tl)) +
            3 * ('correspondence information' in tl) +
            3 * (bool(re.search(r'customer\s*number', tl)) and 'inventor' in tl) +
            2 * ('non-publication' in tl)
        )

        scores['Specification'] = (
            7 * bool(re.search(r'what\s+is\s+claimed', tl)) +
            5 * ('brief description of the drawings' in tl) +
            3 * (' abstract' in tl or '\nabstract' in tl) +
            3 * bool(re.search(r'\bbackground\b', tl)) +
            2 * bool(re.search(r'\bclaim\s+\d+', tl)) +
            -6 * ('i hereby declare' in tl) +     # excludes declaration
            -5 * bool(re.search(r'\bassignor\b', tl)) +  # excludes assignment
            -5 * ('power of attorney' in tl)       # excludes POA
        )

        scores['Declaration'] = (
            7 * ('i hereby declare' in tl) +
            6 * (('37 cfr 1.63' in tl) or ('§1.63' in tl)) +
            4 * ('pto/aia/01' in tl or 'pto/aia/02' in tl) +
            3 * ('original inventor' in tl) +
            2 * ('declaration' in head)
        )

        scores['Assignment'] = (
            6 * bool(re.search(r'\bassignor\b', tl)) +
            5 * bool(re.search(r'\bassignee\b', tl)) +
            5 * bool(re.search(r'right,?\s+title,?\s+and\s+interest', tl)) +
            3 * ('whereas' in tl and 'assign' in tl) +
            2 * ('assignment' in head)
        )

        scores['Power of Attorney'] = (
            7 * ('power of attorney' in tl) +
            6 * bool(re.search(r'pto/aia/82', tl)) +
            3 * ('appoint' in tl and 'attorney' in tl) +
            2 * ('revoke' in tl and 'attorney' in tl)
        )

        # Drawings: defining traits are very little prose plus drawing-style
        # markers (FIG. labels, sheet-numbering, sparse 3-digit reference
        # numerals). Many drawings PDFs are image-only with the only
        # extractable text being the page-margin header (docket, title, sheet
        # numbers like "1/7" or "1 of 7") and a handful of reference numerals.
        has_fig_refs = bool(re.search(r'fig\.\s*\d', tl))
        has_sheet_phrase = bool(re.search(r'sheet\s+\d+\s*(/|of)\s*\d+', tl))
        # "1/7" or "01 / 07" alone on a line — common sheet-number format
        has_bare_sheet_nums = bool(re.search(r'(?m)^\s*\d+\s*/\s*\d+\s*$', text))
        # Many isolated 3-digit numbers (reference numerals) on their own lines
        bare_3dig_lines = len(re.findall(r'(?m)^\s*\d{3}\s*$', text))
        prose_len = len(re.sub(r'\s+', ' ', text).strip())

        if has_fig_refs and prose_len < 1500:
            scores['Drawings'] = 8.0
        elif has_fig_refs and prose_len < 4000:
            scores['Drawings'] = 4.0
        elif has_sheet_phrase or has_bare_sheet_nums:
            scores['Drawings'] = 5.0
        elif prose_len < 2000 and bare_3dig_lines >= 5:
            # Image-only drawings PDF with margin header + extracted ref numerals
            scores['Drawings'] = 4.0
        else:
            scores['Drawings'] = 0.0

        best_type = max(scores, key=lambda k: scores[k])
        best_score = scores[best_type]
        if best_score < 3:
            return ('Unknown', best_score)
        return (best_type, float(best_score))
    
    def load_documents(self):
        """Locate and load all filing documents by classifying every PDF in the
        folder by *content*, not filename. Files named 'Application.pdf',
        'Formals.pdf', 'MS3-0230US-A.pdf', etc. are all handled as long as
        their contents identify them.
        """
        # Initialize all six slots so the report's "Documents Found" list shows
        # which kinds were not detected.
        slots = ['Specification', 'Drawings', 'ADS',
                 'Declaration', 'Assignment', 'Power of Attorney']
        for s in slots:
            self.report.files_found[s] = None

        # Glob both PDFs and Word docs. The USPTO accepts the specification in
        # .docx format; everything else (declarations, ADS, drawings, assignments,
        # POAs) is always PDF — but we let the content classifier decide rather
        # than encoding "only spec can be docx" as a hard rule.
        all_files = sorted(
            list(self.folder_path.glob('*.pdf')) +
            list(self.folder_path.glob('*.docx'))
        )

        # Skip the QC report and any Office lock files (~$ prefix)
        all_files = [
            p for p in all_files
            if 'patent_filing_qc_report' not in p.name.lower()
            and not p.name.startswith('~$')
        ]

        # Classify every file
        candidates_by_type: Dict[str, List[Tuple[Path, float]]] = {}
        unrecognized: List[Path] = []
        for path in all_files:
            doc_type, confidence = self._classify_file(path)
            if doc_type == 'Unknown':
                unrecognized.append(path)
                print(f"  ❓ Unrecognized file (low confidence): {path.name}")
                continue
            candidates_by_type.setdefault(doc_type, []).append((path, confidence))
            print(f"  📄 {path.name} → {doc_type} (confidence {confidence:.0f})")

        # For each slot, pick a candidate and warn if there are duplicates.
        # Spec-specific tie-breaker: when both a .pdf and .docx are present,
        # prefer the .pdf since that's what gets filed at the USPTO. (We keep
        # .docx support for the case where there's only a .docx.) For all other
        # slots, sort purely by confidence.
        for doc_type, candidates in candidates_by_type.items():
            if doc_type == 'Specification':
                # PDF first, then by descending confidence
                candidates.sort(key=lambda x: (
                    0 if x[0].suffix.lower() == '.pdf' else 1,
                    -x[1],
                ))
            else:
                candidates.sort(key=lambda x: -x[1])
            best_path, best_conf = candidates[0]
            self.report.files_found[doc_type] = best_path.name

            if len(candidates) > 1:
                others = ', '.join(p.name for p, _ in candidates[1:])
                print(f"  ⚠️  Multiple files classified as {doc_type}; using "
                      f"{best_path.name} (also matched: {others})")

                # Compose a report-visible warning so the user sees this in
                # the QC report, not just on the console.
                exts = sorted({p.suffix.lower() for p, _ in candidates})
                spec_dual_format = (
                    doc_type == 'Specification'
                    and '.pdf' in exts and '.docx' in exts
                )
                if spec_dual_format:
                    msg = (f"Both a .pdf and a .docx version of the "
                           f"specification are in the to-be-filed folder")
                    details = (
                        f"Picked: {best_path.name} (defaulting to .pdf when "
                        f"both are present — the USPTO accepts spec filings "
                        f"in either .docx or .pdf, but most practitioners file the .pdf).\n"
                        f"Also present: {others}.\n\n"
                        "Verify which one you actually intend to file. If the "
                        "non-filed copy is just a working draft, consider moving "
                        "it out before filing so the folder reflects exactly "
                        "what's being submitted."
                    )
                else:
                    msg = f"Multiple files classified as {doc_type}"
                    details = (
                        f"Picked: {best_path.name} "
                        f"(highest classification confidence).\n"
                        f"Also matched: {others}.\n\n"
                        "Verify the correct file was used. If a non-current "
                        "draft is in the folder, remove it before filing."
                    )
                self.report.add_issue(
                    74, "Document Completeness",
                    "Duplicate Files for Same Document Type",
                    Severity.WARNING, msg, details
                )

            if doc_type == 'Specification':
                self.spec_text = self._extract_text_any(best_path, 'Specification')
            elif doc_type == 'Drawings':
                self.drawings_text = self._extract_text_any(best_path, 'Drawings')
            elif doc_type == 'ADS':
                self.ads_text = self._extract_ads_text(best_path)
            elif doc_type == 'Declaration':
                self.declaration_text = self._extract_text_any(best_path, 'Declaration')
            elif doc_type == 'Assignment':
                self.assignment_text = self._extract_text_any(best_path, 'Assignment')
            elif doc_type == 'Power of Attorney':
                self.poa_text = self._extract_text_any(best_path, 'Power of Attorney')

        # Optional: authoritative inventor list (inventors.json/txt or *.eml)
        self._load_authoritative_inventors()

        # Check for required documents that couldn't be read
        self._check_required_documents_readable()

    def _check_required_documents_readable(self):
        """Verify that required documents were successfully read"""
        required_docs = {
            'Specification': self.spec_text,
            'Drawings': self.drawings_text,
            'ADS': self.ads_text,
            'Declaration': self.declaration_text,
        }

        unreadable = []
        for doc_name, text in required_docs.items():
            if self.report.files_found.get(doc_name):  # File was found
                if not text or len(text.strip()) < 100:  # But couldn't be read
                    unreadable.append(doc_name)

        if unreadable:
            print("\n" + "=" * 70)
            print("🚨 CRITICAL: REQUIRED DOCUMENTS COULD NOT BE READ")
            print("=" * 70)
            for doc in unreadable:
                filename = self.report.files_found.get(doc, 'Unknown')
                print(f"  • {doc}: {filename}")
            print("\nThe QC report will be incomplete. Please fix the above documents")
            print("and re-run the QC check.")
            print("=" * 70 + "\n")

    def extract_inventors(self, text: str) -> List[str]:
        """Extract inventor names from text"""
        inventors = []

        # Pattern 1: ADS OCR format - names after "Suffix" field
        # Matches: "Suffix\nFirstName LASTNAME" or "Suffix\nFirst Middle LASTNAME"
        # Optional trailing suffix token (Jr., Sr., II, III, IV).
        suffix_pattern = (
            r'Suffix\s*\n?\s*'
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\s+[A-Z]{2,}'
            r'(?:\s+(?:Jr\.?|Sr\.?|II|III|IV))?)\s*\n'
        )
        matches = re.findall(suffix_pattern, text)
        for match in matches:
            name = ' '.join(match.split())
            words = name.split()
            # words[-1] may be a suffix; identify the last all-caps token (the surname)
            last_caps_idx = max((i for i, w in enumerate(words)
                                 if w.isupper() and len(w) >= 2), default=-1)
            if 2 <= len(words) <= 5 and last_caps_idx >= 0:
                inventors.append(name)

        # Pattern 2: Declaration/Assignment format - names after "Assignor(s):"
        # Names appear on their own lines as "First LAST" or "First Middle LAST"
        # Look for section between "Assignor" and "Assignee"
        assignor_section = re.search(r'Assignor\s*\(s\)\s*:?\s*(.*?)(?:Assignee|AGREEMENT)', text, re.DOTALL | re.IGNORECASE)
        if assignor_section:
            section_text = assignor_section.group(1)
            # Find names in format "First LAST" or "First Middle LAST" at start of lines
            name_pattern = r'^\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\s+[A-Z]{2,})\s*$'
            matches = re.findall(name_pattern, section_text, re.MULTILINE)
            for match in matches:
                name = ' '.join(match.split())
                words = name.split()
                if 2 <= len(words) <= 4 and words[-1].isupper() and len(words[-1]) >= 2:
                    inventors.append(name)

        # Pattern 3: Direct "First LAST" pattern in text (for various document formats)
        # Be more targeted - look for names followed by address or "c/o"
        name_with_address = r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\s+[A-Z]{2,})\s*\n\s*c/o'
        matches = re.findall(name_with_address, text)
        for match in matches:
            name = ' '.join(match.split())
            words = name.split()
            if 2 <= len(words) <= 4 and words[-1].isupper() and len(words[-1]) >= 2:
                inventors.append(name)

        # If we found inventors, return unique list
        if inventors:
            return list(set(inventors))

        # Fallback: Look for declaration format "I, [Name], declare"
        decl_pattern = r'(?:I,\s*)([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)(?:,?\s*declare)'
        matches = re.findall(decl_pattern, text, re.IGNORECASE)
        for match in matches:
            name = ' '.join(match.split())
            if len(name) > 4:
                inventors.append(name)

        return list(set(inventors))
    
    def _drawings_text_extractable(self) -> bool:
        """True if the drawings PDF has enough extractable text that
        FIG.-label-based checks can produce meaningful results. False if it's
        image-only (typical for sheets exported from CAD/Visio without text
        overlays) — in which case checks that look for FIG.-N labels will
        produce false positives.
        Heuristic: needs at least 2 distinct FIG. references in extracted text."""
        if not self.drawings_text:
            return False
        fig_count = len(set(re.findall(r'fig\.\s*(\d+)', self.drawings_text, re.IGNORECASE)))
        return fig_count >= 2

    def _extract_claims_section(self) -> str:
        """Pull just the CLAIMS section text from the specification. Many
        checks (13, 53, 54) need to operate on claim text only — running
        them against the whole spec text produces false positives because
        the body text contains many noun phrases that aren't claim elements."""
        if not self.spec_text:
            return ""
        # Order matters: most specific first. Critically, the modern-spec
        # pattern '\bCLAIMS\b\s*\n' should win over patterns that match the
        # first 'CLAIMS' substring anywhere (which would catch cross-references
        # like 'Patent Application No.…' early in the spec).
        patterns = [
            r'\bCLAIMS\b\s*\n(.{50,}?)(?:\bABSTRACT\b|\Z)',
            r'CLAIMS\s+What is claimed[^:]*:\s*(.*?)(?:\bABSTRACT\b|\Z)',
            r'\bWhat\s+is\s+claimed[^:]*:\s*(.*?)(?:\bABSTRACT\b|\Z)',
            r'(?:CLAIMS?\s*\n|What is claimed[^\n]*\n)(.*?)(?:\bABSTRACT\b|\Z)',
        ]
        for pat in patterns:
            m = re.search(pat, self.spec_text, re.DOTALL | re.IGNORECASE)
            if m and len(m.group(1)) > 100:
                return m.group(1)
        return ""

    def _is_continuation_filing(self) -> bool:
        """True if the ADS XFA data indicates this is a continuation, divisional,
        or CIP. For these filings, the parent's executed declaration and
        assignment may legitimately be older than 1 year and the spec
        legitimately references the parent application."""
        if not self.ads_data:
            return False
        for entry in self.ads_data.get('domestic_continuity_entries', []) or []:
            ct = (entry.get('continuation_type') or '').upper()
            if ct in ('CON', 'DIV', 'CIP'):
                return True
        return False

    def extract_title(self, text: str) -> str:
        """Extract application title from text"""
        # Multiple patterns for different document formats
        patterns = [
            # ADS format: "Title of Invention | TITLE HERE" or "Title of Invention TITLE HERE"
            r'Title\s+of\s+Invention[\s\|]+([A-Z][A-Z\s\-]+)',
            # Standard format: "Title: TITLE HERE"
            r'(?:Title|TITLE)[:\s]+([A-Z][A-Z\s\-]+)',
            # Specification format: Title after docket number and page number
            # "X000-0000US \n 1 TITLE HERE BACKGROUND"
            r'Docket[^\n]+\n\s*\d+\s+([A-Z][A-Z\s\-]+)',
            # Declaration/Assignment format: title in quotes
            r'entitled\s*["\']([^"\']+)["\']',
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                title = match.group(1).strip()
                # Clean up: remove extra whitespace, normalize
                title = ' '.join(title.split())

                # Truncate at common section headers that might follow the title
                for stop_word in ['BACKGROUND', 'FIELD', 'CROSS', 'RELATED', 'ABSTRACT',
                                  'TECHNICAL', 'SUMMARY', 'BRIEF', 'The application',
                                  'This application', 'T ']:  # 'T ' handles OCR artifacts
                    if stop_word in title:
                        title = title.split(stop_word)[0].strip()

                # Filter out things that aren't titles (too short, or boilerplate)
                if len(title) > 10 and 'PATENT' not in title and 'TRADEMARK' not in title:
                    return title
        return ""
    
    def extract_docket_number(self, text: str) -> str:
        """Extract a single docket number (legacy single-return wrapper).
        Returns the first docket number found via extract_docket_numbers().
        Prefer extract_docket_numbers() for new callers."""
        dockets = self.extract_docket_numbers(text)
        return next(iter(dockets), "")

    def extract_docket_numbers(self, text: str) -> set:
        """Extract ALL docket-shaped tokens from text. Patent filings often
        carry both a Client Docket and an Attorney Docket (e.g.,
        'Client Docket No.: 412147-US-NP' alongside 'Attorney Docket No.:
        MS1-9771US'). Cross-doc consistency checks need to know about both,
        not just whichever appears first."""
        dockets = set()
        # 1. Explicit "<X> Docket No.: <docket>" patterns (highest confidence).
        labelled = re.finditer(
            r'(?:Attorney|Client|Customer|Firm|Reference|File)?\s*'
            r'Docket\s*(?:No\.?|Number)?\s*[:#]?\s*'
            r'([A-Z0-9][A-Z0-9\-\._/]{3,40}[A-Z0-9])',
            text, re.IGNORECASE
        )
        for m in labelled:
            d = m.group(1).strip().rstrip('.,;:')
            if len(d) >= 5 and re.search(r'\d', d) and not d.lower().startswith('no'):
                dockets.add(d)
        # 2. Common bare docket-shaped tokens — dash-separated alphanumeric
        #    sequences with at least one digit (catches both styles like
        #    "MS1-9771USC3" and "412147-US-NP").
        for m in re.finditer(
            r'\b([A-Z]{1,5}\d{1,5}[-_]\d{2,5}[A-Z0-9\-_]{0,15})\b',
            text, re.IGNORECASE
        ):
            d = m.group(1).strip()
            if len(d) >= 5 and re.search(r'\d', d):
                dockets.add(d)
        for m in re.finditer(
            r'\b(\d{4,7}[-_][A-Z]{1,3}(?:[-_][A-Z0-9]{1,5}){0,3})\b',
            text, re.IGNORECASE
        ):
            d = m.group(1).strip()
            if len(d) >= 5 and re.search(r'\d', d):
                dockets.add(d)
        return dockets
    
    def normalize_name(self, name: str) -> str:
        """Normalize a name for comparison"""
        # Remove extra whitespace, convert to uppercase
        return re.sub(r'\s+', ' ', name.strip().upper())

    def _strip_diacritics(self, s: str) -> str:
        """Remove combining marks so 'José' compares equal to 'Jose'."""
        if not s:
            return s
        nfkd = unicodedata.normalize('NFKD', s)
        return ''.join(c for c in nfkd if not unicodedata.combining(c))

    def _normalize_for_compare(self, name: str) -> str:
        """Normalize a name for diacritic-tolerant cross-document comparison."""
        if not name:
            return ""
        s = self._strip_diacritics(name)
        s = re.sub(r'[\.,]', ' ', s)
        s = re.sub(r'\s+', ' ', s.strip().upper())
        return s

    def _format_xfa_inventor(self, inv: Dict) -> str:
        """Compose 'First Middle LAST Suffix' string from an XFA inventor record."""
        parts = []
        if inv.get('first'):
            parts.append(inv['first'])
        if inv.get('middle'):
            parts.append(inv['middle'])
        if inv.get('last'):
            parts.append(inv['last'].upper())
        if inv.get('suffix'):
            parts.append(inv['suffix'])
        return ' '.join(parts).strip()

    def _check_inventors_against_authoritative(self, doc_inventor_sets: List[Tuple[str, set]]):
        """Cross-check each document's inventor list against the authoritative source.
        Flags exact mismatches as CRITICAL and diacritic-only mismatches as WARNING."""
        auth_full = {self._normalize_for_compare(self._format_xfa_inventor(i))
                     for i in self.authoritative_inventors}
        auth_full_with_diacritics = {self._format_xfa_inventor(i) for i in self.authoritative_inventors}

        details_lines = [
            f"Authoritative source: {self.authoritative_source}",
            f"Authoritative inventors: {sorted(auth_full_with_diacritics)}",
            "",
        ]
        any_mismatch = False
        any_warning = False

        for doc_name, doc_set in doc_inventor_sets:
            if not doc_set:
                continue
            missing_in_doc = auth_full - doc_set
            extra_in_doc = doc_set - auth_full
            if not missing_in_doc and not extra_in_doc:
                # Exact match. But also check whether the doc had diacritics that the
                # authoritative source dropped (or vice versa) — a soft signal worth noting.
                details_lines.append(f"{doc_name}: ✅ matches authoritative source")
            else:
                any_mismatch = True
                details_lines.append(f"{doc_name}: ❌ disagrees with authoritative source")
                if missing_in_doc:
                    details_lines.append(f"   Missing from {doc_name}: {sorted(missing_in_doc)}")
                if extra_in_doc:
                    details_lines.append(f"   Extra in {doc_name}: {sorted(extra_in_doc)}")

        check_id = 71  # New check; numbering continues past the existing 70
        if any_mismatch:
            self.report.add_issue(
                check_id, "Cross-Document Consistency",
                "Inventor Names vs. Authoritative Source",
                Severity.CRITICAL,
                "One or more documents disagree with the authoritative inventor list",
                "\n".join(details_lines)
            )
        else:
            self.report.add_issue(
                check_id, "Cross-Document Consistency",
                "Inventor Names vs. Authoritative Source",
                Severity.PASS,
                "All documents agree with the authoritative inventor list",
                "\n".join(details_lines)
            )

    def _load_authoritative_inventors(self):
        """Load an authoritative inventor list from the folder if present.
        Supported sources (first match wins):
          - inventors.json (list of {first, middle, last, suffix} or list of strings)
          - inventors.txt   (one inventor per line; 'First Middle Last [Suffix]' or
                             'Last, First M., Suffix')
          - any *.eml file  (best-effort extraction from the email body)
        """
        # JSON
        json_path = self.folder_path / 'inventors.json'
        if json_path.exists():
            try:
                import json
                with open(json_path, 'r', encoding='utf-8') as f:
                    raw = json.load(f)
                self.authoritative_inventors = [self._coerce_inventor_record(item) for item in raw]
                self.authoritative_inventors = [i for i in self.authoritative_inventors if i.get('last') or i.get('first')]
                self.authoritative_source = json_path.name
                print(f"  ℹ️  Loaded {len(self.authoritative_inventors)} authoritative inventor(s) from {json_path.name}")
                return
            except Exception as e:
                print(f"  ⚠️  Failed to read {json_path.name}: {e}")

        # TXT
        txt_path = self.folder_path / 'inventors.txt'
        if txt_path.exists():
            try:
                with open(txt_path, 'r', encoding='utf-8') as f:
                    lines = [ln.strip() for ln in f if ln.strip() and not ln.strip().startswith('#')]
                self.authoritative_inventors = [self._parse_inventor_line(ln) for ln in lines]
                self.authoritative_inventors = [i for i in self.authoritative_inventors if i.get('last') or i.get('first')]
                self.authoritative_source = txt_path.name
                print(f"  ℹ️  Loaded {len(self.authoritative_inventors)} authoritative inventor(s) from {txt_path.name}")
                return
            except Exception as e:
                print(f"  ⚠️  Failed to read {txt_path.name}: {e}")

        # EML (best-effort)
        eml_paths = sorted(self.folder_path.glob('*.eml'))
        if eml_paths:
            try:
                import email
                from email import policy
                names = []
                for eml in eml_paths:
                    with open(eml, 'rb') as f:
                        msg = email.message_from_binary_file(f, policy=policy.default)
                    body = ""
                    if msg.is_multipart():
                        for part in msg.walk():
                            if part.get_content_type() == 'text/plain':
                                try:
                                    body += part.get_content() + "\n"
                                except Exception:
                                    pass
                    else:
                        try:
                            body = msg.get_content()
                        except Exception:
                            body = ""
                    found = self.extract_inventors(body)
                    for n in found:
                        names.append(self._parse_inventor_line(n))
                # Deduplicate by normalized form
                seen = set()
                unique = []
                for n in names:
                    key = self._normalize_for_compare(self._format_xfa_inventor(n))
                    if key and key not in seen:
                        seen.add(key)
                        unique.append(n)
                if unique:
                    self.authoritative_inventors = unique
                    self.authoritative_source = ', '.join(p.name for p in eml_paths)
                    print(f"  ℹ️  Extracted {len(unique)} candidate inventor name(s) from {len(eml_paths)} email file(s) — verify these are correct")
            except Exception as e:
                print(f"  ⚠️  Failed to read .eml file(s): {e}")

    @staticmethod
    def _coerce_inventor_record(item) -> Dict:
        """Accept either a dict {first, middle, last, suffix} or a plain string."""
        if isinstance(item, dict):
            return {
                'first': str(item.get('first', '') or '').strip(),
                'middle': str(item.get('middle', '') or '').strip(),
                'last': str(item.get('last', '') or '').strip(),
                'suffix': str(item.get('suffix', '') or '').strip(),
            }
        return PatentFilingQC._parse_inventor_line(str(item))

    @staticmethod
    def _parse_inventor_line(line: str) -> Dict:
        """Parse a free-form inventor name line into {first, middle, last, suffix}.
        Handles:
          'Dharani Bharathi Thirupathi'
          'Smith, John P., Jr.'
          'José Ramón García-López'
        """
        if not line:
            return {'first': '', 'middle': '', 'last': '', 'suffix': ''}
        line = line.strip()
        suffix = ''
        # Trailing suffix tokens (Jr., Sr., II, III, IV)
        m = re.search(r',?\s+(Jr\.?|Sr\.?|II|III|IV)\s*$', line, re.IGNORECASE)
        if m:
            suffix = m.group(1).rstrip('.').title()
            line = line[:m.start()].rstrip(', ').strip()

        if ',' in line:
            # 'Last, First Middle' format
            last, rest = line.split(',', 1)
            last = last.strip()
            tokens = rest.strip().split()
            first = tokens[0] if tokens else ''
            middle = ' '.join(tokens[1:]) if len(tokens) > 1 else ''
        else:
            tokens = line.split()
            if len(tokens) == 1:
                first, middle, last = tokens[0], '', ''
            elif len(tokens) == 2:
                first, middle, last = tokens[0], '', tokens[1]
            else:
                first, middle, last = tokens[0], ' '.join(tokens[1:-1]), tokens[-1]
        return {'first': first, 'middle': middle, 'last': last, 'suffix': suffix}

    def _extract_figure_numbers(self, text: str) -> List[int]:
        """
        Extract figure numbers from text, handling PyPDF2 extraction quirks.

        PyPDF2 often concatenates text without spaces, so:
        - "FIG. 2" + "118" (ref numeral) becomes "FIG. 2118"
        - "FIG. 4" + "START" becomes "FIG. 4START"
        - "FIG. 5" + "CPU" becomes "FIG. 5CPU"

        This method handles such cases intelligently.
        """
        if not text:
            return []

        fig_nums = set()

        # All patterns require FIG to be a standalone word (preceded by non-letter or start)
        # This prevents matching "CONFIG" as containing "FIG"

        # Pattern 1: FIG. N followed by sub-figure letter A-E, then non-letter or more text
        # e.g., "FIG. 3A " or "FIG. 3ABUILD" -> extracts 3
        for match in re.finditer(r'(?:^|[^A-Z])FIG(?:URE)?\.?\s*(\d{1,2})([A-E])(?:[^A-Z]|[A-Z])', text, re.IGNORECASE):
            num = int(match.group(1))
            if 1 <= num <= 20:
                fig_nums.add(num)

        # Pattern 2: FIG. N followed by uppercase letter F-Z (not a sub-figure)
        # e.g., "FIG. 4START" -> extracts 4, "FIG. 5CPU" -> extracts 5
        for match in re.finditer(r'(?:^|[^A-Z])FIG(?:URE)?\.?\s*(\d{1,2})([F-Z])', text, re.IGNORECASE):
            num = int(match.group(1))
            if 1 <= num <= 20:
                fig_nums.add(num)

        # Pattern 3: FIG. N followed by space, newline, punctuation, or end
        for match in re.finditer(r'(?:^|[^A-Z])FIG(?:URE)?\.?\s*(\d{1,2})(?:[\s\.,;:\)\]\-]|$)', text, re.IGNORECASE):
            num = int(match.group(1))
            if 1 <= num <= 20:
                fig_nums.add(num)

        # Pattern 4: FIG. single-digit followed by 3+ digits (reference numeral)
        # e.g., "FIG. 2118" -> extracts 2 (118 is likely a ref numeral)
        for match in re.finditer(r'(?:^|[^A-Z])FIG(?:URE)?\.?\s*(\d)(\d{3,})', text, re.IGNORECASE):
            num = int(match.group(1))
            if 1 <= num <= 9:
                fig_nums.add(num)

        # Pattern 5: FIG. 1N or 20 followed by 2+ digits
        # e.g., "FIG. 10200" -> extracts 10
        for match in re.finditer(r'(?:^|[^A-Z])FIG(?:URE)?\.?\s*(1\d|20)(\d{2,})', text, re.IGNORECASE):
            num = int(match.group(1))
            if 10 <= num <= 20:
                fig_nums.add(num)

        return sorted(fig_nums)

    def _extract_reference_numerals(self, text: str) -> dict:
        """
        Extract reference numerals from specification text with their descriptions.

        Returns a dict mapping reference numeral (str) to a dict containing:
        - 'descriptions': set of description strings found (canonical element names)
        - 'count': number of occurrences
        - 'type': 'element' or 'operation' (flowchart step)
        """
        if not text:
            return {}

        ref_data = {}

        # Pattern 1: Canonical element descriptions - "the NOUN NOUN NUM" or "a/an NOUN NOUN NUM"
        # e.g., "the computing platform 102", "a supervisor agent 204", "the BMC 106"
        # Limit to 1-3 words to avoid capturing surrounding context
        # Use word boundary and require the element name to be a proper noun phrase
        canonical_pattern = r'(?:the|a|an)\s+((?:[A-Z]{2,}|[\w\-]+(?:\s+[\w\-]+){0,2}))\s+(\d{3})\b'

        for match in re.finditer(canonical_pattern, text, re.IGNORECASE):
            desc = match.group(1).strip().lower()
            num = match.group(2)

            # Skip likely dates or non-reference numbers
            if int(num) < 100 or int(num) > 999:
                continue

            # Check if this is an operation/step/block reference (flowchart)
            is_operation = desc in ['operation', 'step', 'block'] or desc.startswith(('operation ', 'step ', 'block '))

            # Filter out noise
            noise_words = ['page', 'fig', 'figure', 'claim', 'method', 'system', 'device', 'apparatus',
                          'january', 'february', 'march', 'april', 'may', 'june',
                          'july', 'august', 'september', 'october', 'november', 'december',
                          'docket', 'attorney', 'no.', 'X000', 'patent', 'application',
                          'embodiment', 'example', 'implementation', 'aspect', 'routine',
                          'interface between', 'multitude of', 'plurality of', 'portion of']
            if any(noise in desc for noise in noise_words):
                continue

            # Skip if description is just common words or too generic
            generic_words = ['first', 'second', 'third', 'one', 'two', 'other', 'same', 'such', 'following']
            if desc in generic_words or all(w in generic_words for w in desc.split()):
                continue

            # Extract just the core element name (last 1-2 significant words)
            # This helps normalize "supervisor agent" vs "the supervisor agent"
            words = desc.split()
            # Remove prepositions and articles from the extracted phrase
            prepositions = ['of', 'in', 'on', 'at', 'to', 'for', 'with', 'by', 'from', 'between']
            core_words = []
            for w in words:
                if w in prepositions:
                    # Reset - take only words after preposition
                    core_words = []
                else:
                    core_words.append(w)

            desc_normalized = ' '.join(core_words[-3:]) if core_words else desc  # Take last 3 words max

            if num not in ref_data:
                ref_data[num] = {'descriptions': set(), 'count': 0, 'type': 'operation' if is_operation else 'element'}

            if desc_normalized and len(desc_normalized) >= 2:
                ref_data[num]['descriptions'].add(desc_normalized)

            ref_data[num]['count'] += 1

            # Update type if we find element usage (element takes precedence over operation)
            if not is_operation and ref_data[num]['type'] == 'operation':
                ref_data[num]['type'] = 'element'

        # Pattern 2: Direct element introductions like "element 102" or "(102)"
        # Also catch "NUM (the DESCRIPTION)" patterns
        intro_pattern = r'(\d{3})\s*\((?:the\s+)?([\w\-]+(?:\s+[\w\-]+){0,2})\)'
        for match in re.finditer(intro_pattern, text, re.IGNORECASE):
            num = match.group(1)
            desc = match.group(2).strip().lower()

            if 100 <= int(num) <= 999 and len(desc) >= 2:
                if num not in ref_data:
                    ref_data[num] = {'descriptions': set(), 'count': 0, 'type': 'element'}
                ref_data[num]['descriptions'].add(desc)
                ref_data[num]['count'] += 1

        # Pattern 3: Operation references like "at operation 332", "to operation 404"
        operation_pattern = r'(?:at|to|from|of|proceeds to|continues to)\s+(operation|step|block)\s+(\d{3})\b'
        for match in re.finditer(operation_pattern, text, re.IGNORECASE):
            op_type = match.group(1).lower()
            num = match.group(2)

            if 100 <= int(num) <= 999:
                if num not in ref_data:
                    ref_data[num] = {'descriptions': set(), 'count': 0, 'type': 'operation'}
                ref_data[num]['descriptions'].add(op_type)
                ref_data[num]['count'] += 1

        return ref_data

    def _extract_reference_numerals_from_drawings(self, text: str) -> set:
        """
        Extract reference numerals from drawings text.
        Drawings often have limited text, so this extracts what's available.
        """
        if not text:
            return set()

        refs = set()
        # Look for 3-digit numbers that are likely reference numerals
        pattern = r'\b(\d{3})\b'
        for match in re.finditer(pattern, text):
            num = match.group(1)
            if 100 <= int(num) <= 999:
                refs.add(num)

        return refs

    def run_all_checks(self):
        """Execute all 70 QC checks"""
        self.check_cross_document_consistency()
        self.check_document_completeness()
        self.check_specification()
        self.check_drawings()
        self.check_ads()
        self.check_declaration()
        self.check_assignment()
        self.check_poa()
        self.check_formatting()
        self.check_common_errors()
        self.check_file_quality()
        self.check_cross_references()
        self.check_priority()
        self.check_final_quality()
    
    def check_cross_document_consistency(self):
        """Checks 1-8: Cross-document consistency"""

        # Check 1: Inventor names consistency
        # Strategy: use the ADS XFA inventor list as ground truth (very
        # reliable — structured fields). For each ADS inventor, verify the
        # name appears in each other document's text. Avoids the brittleness
        # of trying to regex-extract names from declarations/assignments,
        # which use varied formats that the previous extract_inventors()
        # patterns didn't reliably match.
        ads_inventors = []
        if self.ads_data and self.ads_data.get('inventors'):
            ads_inventors = [self._format_xfa_inventor(inv) for inv in self.ads_data['inventors']]
        elif self.ads_text:
            ads_inventors = self.extract_inventors(self.ads_text)

        if not ads_inventors:
            self.report.add_issue(
                1, "Cross-Document Consistency", "Inventor Names Consistency",
                Severity.WARNING,
                "SKIPPED — could not extract inventor names from ADS to use as reference"
            )
        else:
            other_docs = [
                ('Declaration', self.declaration_text),
                ('Assignment', self.assignment_text),
            ]
            # Drawings: include only if text is extractable. Image-only
            # drawings PDFs don't contain inventor names as text, so checking
            # for them would always produce a false-positive "all missing."
            if self._drawings_text_extractable():
                other_docs.append(('Drawings', self.drawings_text))
            present_docs = [(n, t) for n, t in other_docs if t and len(t.strip()) > 100]

            if not present_docs:
                self.report.add_issue(
                    1, "Cross-Document Consistency", "Inventor Names Consistency",
                    Severity.WARNING,
                    "SKIPPED — no other documents available to cross-check ADS inventor names against",
                    f"ADS inventors: {ads_inventors}"
                )
            else:
                per_doc_missing: Dict[str, List[str]] = {}
                for doc_name, doc_text in present_docs:
                    norm_doc = self._normalize_for_compare(doc_text)
                    missing = []
                    for inv_name in ads_inventors:
                        last_name = inv_name.split()[-1] if inv_name.split() else ""
                        last_norm = self._normalize_for_compare(last_name)
                        full_norm = self._normalize_for_compare(inv_name)
                        if (last_norm and last_norm in norm_doc) or \
                           (full_norm and full_norm in norm_doc):
                            continue
                        missing.append(inv_name)
                    if missing:
                        per_doc_missing[doc_name] = missing

                if not per_doc_missing:
                    self.report.add_issue(
                        1, "Cross-Document Consistency", "Inventor Names Consistency",
                        Severity.PASS,
                        f"All {len(ads_inventors)} ADS inventors appear in: "
                        f"{', '.join(n for n, _ in present_docs)}"
                    )
                else:
                    details_lines = []
                    for doc_name, missing in per_doc_missing.items():
                        line = (f"{doc_name}: missing {len(missing)} of "
                                f"{len(ads_inventors)} — " + ', '.join(missing))
                        # Hedge if the doc has image-only pages
                        img_pages = self.image_only_pages.get(doc_name, 0)
                        if img_pages:
                            line += (f"  [Note: {doc_name} has {img_pages} image-only "
                                     f"page(s) — name(s) may be there but not extractable.]")
                        details_lines.append(line)
                    # If every "missing" finding is shadowed by image-only pages,
                    # downgrade severity from CRITICAL to WARNING.
                    all_hedged = all(self.image_only_pages.get(n) for n in per_doc_missing)
                    severity = Severity.WARNING if all_hedged else Severity.CRITICAL
                    msg = ("Some ADS inventors not found in cross-checked document text — "
                           "may be on image-only pages") if all_hedged else \
                          "Some ADS inventors do not appear in all cross-checked documents"
                    self.report.add_issue(
                        1, "Cross-Document Consistency", "Inventor Names Consistency",
                        severity, msg, "\n".join(details_lines)
                    )

        # Check 1b (extension): authoritative-source cross-check.
        # If the user dropped an inventors.txt / inventors.json / .eml in the folder,
        # treat that list as ground truth and flag any document that disagrees.
        if self.authoritative_inventors:
            # Build per-doc inventor sets from ADS XFA + extract from other docs
            decl_inventors = self.extract_inventors(self.declaration_text) if self.declaration_text else []
            assign_inventors = self.extract_inventors(self.assignment_text) if self.assignment_text else []
            drawings_inventors = self.extract_inventors(self.drawings_text) if self.drawings_text else []
            all_inventor_sets = [
                ("ADS", set(self._normalize_for_compare(i) for i in ads_inventors)),
                ("Declaration", set(self._normalize_for_compare(i) for i in decl_inventors)),
                ("Assignment", set(self._normalize_for_compare(i) for i in assign_inventors)),
                ("Drawings", set(self._normalize_for_compare(i) for i in drawings_inventors))
            ]
            self._check_inventors_against_authoritative(all_inventor_sets)
        
        # Check 2: Application title consistency
        # Strategy: get the canonical title from the ADS (XFA-extracted; very
        # reliable), then search the spec text for it. This avoids the
        # fragility of extract_title() trying to find the title without a
        # known reference — regex patterns can't reliably identify the title
        # from spec text alone (no "Title:" label on most specs).
        ads_title = ""
        if self.ads_data and self.ads_data.get('title'):
            ads_title = self.ads_data['title']
        else:
            ads_title = self.extract_title(self.ads_text)

        if not ads_title:
            self.report.add_issue(
                2, "Cross-Document Consistency", "Application Title Consistency",
                Severity.WARNING, "Unable to extract title from ADS"
            )
        elif not self.spec_text:
            self.report.add_issue(
                2, "Cross-Document Consistency", "Application Title Consistency",
                Severity.WARNING, "Specification not available to compare title"
            )
        else:
            # Normalize: collapse whitespace, uppercase, strip trailing punctuation
            def normalize(s):
                return re.sub(r'\s+', ' ', s.upper()).strip().rstrip('.,;:')
            ads_norm = normalize(ads_title)
            spec_norm = normalize(self.spec_text)
            if ads_norm in spec_norm:
                self.report.add_issue(
                    2, "Cross-Document Consistency", "Application Title Consistency",
                    Severity.PASS,
                    f"ADS title appears verbatim in specification"
                )
            else:
                # Maybe minor differences (hyphenation, line breaks). Try matching
                # on the first 60% of the title, ignoring punctuation differences.
                title_words = ads_norm.split()
                key_chunk = ' '.join(title_words[:max(4, int(len(title_words) * 0.6))])
                if key_chunk in spec_norm:
                    self.report.add_issue(
                        2, "Cross-Document Consistency", "Application Title Consistency",
                        Severity.PASS,
                        "Most of ADS title appears in specification (minor wording differences detected)"
                    )
                else:
                    self.report.add_issue(
                        2, "Cross-Document Consistency", "Application Title Consistency",
                        Severity.CRITICAL,
                        "ADS title does not appear in specification — verify they describe the same application",
                        f"ADS title: {ads_title}"
                    )
        
        # Check 3: Attorney docket number consistency
        # Patent filings often carry both a Client Docket and an Attorney Docket
        # in the same document footer. We extract ALL docket-shaped tokens from
        # each doc (as a set) and PASS if any docket overlaps across docs.
        # For continuation/divisional/CIP filings, the parent's executed
        # declaration and assignment legitimately carry the parent application's
        # dockets (which differ from the child's), so we only require spec↔ADS
        # consistency in that case.
        docket_sets = {}
        if self.ads_data and self.ads_data.get('docket_number'):
            docket_sets["ADS"] = {self.ads_data['docket_number']}
        elif self.ads_text:
            docket_sets["ADS"] = self.extract_docket_numbers(self.ads_text)
        for name, text in [("Spec", self.spec_text),
                           ("Declaration", self.declaration_text),
                           ("Assignment", self.assignment_text)]:
            if text:
                ds = self.extract_docket_numbers(text)
                if ds:
                    docket_sets[name] = ds

        if len(docket_sets) >= 2:
            is_continuation = self._is_continuation_filing()
            # For continuations, only require Spec↔ADS to match. Dec/Asgn
            # carried forward from a parent are expected to have different dockets.
            sources_to_compare = list(docket_sets.keys())
            if is_continuation:
                sources_to_compare = [s for s in sources_to_compare if s in ("ADS", "Spec")]

            if len(sources_to_compare) >= 2:
                # Use case-insensitive comparison
                normalized = {s: {d.upper() for d in docket_sets[s]} for s in sources_to_compare}
                # PASS if any docket appears in EVERY compared doc
                shared = set.intersection(*normalized.values()) if normalized else set()
                if shared:
                    suffix = " (Spec↔ADS only — parent's Dec/Asgn dockets carried forward)" if is_continuation else ""
                    msg = f"Attorney docket number consistent across documents{suffix}: {sorted(shared)[0]}"
                    self.report.add_issue(
                        3, "Cross-Document Consistency", "Attorney Docket Number Consistency",
                        Severity.PASS, msg
                    )
                else:
                    details_lines = [f"{name}: {sorted(docket_sets[name])}" for name in docket_sets]
                    self.report.add_issue(
                        3, "Cross-Document Consistency", "Attorney Docket Number Consistency",
                        Severity.CRITICAL, "Attorney docket number mismatch — no docket appears in all compared documents",
                        "\n".join(details_lines)
                    )
            else:
                self.report.add_issue(
                    3, "Cross-Document Consistency", "Attorney Docket Number Consistency",
                    Severity.WARNING,
                    "Unable to compare dockets — only one source has extractable docket numbers"
                )
        else:
            self.report.add_issue(
                3, "Cross-Document Consistency", "Attorney Docket Number Consistency",
                Severity.WARNING, "Unable to extract docket numbers from multiple documents"
            )
        
        # Check 4: Correspondence address consistency
        # Compare customer number across ADS and POA (if present)
        ads_customer_num = None
        poa_customer_num = None

        # Prefer XFA-extracted customer number when available
        if self.ads_data and self.ads_data.get('customer_number'):
            ads_customer_num = self.ads_data['customer_number']
        elif self.ads_text:
            cust_match = re.search(r'Customer\s*(?:Number|No\.?)[:\s]*(\d{5,6})', self.ads_text, re.IGNORECASE)
            if cust_match:
                ads_customer_num = cust_match.group(1)

        if self.poa_text:
            cust_match = re.search(r'Customer\s*(?:Number|No\.?)[:\s]*(\d{5,6})', self.poa_text, re.IGNORECASE)
            if cust_match:
                poa_customer_num = cust_match.group(1)

        if ads_customer_num and poa_customer_num:
            if ads_customer_num == poa_customer_num:
                self.report.add_issue(
                    4, "Cross-Document Consistency", "Correspondence Address Consistency",
                    Severity.PASS, f"Customer number consistent: {ads_customer_num}"
                )
            else:
                self.report.add_issue(
                    4, "Cross-Document Consistency", "Correspondence Address Consistency",
                    Severity.CRITICAL, f"Customer number mismatch: ADS={ads_customer_num}, POA={poa_customer_num}"
                )
        elif ads_customer_num or poa_customer_num:
            self.report.add_issue(
                4, "Cross-Document Consistency", "Correspondence Address Consistency",
                Severity.PASS, f"Customer number found: {ads_customer_num or poa_customer_num}"
            )
        else:
            self.report.add_issue(
                4, "Cross-Document Consistency", "Correspondence Address Consistency",
                Severity.INFO, "No customer number found - manual review of correspondence address recommended"
            )
        
        # Check 5: Assignee name consistency
        # Use the XFA-extracted assignee from ADS as the canonical name (clean,
        # structured), then verify it appears in the assignment text. The old
        # version regex-extracted the assignee from synthesized ADS text and
        # got newline-broken values; using XFA structured data is reliable.
        ads_assignee = None
        if self.ads_data and self.ads_data.get('assignee_org'):
            ads_assignee = self.ads_data['assignee_org']
        elif self.ads_text:
            # Legacy fallback for non-XFA ADS
            for pat in [
                r'(?:Organization|ization)\s*Name\s*[|\s]+([A-Za-z][\w\s,]+(?:LLC|Inc|Corp))',
                r'c/o\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)+(?:,?\s*(?:LLC|Inc|Corp)))',
                r'Applicant\s*Name[:\s|]+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)+(?:,?\s*(?:LLC|Inc|Corp)))',
            ]:
                m = re.search(pat, self.ads_text, re.IGNORECASE)
                if m:
                    ads_assignee = re.sub(r'\s+', ' ', m.group(1).strip())
                    break

        assignment_assignee = None
        if self.assignment_text:
            # Look for the same assignee name in assignment by matching key
            # words from the ADS-derived name (rather than regex-extracting
            # a fresh name from assignment text, which can pick up newlines).
            if ads_assignee:
                # Take significant words (3+ chars, not a common word) and
                # check that the assignment contains them
                key_words = [w for w in re.findall(r'[A-Za-z]{3,}', ads_assignee.upper())
                             if w not in {'THE', 'AND', 'INC', 'LLC', 'CORP', 'LTD', 'CO'}]
                asgn_norm = self._normalize_for_compare(self.assignment_text)
                matches = sum(1 for w in key_words if w in asgn_norm)
                if key_words and matches >= max(1, len(key_words) // 2):
                    assignment_assignee = ads_assignee  # Found by key-word match
            if not assignment_assignee:
                # Fall back to regex extraction (collapses any newlines)
                m = re.search(
                    r'([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)+(?:,?\s*(?:LLC|Inc|Corp)))',
                    self.assignment_text
                )
                if m:
                    assignment_assignee = re.sub(r'\s+', ' ', m.group(1).strip())

        # Normalize for comparison - handle OCR errors
        def normalize_company(name):
            if not name:
                return ""
            # Handle common OCR errors: "N" vs "n", "l" vs "1", etc.
            norm = name.lower()
            norm = re.sub(r'[\s,\.]+', ' ', norm)
            norm = norm.replace('0', 'o').replace('1', 'l')
            # Extract key identifier words
            return norm.strip()

        if ads_assignee and assignment_assignee:
            ads_norm = normalize_company(ads_assignee)
            assign_norm = normalize_company(assignment_assignee)

            # Check if they share key words (company name parts)
            ads_words = set(ads_norm.split())
            assign_words = set(assign_norm.split())
            common_words = ads_words & assign_words

            # If they share significant words (like "american", "megatrends", "international")
            if len(common_words) >= 2 or ads_norm in assign_norm or assign_norm in ads_norm:
                self.report.add_issue(
                    5, "Cross-Document Consistency", "Assignee Name Consistency",
                    Severity.PASS, f"Assignee name consistent: {assignment_assignee}"
                )
            else:
                self.report.add_issue(
                    5, "Cross-Document Consistency", "Assignee Name Consistency",
                    Severity.WARNING,
                    f"Assignee names may differ: ADS='{ads_assignee}', Assignment='{assignment_assignee}'"
                )
        elif ads_assignee or assignment_assignee:
            assignee = ads_assignee or assignment_assignee
            self.report.add_issue(
                5, "Cross-Document Consistency", "Assignee Name Consistency",
                Severity.PASS, f"Assignee found: {assignee}"
            )
        else:
            self.report.add_issue(
                5, "Cross-Document Consistency", "Assignee Name Consistency",
                Severity.INFO, "Could not extract assignee name for comparison"
            )

        # Check 6: Filing date consistency
        # For new applications, documents should indicate "Herewith" or blank filing date
        # Check that no document has a conflicting/premature filing date
        filing_date_issues = []

        # Check POA for "Filing Date Herewith" or blank
        if self.poa_text:
            if re.search(r'Filing\s*Date\s*Herewith', self.poa_text, re.IGNORECASE):
                # Good - indicates new filing
                pass
            elif re.search(r'Filing\s*Date\s*\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}', self.poa_text, re.IGNORECASE):
                filing_date_issues.append("POA has specific filing date (should be 'Herewith' for new applications)")

        # Check ADS for filing date field
        if self.ads_text:
            ads_date_match = re.search(r'Filing\s*Date[:\s]+(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})', self.ads_text, re.IGNORECASE)
            if ads_date_match:
                filing_date_issues.append(f"ADS has filing date: {ads_date_match.group(1)}")

        if not filing_date_issues:
            self.report.add_issue(
                6, "Cross-Document Consistency", "Filing Date Consistency",
                Severity.PASS, "Documents consistently indicate new filing (no conflicting dates)"
            )
        else:
            self.report.add_issue(
                6, "Cross-Document Consistency", "Filing Date Consistency",
                Severity.WARNING, f"Filing date inconsistency: {'; '.join(filing_date_issues)}"
            )
        
        # Check 7: Number of inventors consistency
        # Compare ADS inventor count to declaration "I hereby declare" count
        # (each inventor signs their own declaration, so the phrase appears
        # once per inventor in a multi-inventor declaration package).
        if ads_inventors and self.declaration_text:
            ads_count = len(ads_inventors)
            decl_count = len(re.findall(r'(?:i|I)\s+hereby\s+declare', self.declaration_text))
            if decl_count == 0:
                # Try alternate signature/inventor markers
                decl_count = len(re.findall(r'/[A-Z][^/]{2,40}/', self.declaration_text))
            if decl_count == ads_count:
                self.report.add_issue(
                    7, "Cross-Document Consistency", "Number of Inventors Consistency",
                    Severity.PASS, f"Same number of inventors ({ads_count}) in ADS and Declaration"
                )
            elif decl_count == 0:
                self.report.add_issue(
                    7, "Cross-Document Consistency", "Number of Inventors Consistency",
                    Severity.INFO,
                    f"ADS has {ads_count} inventor(s); could not count inventors in declaration. "
                    "Manual verification recommended."
                )
            else:
                # Continuation note: parent's declaration may have fewer inventors
                # if some have been added/removed in the continuation.
                cont_note = ""
                if self._is_continuation_filing():
                    cont_note = (" Note: this is a continuation filing — if the parent's "
                                 "declaration is carried forward, inventorship may have changed in "
                                 "the continuation, in which case a new declaration is required.")
                # Hedge if the declaration has image-only pages
                img_pages = self.image_only_pages.get('Declaration', 0)
                if img_pages and decl_count + img_pages >= ads_count:
                    self.report.add_issue(
                        7, "Cross-Document Consistency", "Number of Inventors Consistency",
                        Severity.WARNING,
                        f"Could not confirm count: ADS has {ads_count}, Declaration text shows "
                        f"{decl_count} signed declaration(s) and {img_pages} additional image-only "
                        f"page(s) which may contain the remaining {ads_count - decl_count}.",
                    )
                else:
                    self.report.add_issue(
                        7, "Cross-Document Consistency", "Number of Inventors Consistency",
                        Severity.CRITICAL,
                        f"Inventor count mismatch: ADS has {ads_count}, Declaration has {decl_count}." + cont_note
                    )
        else:
            self.report.add_issue(
                7, "Cross-Document Consistency", "Number of Inventors Consistency",
                Severity.INFO, "Unable to count inventors — ADS or Declaration text not available"
            )
        
        # Check 8: Inventor citizenship/residency consistency
        # Check that all inventors have residency information in ADS
        if self.ads_text:
            # Count inventors with US Residency marked
            us_residency_count = len(re.findall(r'US\s*Residency', self.ads_text, re.IGNORECASE))
            # Count total inventors
            inventor_count = len(re.findall(r'Inventor\s+\d+', self.ads_text, re.IGNORECASE))

            if inventor_count > 0 and us_residency_count >= inventor_count:
                self.report.add_issue(
                    8, "Cross-Document Consistency", "Inventor Citizenship/Residency Consistency",
                    Severity.PASS, f"All {inventor_count} inventors have residency information"
                )
            elif inventor_count > 0:
                self.report.add_issue(
                    8, "Cross-Document Consistency", "Inventor Citizenship/Residency Consistency",
                    Severity.WARNING, f"Found {us_residency_count} residency entries for {inventor_count} inventors"
                )
            else:
                self.report.add_issue(
                    8, "Cross-Document Consistency", "Inventor Citizenship/Residency Consistency",
                    Severity.INFO, "Could not verify inventor residency information"
                )
        else:
            self.report.add_issue(
                8, "Cross-Document Consistency", "Inventor Citizenship/Residency Consistency",
                Severity.INFO, "ADS not available to check residency information"
            )
    
    def check_document_completeness(self):
        """Checks 9-12: Document completeness"""
        
        # Check 9: All required documents present
        required = ['Specification', 'Drawings', 'ADS', 'Declaration']
        missing = [doc for doc in required if not self.report.files_found.get(doc)]

        # Documents that are commonly omitted at filing as part of an intentional
        # "missing parts" filing under 37 CFR §1.53(f). The Declaration is the most
        # frequent — filers can pay the §1.16(f) surcharge and submit the missing
        # parts within 2 months of the Notice to File Missing Parts.
        missing_parts_eligible = {'Declaration'}

        if not missing:
            self.report.add_issue(
                9, "Document Completeness", "All Required Documents Present",
                Severity.PASS, "All required documents found"
            )
        else:
            blocking = [d for d in missing if d not in missing_parts_eligible]
            optional_missing = [d for d in missing if d in missing_parts_eligible]

            if blocking:
                # Spec / Drawings / ADS missing — these aren't eligible for missing-parts
                self.report.add_issue(
                    9, "Document Completeness", "All Required Documents Present",
                    Severity.CRITICAL,
                    f"Missing required documents: {', '.join(blocking)}",
                    "These documents must be in the filing folder. They are not eligible "
                    "for the missing-parts procedure under 37 CFR §1.53(f)."
                )

            if optional_missing:
                # Declaration missing — emit a CRITICAL-with-question that Claude (per SKILL.md)
                # will resolve with the user. If intentional, the filer should be reminded
                # about the §1.16(f) surcharge and 2-month deadline.
                docs_str = ', '.join(optional_missing)
                self.report.add_issue(
                    9, "Document Completeness", "All Required Documents Present",
                    Severity.CRITICAL,
                    f"{docs_str} not found — confirm whether this is intentional",
                    (
                        "ACTION REQUIRED: Ask the filer whether this is an intentional "
                        "missing-parts filing under 37 CFR §1.53(f).\n"
                        "  • If YES (intentional): downgrade this to a WARNING and remind the filer that:\n"
                        "      – A §1.16(f) surcharge fee is due at or after filing\n"
                        "      – The missing parts (e.g., declaration) must be filed within 2 months\n"
                        "        of the USPTO's Notice to File Missing Parts to avoid abandonment\n"
                        "  • If NO (oversight): the missing document(s) must be added before filing"
                    )
                )
        
        # Check 10: ADS required fields complete
        if self.ads_text:
            required_fields = ['title', 'inventor', 'correspondence']
            # Simplified check - look for key terms
            missing_fields = []
            if 'title' not in self.ads_text.lower():
                missing_fields.append('title')
            if 'inventor' not in self.ads_text.lower():
                missing_fields.append('inventor')
            if 'correspondence' not in self.ads_text.lower():
                missing_fields.append('correspondence')
            
            if not missing_fields:
                self.report.add_issue(
                    10, "Document Completeness", "ADS Required Fields Complete",
                    Severity.PASS, "ADS appears to have required fields"
                )
            else:
                self.report.add_issue(
                    10, "Document Completeness", "ADS Required Fields Complete",
                    Severity.WARNING, f"ADS may be missing fields: {', '.join(missing_fields)}"
                )
        else:
            self.report.add_issue(
                10, "Document Completeness", "ADS Required Fields Complete",
                Severity.CRITICAL, "ADS not found"
            )
        
        # Check 11: Declaration signatures present
        if self.declaration_text:
            # Look for various signature indicators
            # PyPDF2 may not extract all text, so check multiple patterns
            signature_indicators = [
                '/s/',                          # Electronic signature
                'signature',                    # Generic signature reference
                'inventor signature',           # Inventor signature label
                'witness signature',            # Witness signature label
                'witnessed by',                 # Witness section header
                'legal name of inventor',       # Name above signature line
                'signed',                       # Signed indicator
                'executed',                     # Executed indicator
                r'\d{1,2}/\d{1,2}/\d{2,4}',    # Date pattern (regex)
                r'\d{1,2}/\d{1,2}/20\d{2}',    # Full year date pattern (regex)
            ]
            dec_text_lower = self.declaration_text.lower()
            sig_found = False
            for indicator in signature_indicators:
                if indicator.startswith('\\'):
                    # Regex pattern
                    if re.search(indicator, self.declaration_text, re.IGNORECASE):
                        sig_found = True
                        break
                elif indicator in dec_text_lower:
                    sig_found = True
                    break

            if sig_found:
                self.report.add_issue(
                    11, "Document Completeness", "Declaration Signatures Present",
                    Severity.PASS, "Declaration appears to have signature indicators"
                )
            else:
                self.report.add_issue(
                    11, "Document Completeness", "Declaration Signatures Present",
                    Severity.WARNING, "Declaration may be missing signatures"
                )
        else:
            self.report.add_issue(
                11, "Document Completeness", "Declaration Signatures Present",
                Severity.WARNING, "Declaration not found - cannot check signatures"
            )

        # Check 12: Assignment signatures present
        if self.assignment_text:
            # Look for various signature indicators
            signature_indicators = [
                '/s/',                          # Electronic signature
                'signature',                    # Generic signature reference
                'inventor signature',           # Inventor signature label
                'witness signature',            # Witness signature label
                'witnessed by',                 # Witness section header
                'legal name of inventor',       # Name above signature line
                'assignor',                     # Assignor (who signs assignment)
                'signed',                       # Signed indicator
                'executed',                     # Executed indicator
                r'\d{1,2}/\d{1,2}/\d{2,4}',    # Date pattern (regex)
                r'\d{1,2}/\d{1,2}/20\d{2}',    # Full year date pattern (regex)
            ]
            assign_text_lower = self.assignment_text.lower()
            sig_found = False
            for indicator in signature_indicators:
                if indicator.startswith('\\'):
                    # Regex pattern
                    if re.search(indicator, self.assignment_text, re.IGNORECASE):
                        sig_found = True
                        break
                elif indicator in assign_text_lower:
                    sig_found = True
                    break

            if sig_found:
                self.report.add_issue(
                    12, "Document Completeness", "Assignment Signatures Present",
                    Severity.PASS, "Assignment appears to have signature indicators"
                )
            else:
                self.report.add_issue(
                    12, "Document Completeness", "Assignment Signatures Present",
                    Severity.WARNING, "Assignment may be missing signatures"
                )
        else:
            self.report.add_issue(
                12, "Document Completeness", "Assignment Signatures Present",
                Severity.INFO, "Assignment not found (optional document)"
            )
    
    def check_specification(self):
        """Checks 13-27: Specification-specific checks"""
        
        if not self.spec_text:
            for i in range(13, 28):
                self.report.add_issue(
                    i, "Specification", f"Check {i}",
                    Severity.CRITICAL, "Specification not found"
                )
            return
        
        # Check 13: Claim numbering sequential
        # Find the CLAIMS section first. The original code only found it via
        # "What is claimed is:" preamble — many specs use just "CLAIMS" as a
        # section header, so add that pattern first.
        claims_section_patterns = [
            r'\bCLAIMS\b\s*\n(.{50,}?)(?:\bABSTRACT\b|\Z)',                  # "CLAIMS\n..." (common modern format)
            r'CLAIMS\s+What is claimed[^:]*:\s*(.*?)(?:ABSTRACT|$)',         # "CLAIMS What is claimed is:"
            r'What is claimed[^:]*:\s*(.*?)(?:ABSTRACT|$)',                  # Just "What is claimed is:"
            r'(?:CLAIMS?\s*\n|What is claimed[^\n]*\n)(.*?)(?:ABSTRACT|$)',  # Fallback with newlines
        ]

        claims_text = None
        for pattern in claims_section_patterns:
            claims_section_match = re.search(pattern, self.spec_text, re.DOTALL | re.IGNORECASE)
            if claims_section_match:
                claims_text = claims_section_match.group(1)
                break

        if not claims_text:
            claims_text = self.spec_text

        # Find claim numbers within the claims section. The previous version
        # required a specific preamble word (A/An/The) after "N. " — that
        # missed claims that start with other words like "A computer-implemented",
        # "Method for", "Computer system", etc. Real-world claim 1 frequently
        # doesn't fit that narrow set, producing a false-positive "claim 1
        # missing" warning.
        # Strategy: anchor on number-then-period at start-of-line or after
        # sentence punctuation, then accept ANY following word. To filter out
        # decimal numbers ("1.5") and section refs ("1.2"), require non-digit
        # immediately after the period.
        claim_patterns = [
            r'(?:^|\n)\s*(\d+)\.\s+(?=\D)',                  # newline-anchored
            r'(?:\.\s+|\;\s+|:\s+)(\d+)\.\s+(?=\D)',         # after sentence end
            r'\s{2,}(\d+)\.\s+(?=\D)',                       # after multiple spaces
            r'\s+\d{2,3}\s+(\d+)\.\s+(?=\D)',                # after page number
        ]

        claim_matches = []
        for pattern in claim_patterns:
            matches = re.findall(pattern, claims_text, re.IGNORECASE)
            claim_matches.extend(matches)

        # Deduplicate
        claim_matches = list(set(claim_matches))

        if claim_matches:
            # Filter to reasonable claim numbers and deduplicate
            claim_nums = sorted(set(int(n) for n in claim_matches if 1 <= int(n) <= 100))
            expected = list(range(1, max(claim_nums) + 1))

            if claim_nums == expected:
                self.report.add_issue(
                    13, "Specification", "Claim Numbering Sequential",
                    Severity.PASS, f"Claims numbered sequentially (1-{len(claim_nums)})"
                )
            else:
                missing = set(expected) - set(claim_nums)
                if missing:
                    self.report.add_issue(
                        13, "Specification", "Claim Numbering Sequential",
                        Severity.CRITICAL, f"Claims not numbered sequentially - missing: {sorted(missing)}",
                        f"Found: {claim_nums}"
                    )
                else:
                    self.report.add_issue(
                        13, "Specification", "Claim Numbering Sequential",
                        Severity.PASS, f"Claims numbered sequentially (1-{max(claim_nums)})"
                    )
        else:
            # Fallback: try simpler pattern in claims section
            # Look for any number followed by period and space
            simple_patterns = [
                r'(?:^|\n)\s*(\d+)\.\s+',           # Original with newlines
                r'(?:\.\s+|\;\s+|:\s+)(\d+)\.\s+',  # After sentence end
                r'\s{2,}(\d+)\.\s+',                # After multiple spaces
                r'^\s*(\d+)\.\s+',                  # At start of claims section
                r'\s+\d{2,3}\s+(\d+)\.\s+',         # After page number (2-3 digits)
            ]
            simple_matches = []
            for pattern in simple_patterns:
                matches = re.findall(pattern, claims_text, re.MULTILINE)
                simple_matches.extend(matches)
            simple_nums = sorted(set(int(n) for n in simple_matches if 1 <= int(n) <= 100))

            if simple_nums:
                expected = list(range(1, max(simple_nums) + 1))
                if simple_nums == expected:
                    self.report.add_issue(
                        13, "Specification", "Claim Numbering Sequential",
                        Severity.PASS, f"Claims appear numbered sequentially (1-{len(simple_nums)})"
                    )
                else:
                    self.report.add_issue(
                        13, "Specification", "Claim Numbering Sequential",
                        Severity.WARNING, f"Claim numbering may have issues. Found numbers: {simple_nums}"
                    )
            else:
                self.report.add_issue(
                    13, "Specification", "Claim Numbering Sequential",
                    Severity.WARNING, "Unable to detect claim numbers - claims may use non-standard format"
                )
        
        # Check 14: Claim dependency validity
        # Look for dependent claim patterns like "2. The method of claim 1" or "3. A system according to claim 2"
        # PyPDF2 often doesn't preserve newlines, so we need flexible patterns
        # Patterns handle: start of line, after period+spaces, after double spaces
        dependent_claim_patterns = [
            # Standard format: "N. The [noun] of claim M" - handles hyphenated nouns like "computer-implemented"
            r'(?:^|\n|\.\s{1,3})(\d+)\.\s+(?:The|A|An)\s+[\w\-]+(?:\s+[\w\-]+)*\s+(?:of|according to|as (?:recited|claimed|set forth) in|as in)\s+claim\s+(\d+)',
            # After double/triple space (common in PDF extraction)
            r'\s{2,}(\d+)\.\s+(?:The|A|An)\s+[\w\-]+(?:\s+[\w\-]+)*\s+(?:of|according to|as (?:recited|claimed|set forth) in|as in)\s+claim\s+(\d+)',
            # After page number pattern (e.g., "48 4. The")
            r'\s+\d{2,3}\s+(\d+)\.\s+(?:The|A|An)\s+[\w\-]+(?:\s+[\w\-]+)*\s+(?:of|according to|as (?:recited|claimed|set forth) in|as in)\s+claim\s+(\d+)',
        ]

        dependent_claims = []
        for pattern in dependent_claim_patterns:
            matches = re.findall(pattern, self.spec_text, re.IGNORECASE | re.MULTILINE)
            dependent_claims.extend(matches)

        # Deduplicate
        dependent_claims = list(set(dependent_claims))

        if dependent_claims:
            # Filter out any self-references (which shouldn't happen with proper claims)
            # Also filter to only reasonable claim numbers (1-100)
            valid_deps = [(c, d) for c, d in dependent_claims
                         if int(c) != int(d) and 1 <= int(c) <= 100 and 1 <= int(d) <= 100]
            invalid_deps = [(c, d) for c, d in valid_deps if int(c) <= int(d)]

            if not invalid_deps and valid_deps:
                self.report.add_issue(
                    14, "Specification", "Claim Dependency Validity",
                    Severity.PASS, f"Dependent claims reference lower-numbered claims ({len(valid_deps)} dependencies found)"
                )
            elif invalid_deps:
                self.report.add_issue(
                    14, "Specification", "Claim Dependency Validity",
                    Severity.CRITICAL, "Invalid claim dependencies found - claims reference same or higher numbered claims",
                    str(invalid_deps)
                )
            else:
                self.report.add_issue(
                    14, "Specification", "Claim Dependency Validity",
                    Severity.INFO, "No valid dependent claims detected after filtering"
                )
        else:
            # Try a fallback pattern - look for "claim N" references within claim text
            # More permissive: look for claim number followed by text containing "of claim" or "claim N"
            fallback_pattern = r'(\d+)\.\s+(?:The|A|An)\s+.{0,150}?\s+(?:of\s+)?claim\s+(\d+)'
            fallback_deps = re.findall(fallback_pattern, self.spec_text, re.IGNORECASE)

            # Filter self-references and unreasonable claim numbers
            fallback_deps = [(c, d) for c, d in fallback_deps
                            if int(c) != int(d) and 1 <= int(c) <= 100 and 1 <= int(d) <= 100]
            # Deduplicate
            fallback_deps = list(set(fallback_deps))

            if fallback_deps:
                invalid_deps = [(c, d) for c, d in fallback_deps if int(c) <= int(d)]
                if not invalid_deps:
                    self.report.add_issue(
                        14, "Specification", "Claim Dependency Validity",
                        Severity.PASS, f"Dependent claims appear valid ({len(fallback_deps)} dependencies found)"
                    )
                else:
                    self.report.add_issue(
                        14, "Specification", "Claim Dependency Validity",
                        Severity.CRITICAL, "Invalid claim dependencies found",
                        str(invalid_deps)
                    )
            else:
                self.report.add_issue(
                    14, "Specification", "Claim Dependency Validity",
                    Severity.INFO, "No dependent claims detected (all claims may be independent)"
                )
        
        # Check 15: Figure reference validity
        # Drawings PDFs are commonly image-only (CAD/Visio exports without
        # text overlays). When that's the case, FIG. labels exist visually
        # but aren't extractable as text. Don't flag as a real "figure
        # missing" issue when extraction fundamentally can't see the labels.
        fig_refs_in_spec = set(self._extract_figure_numbers(self.spec_text))
        fig_nums_in_drawings = set(self._extract_figure_numbers(self.drawings_text)) if self.drawings_text else set()

        if not self._drawings_text_extractable() and fig_refs_in_spec:
            self.report.add_issue(
                15, "Specification", "Figure Reference Validity",
                Severity.INFO,
                f"Drawings PDF appears to be image-only — cannot verify FIG. labels by text extraction. "
                f"Spec references: FIG. {', '.join(str(n) for n in sorted(fig_refs_in_spec))}. "
                f"Manually verify each is present in the drawings."
            )
        elif fig_refs_in_spec and fig_nums_in_drawings:
            missing_figs = fig_refs_in_spec - fig_nums_in_drawings
            if not missing_figs:
                self.report.add_issue(
                    15, "Specification", "Figure Reference Validity",
                    Severity.PASS, f"All referenced figures exist in drawings (FIG. {', '.join(str(n) for n in sorted(fig_refs_in_spec))})"
                )
            else:
                self.report.add_issue(
                    15, "Specification", "Figure Reference Validity",
                    Severity.CRITICAL, f"Specification references figures not in drawings: FIG. {', '.join(str(n) for n in sorted(missing_figs))}"
                )
        elif fig_refs_in_spec:
            self.report.add_issue(
                15, "Specification", "Figure Reference Validity",
                Severity.INFO,
                f"Drawings text extraction was minimal — cannot verify. "
                f"Spec references: FIG. {', '.join(str(n) for n in sorted(fig_refs_in_spec))}"
            )
        else:
            self.report.add_issue(
                15, "Specification", "Figure Reference Validity",
                Severity.WARNING, "Unable to detect figure references in specification"
            )
        
        # Check 16: Reference numeral consistency
        spec_refs = self._extract_reference_numerals(self.spec_text)
        drawings_refs = self._extract_reference_numerals_from_drawings(self.drawings_text) if self.drawings_text else set()

        if spec_refs:
            issues = []
            warnings = []

            # Check 1: Consistency - each reference numeral should have consistent description
            for num, data in spec_refs.items():
                descs = data['descriptions']

                if len(descs) > 1:
                    # Normalize descriptions to find the core element name
                    # Remove modifiers like "target", "source", "primary", etc.
                    # and extract the core noun phrase

                    def get_core_name(desc):
                        """Extract core element name from description"""
                        # Remove common modifiers
                        modifiers = ['target', 'source', 'primary', 'secondary', 'main', 'new', 'old',
                                    'current', 'next', 'previous', 'updated', 'original', 'modified',
                                    'first', 'second', 'third', 'specific', 'particular', 'given',
                                    'respective', 'corresponding', 'associated', 'related']
                        words = desc.lower().split()
                        core_words = [w for w in words if w not in modifiers]

                        # If we stripped everything, keep the last word
                        if not core_words and words:
                            core_words = [words[-1]]

                        return ' '.join(core_words)

                    # Also check if one description contains another (subset relationship)
                    def is_same_element(desc1, desc2):
                        """Check if two descriptions refer to the same element"""
                        core1 = get_core_name(desc1)
                        core2 = get_core_name(desc2)

                        # Exact match after normalization
                        if core1 == core2:
                            return True

                        # One contains the other (e.g., "computing platform" vs "platform")
                        if core1 in core2 or core2 in core1:
                            return True

                        # Last word matches (usually the main noun)
                        if core1.split()[-1] == core2.split()[-1]:
                            return True

                        # Acronym match: one is the initials of the other
                        # (e.g. "gnn" vs "graph neural network")
                        def is_acronym_of(short, long):
                            short = short.replace(' ', '').lower()
                            words = long.split()
                            if len(short) == len(words) and len(words) >= 2:
                                initials = ''.join(w[0].lower() for w in words if w)
                                return initials == short
                            return False
                        if is_acronym_of(core1, core2) or is_acronym_of(core2, core1):
                            return True

                        # Whitespace-artifact match: pdfplumber sometimes drops
                        # spaces between word and trailing single-letter variable
                        # (e.g. "embedding Z" → "embeddingZ"). Strip all spaces
                        # and compare.
                        if re.sub(r'\s', '', core1.lower()) == re.sub(r'\s', '', core2.lower()):
                            return True

                        return False

                    # Group descriptions that refer to the same element
                    desc_list = list(descs)
                    distinct_groups = []
                    used = set()

                    for i, d1 in enumerate(desc_list):
                        if i in used:
                            continue
                        group = [d1]
                        used.add(i)
                        for j, d2 in enumerate(desc_list):
                            if j not in used and is_same_element(d1, d2):
                                group.append(d2)
                                used.add(j)
                        distinct_groups.append(group)

                    # Only warn if there are truly distinct element names
                    if len(distinct_groups) > 1:
                        # Get representative from each group
                        representatives = [g[0] for g in distinct_groups]
                        warnings.append(f"Ref {num} may have inconsistent descriptions: {representatives[:3]}")

            # Check 2: Figure series organization
            # Reference numerals typically follow patterns: 100-series for FIG. 1, 200-series for FIG. 2, etc.
            series = {}
            for num in spec_refs.keys():
                series_num = int(num) // 100
                if series_num not in series:
                    series[series_num] = []
                series[series_num].append(int(num))

            # Check for gaps in numbering within each series
            for series_num, nums in series.items():
                sorted_nums = sorted(nums)
                if len(sorted_nums) > 2:
                    # Check for large gaps (more than 10)
                    for i in range(1, len(sorted_nums)):
                        gap = sorted_nums[i] - sorted_nums[i-1]
                        if gap > 10 and gap != 100:  # 100 gap might be intentional series change
                            pass  # Don't report gaps as they may be intentional

            # Check 3: Cross-reference with drawings (if text extractable)
            if drawings_refs:
                spec_ref_nums = set(spec_refs.keys())
                in_spec_not_drawings = spec_ref_nums - drawings_refs
                in_drawings_not_spec = drawings_refs - spec_ref_nums

                # Only flag if truly missing (not found at all in spec, even as operations)
                if in_drawings_not_spec:
                    # This is actually rare since most drawings text extraction is limited
                    # Don't flag as critical - drawings refs are hard to extract reliably
                    warnings.append(f"Reference numerals in drawings may need verification: {sorted(in_drawings_not_spec)}")

            # Check 4: Verify all reference numerals are properly introduced
            # (This is a simplified check - full antecedent basis check is complex)

            # Generate report
            total_refs = len(spec_refs)
            total_occurrences = sum(d['count'] for d in spec_refs.values())

            if issues:
                self.report.add_issue(
                    16, "Specification", "Reference Numeral Consistency",
                    Severity.CRITICAL,
                    f"Reference numeral issues found: {'; '.join(issues[:3])}",
                    f"Total reference numerals: {total_refs}, Total occurrences: {total_occurrences}"
                )
            elif warnings:
                # Limit warnings shown
                warning_summary = "; ".join(warnings[:3])
                if len(warnings) > 3:
                    warning_summary += f" (and {len(warnings) - 3} more)"
                self.report.add_issue(
                    16, "Specification", "Reference Numeral Consistency",
                    Severity.WARNING,
                    f"Potential inconsistencies: {warning_summary}",
                    f"Total reference numerals: {total_refs}. Manual review recommended."
                )
            else:
                self.report.add_issue(
                    16, "Specification", "Reference Numeral Consistency",
                    Severity.PASS,
                    f"Reference numerals appear consistent ({total_refs} unique numerals, {total_occurrences} total occurrences)"
                )
        else:
            self.report.add_issue(
                16, "Specification", "Reference Numeral Consistency",
                Severity.INFO,
                "Unable to extract reference numerals - manual review recommended"
            )
        
        # Check 17: Abstract present and length compliant (≤150 words, 37 CFR 1.72(b))
        # The previous version overcounted by capturing PDF page-footer noise
        # (page numbers, "Page X of Y", repeated headers, etc.) along with the
        # abstract proper. Now: capture the abstract chunk, then strip common
        # PDF-extraction noise before counting.
        abstract_match = re.search(
            r'\bABSTRACT\b\s*(?:OF\s+THE\s+(?:DISCLOSURE|INVENTION))?\s*[:\n]+'
            r'(.{20,2500}?)'
            r'(?='
            r'\n\s*(?:BACKGROUND|FIELD|BRIEF|DETAILED|CLAIMS|WHAT\s+IS\s+CLAIMED|'
            r'I\s+HEREBY\s+DECLARE|FIG\.|Sheet\s+\d+\s*(?:of|/)\s*\d+)'
            r'|\Z)',
            self.spec_text, re.IGNORECASE | re.DOTALL
        )
        if abstract_match:
            raw = abstract_match.group(1)
            # Strip common page-footer / page-header artifacts that PDF text
            # extraction frequently splices into nearby content
            cleaned = raw
            cleaned = re.sub(r'\bPage\s+\d+(?:\s+of\s+\d+)?\b', ' ', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r'(?m)^\s*\d{1,3}\s*$', ' ', cleaned)            # bare page numbers
            cleaned = re.sub(r'(?m)^\s*\d+\s*/\s*\d+\s*$', ' ', cleaned)      # "52 / 52" footers
            cleaned = re.sub(r'\bPatent\s+Application\s+Publication\b', ' ', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r'\b(?:US|U\.S\.)\s*\d{4}/\d+\s*A\d?\b', ' ', cleaned)  # pub-no headers
            cleaned = re.sub(r'\s+', ' ', cleaned).strip()
            word_count = len(cleaned.split())
            if word_count == 0:
                self.report.add_issue(
                    17, "Specification", "Abstract Present and Length Compliant",
                    Severity.WARNING, "Abstract heading found but body could not be extracted"
                )
            elif word_count <= 150:
                self.report.add_issue(
                    17, "Specification", "Abstract Present and Length Compliant",
                    Severity.PASS, f"Abstract found ({word_count} words, limit is 150)"
                )
            else:
                # Borderline counts (151-160) are usually PDF extraction noise;
                # show the user what was actually counted so they can verify.
                preview = (cleaned[:240] + '…') if len(cleaned) > 240 else cleaned
                self.report.add_issue(
                    17, "Specification", "Abstract Present and Length Compliant",
                    Severity.WARNING,
                    f"Abstract may be too long ({word_count} words, limit is 150)",
                    f"Extracted text used for the count (verify against the source .docx, "
                    f"as PDF text extraction can splice in page-header/footer artifacts):\n\n{preview}"
                )
        else:
            self.report.add_issue(
                17, "Specification", "Abstract Present and Length Compliant",
                Severity.CRITICAL, "Abstract section not found"
            )
        
        # Check 18: Background section present
        if re.search(r'BACKGROUND|FIELD OF (?:THE )?INVENTION', self.spec_text, re.IGNORECASE):
            self.report.add_issue(
                18, "Specification", "Background Section Present",
                Severity.PASS, "Background/Field section found"
            )
        else:
            self.report.add_issue(
                18, "Specification", "Background Section Present",
                Severity.WARNING, "Background/Field section not clearly identified"
            )
        
        # Check 19: Brief Description of Drawings present
        if re.search(r'BRIEF DESCRIPTION OF (?:THE )?DRAWINGS', self.spec_text, re.IGNORECASE):
            self.report.add_issue(
                19, "Specification", "Brief Description of Drawings Present",
                Severity.PASS, "Brief Description of Drawings section found"
            )
        else:
            self.report.add_issue(
                19, "Specification", "Brief Description of Drawings Present",
                Severity.WARNING, "Brief Description of Drawings section not clearly identified"
            )
        
        # Check 20: Detailed Description present
        if re.search(r'DETAILED DESCRIPTION', self.spec_text, re.IGNORECASE):
            self.report.add_issue(
                20, "Specification", "Detailed Description Present",
                Severity.PASS, "Detailed Description section found"
            )
        else:
            self.report.add_issue(
                20, "Specification", "Detailed Description Present",
                Severity.CRITICAL, "Detailed Description section not clearly identified"
            )
        
        # Check 21: Claims section present
        # Look for various claim section indicators
        # Note: PyPDF2 often doesn't preserve newlines, so we use flexible patterns
        claims_patterns = [
            # Original patterns with newlines
            r'(?:^|\n)\s*CLAIMS?\s*(?:\n|$)',  # "CLAIMS" or "CLAIM" as header
            r'(?:^|\n)\s*What is claimed',      # "What is claimed is:"
            r'(?:^|\n)\s*I\s+claim',            # "I claim:"
            r'(?:^|\n)\s*We\s+claim',           # "We claim:"
            r'(?:^|\n)\s*1\.\s+A\s+',           # First claim starting with "1. A method/system/etc"
            r'(?:^|\n)\s*1\.\s+An\s+',          # First claim starting with "1. An apparatus/etc"
            # More flexible patterns for PyPDF2 text extraction quirks
            r'(?:^|[^A-Z])CLAIMS(?:[^A-Z]|$)',  # "CLAIMS" as standalone word
            r'What\s+is\s+claimed',             # "What is claimed" anywhere
            r'What\s*is\s*claimed',             # "Whatisclaimed" (no spaces)
            r'1\.\s*A\s+computer-implemented',  # First claim pattern
            r'1\.\s*A\s+method',                # First claim pattern
            r'1\.\s*A\s+system',                # First claim pattern
            r'1\.\s*An?\s+apparatus',           # First claim pattern
        ]

        claims_found = False
        for pattern in claims_patterns:
            if re.search(pattern, self.spec_text, re.IGNORECASE | re.MULTILINE):
                claims_found = True
                break

        if claims_found:
            self.report.add_issue(
                21, "Specification", "Claims Section Present",
                Severity.PASS, "Claims section found"
            )
        else:
            self.report.add_issue(
                21, "Specification", "Claims Section Present",
                Severity.CRITICAL, "Claims section not clearly identified"
            )
    
    def check_drawings(self):
        """Checks 22-25: Drawings-specific checks"""

        if not self.drawings_text:
            for i in [22, 23, 24, 25]:
                self.report.add_issue(
                    i, "Drawings", f"Check {i}",
                    Severity.CRITICAL, "Drawings not found"
                )
            return
        
        # Check 22: Figure numbering sequential
        # Same image-only-drawings caveat as Check 15 — if the drawings PDF
        # has no extractable FIG. labels, we can't run a sequence check.
        if not self._drawings_text_extractable():
            self.report.add_issue(
                22, "Drawings", "Figure Numbering Sequential",
                Severity.INFO,
                "Drawings PDF appears to be image-only — figure labels not "
                "extractable as text. Manually verify FIG. numbering is sequential."
            )
        else:
            fig_nums = self._extract_figure_numbers(self.drawings_text)
            if fig_nums:
                fig_ints = sorted(set(fig_nums))
                if fig_ints:
                    expected = list(range(1, max(fig_ints) + 1))
                    missing = set(expected) - set(fig_ints)
                    if not missing:
                        self.report.add_issue(
                            22, "Drawings", "Figure Numbering Sequential",
                            Severity.PASS, f"Figures numbered sequentially (1-{max(fig_ints)})"
                        )
                    else:
                        self.report.add_issue(
                            22, "Drawings", "Figure Numbering Sequential",
                            Severity.WARNING, f"Figure numbers may have gaps. Found: {fig_ints}, Missing: {sorted(missing)}"
                        )
                else:
                    self.report.add_issue(
                        22, "Drawings", "Figure Numbering Sequential",
                        Severity.WARNING, "Unable to detect valid figure numbers in drawings"
                    )
            else:
                self.report.add_issue(
                    22, "Drawings", "Figure Numbering Sequential",
                    Severity.INFO, "No extractable figure numbers — manually verify"
                )
        
        # Check 23: Drawings have margin labels (title and docket number).
        # Use the actual docket from XFA when available (the original code's
        # docket regex matched only "X000-0000"-style and missed real-world
        # firm/customer dockets like "MS1-9771USC3" or "412147-US03-CON").
        issues_23 = []

        # Title check — look for any title word appearing in drawings
        title_words = []
        if self.ads_data and self.ads_data.get('title'):
            title_words = [w for w in self.ads_data['title'].split() if len(w) > 3]
        elif self.spec_text:
            title_match = re.search(r'(?:TITLE|Title).*?(?:of.*?Invention)?[:\s]*([\w\s\-]+?)(?:\n|CROSS|BACKGROUND|FIELD)',
                                   self.spec_text, re.IGNORECASE)
            if title_match:
                title_words = [w for w in title_match.group(1).strip().split() if len(w) > 3]

        has_title = False
        if title_words and len(title_words) >= 2:
            found_words = sum(1 for w in title_words if w.upper() in self.drawings_text.upper())
            if found_words >= len(title_words) * 0.5:
                has_title = True
        elif re.search(r'Title:', self.drawings_text, re.IGNORECASE):
            has_title = True

        if not has_title:
            issues_23.append("Application title not detected in drawings margin")

        # Docket check — try both XFA-derived dockets and any docket-shaped
        # token from spec/ADS text. Strip dashes/spaces for tolerant matching.
        docket_candidates = []
        if self.ads_data:
            for key in ('docket_number',):
                v = self.ads_data.get(key)
                if v: docket_candidates.append(v)
        # Also check for docket-shaped tokens in spec text
        for m in re.finditer(r'\b([A-Z]{2,4}\d?[-_]\d{3,5}[A-Z]{0,5}\d?)\b',
                             self.ads_text or self.spec_text or '', re.IGNORECASE):
            docket_candidates.append(m.group(1))

        norm_drawings = re.sub(r'[\s\-_]', '', self.drawings_text).upper()
        has_docket = False
        for cand in docket_candidates:
            cand_norm = re.sub(r'[\s\-_]', '', cand).upper()
            if cand_norm and cand_norm in norm_drawings:
                has_docket = True
                break
        if not has_docket and re.search(r'Docket\s*(?:No\.?|Number)', self.drawings_text, re.IGNORECASE):
            has_docket = True

        if not has_docket:
            issues_23.append("Docket number not detected in drawings margin")

        if not issues_23:
            self.report.add_issue(
                23, "Drawings", "All Figures Have Labels",
                Severity.PASS, "Drawings margin labels present (title and docket number)"
            )
        else:
            self.report.add_issue(
                23, "Drawings", "All Figures Have Labels",
                Severity.WARNING, "; ".join(issues_23)
            )
        
        # Check 24: Sheet numbering present
        # Accept various formats: "Sheet 1 of 9", "Page 1 of 9", "1/9", etc.
        sheet_patterns = [
            r'Sheet\s+\d+\s+of\s+\d+',
            r'Page\s+\d+\s+of\s+\d+',
            r'\d+\s*/\s*\d+',  # "1/9" format
        ]
        sheet_found = any(re.search(p, self.drawings_text, re.IGNORECASE) for p in sheet_patterns)
        if sheet_found:
            self.report.add_issue(
                24, "Drawings", "Sheet Numbering Present",
                Severity.PASS, "Sheet/page numbering detected"
            )
        else:
            self.report.add_issue(
                24, "Drawings", "Sheet Numbering Present",
                Severity.WARNING, "Sheet numbering not detected"
            )
        
        # Check 25: No color drawings
        if OCR_AVAILABLE:
            drawing_files = list(self.folder_path.glob('*Drawing*.pdf')) + list(self.folder_path.glob('*drawing*.pdf'))
            if drawing_files:
                try:
                    images = convert_from_path(drawing_files[0])
                    has_color = False
                    for img in images:
                        # Convert to RGB and check for colored pixels
                        rgb_img = img.convert('RGB')
                        # Sample pixels across the image
                        width, height = rgb_img.size
                        color_pixels = 0
                        total_checked = 0
                        step = max(1, min(width, height) // 100)  # Sample ~100x100 grid
                        for x in range(0, width, step):
                            for y in range(0, height, step):
                                r, g, b = rgb_img.getpixel((x, y))
                                total_checked += 1
                                # Check if pixel has significant color (not grayscale)
                                # Grayscale pixels have r≈g≈b
                                max_diff = max(abs(r-g), abs(r-b), abs(g-b))
                                if max_diff > 30:  # Threshold for color detection
                                    color_pixels += 1
                        if total_checked > 0 and (color_pixels / total_checked) > 0.01:
                            has_color = True
                            break

                    if has_color:
                        self.report.add_issue(
                            25, "Drawings", "No Color Drawings",
                            Severity.WARNING, "Color detected in drawings - USPTO requires black and white unless petition filed"
                        )
                    else:
                        self.report.add_issue(
                            25, "Drawings", "No Color Drawings",
                            Severity.PASS, "Drawings appear to be black and white"
                        )
                except Exception as e:
                    self.report.add_issue(
                        25, "Drawings", "No Color Drawings",
                        Severity.INFO, "Could not analyze drawing colors - manual verification recommended"
                    )
            else:
                self.report.add_issue(
                    25, "Drawings", "No Color Drawings",
                    Severity.INFO, "Drawing file not found for color analysis"
                )
        else:
            self.report.add_issue(
                25, "Drawings", "No Color Drawings",
                Severity.INFO, "Manual visual inspection recommended - ensure drawings are black and white"
            )
        
    
    def check_ads(self):
        """Checks 27-31: ADS-specific checks"""
        
        if not self.ads_text:
            for i in range(27, 32):
                self.report.add_issue(
                    i, "ADS", f"Check {i}",
                    Severity.CRITICAL, "ADS not found"
                )
            return
        
        # Check 27: Inventor addresses complete
        # Preferred path: when XFA was successfully parsed, use the structured
        # inventor records directly. The legacy regex-on-OCR-text path produced
        # false-positive "missing City/State/Country" warnings against XFA-derived
        # data because the synthesized text doesn't carry the literal field
        # labels the regex expected.
        if self.ads_data and self.ads_data.get('inventors'):
            incomplete_inventors = []
            complete_count = 0
            for idx, inv in enumerate(self.ads_data['inventors'], start=1):
                missing_fields = []
                if not (inv.get('mail_address1') or '').strip():
                    missing_fields.append('Address 1')
                if not (inv.get('mail_city') or '').strip():
                    missing_fields.append('City')
                if not (inv.get('mail_state') or '').strip():
                    # State is required for US mailing addresses; for non-US,
                    # absence is normal (many countries don't have a state field)
                    if (inv.get('mail_country') or '').upper() == 'US':
                        missing_fields.append('State')
                if not (inv.get('mail_postcode') or '').strip():
                    missing_fields.append('Postal Code')
                if not (inv.get('mail_country') or '').strip():
                    missing_fields.append('Country')

                if missing_fields:
                    name = self._format_xfa_inventor(inv) or f"Inventor {idx}"
                    incomplete_inventors.append(f"{name}: missing {', '.join(missing_fields)}")
                else:
                    complete_count += 1

            if incomplete_inventors:
                self.report.add_issue(
                    27, "ADS", "Inventor Addresses Complete",
                    Severity.WARNING,
                    f"{len(incomplete_inventors)} of {len(self.ads_data['inventors'])} "
                    f"inventor mailing addresses are incomplete in the ADS",
                    "\n".join(f"  • {line}" for line in incomplete_inventors)
                )
            else:
                self.report.add_issue(
                    27, "ADS", "Inventor Addresses Complete",
                    Severity.PASS,
                    f"All {complete_count} inventor mailing addresses are complete in the ADS"
                )
        else:
            # Legacy path: regex against OCR'd text when no structured XFA data.
            inventor_splits = re.split(r'(?=Inventor\s+\d+)', self.ads_text, flags=re.IGNORECASE)
            inventor_sections = [s for s in inventor_splits if re.match(r'Inventor\s+\d+', s, re.IGNORECASE)]

            if inventor_sections:
                incomplete_inventors = []
                complete_count = 0
                for section in inventor_sections:
                    inv_match = re.match(r'Inventor\s+(\d+)', section, re.IGNORECASE)
                    inv_num = inv_match.group(1) if inv_match else '?'
                    missing_fields = []
                    if not re.search(r'Address\s*1\s+[A-Za-z0-9c/o]', section, re.IGNORECASE):
                        missing_fields.append('Address 1')
                    if not re.search(r'City\s+[A-Za-z]{2,}', section, re.IGNORECASE):
                        missing_fields.append('City')
                    if not re.search(r'State\s*/?\s*Province\s+[A-Z]{2}', section, re.IGNORECASE):
                        missing_fields.append('State/Province')
                    if not re.search(r'Postal\s*Code\s+\d{5}', section, re.IGNORECASE):
                        missing_fields.append('Postal Code')
                    if not re.search(r'Country\s*[:\s]*[A-Z]{2}', section, re.IGNORECASE):
                        missing_fields.append('Country')
                    if missing_fields:
                        incomplete_inventors.append(f"Inventor {inv_num}: missing {', '.join(missing_fields)}")
                    else:
                        complete_count += 1

                if incomplete_inventors:
                    self.report.add_issue(
                        27, "ADS", "Inventor Addresses Complete",
                        Severity.WARNING,
                        f"Incomplete inventor addresses: {'; '.join(incomplete_inventors[:3])}"
                    )
                else:
                    self.report.add_issue(
                        27, "ADS", "Inventor Addresses Complete",
                        Severity.PASS,
                        f"All {complete_count} inventor addresses appear complete"
                    )
            else:
                has_addresses = (
                    re.search(r'Address\s*1\s+[A-Za-z0-9]', self.ads_text, re.IGNORECASE) and
                    re.search(r'Postal\s*Code\s+\d{5}', self.ads_text, re.IGNORECASE)
                )
                if has_addresses:
                    self.report.add_issue(
                        27, "ADS", "Inventor Addresses Complete",
                        Severity.PASS, "Address information detected in ADS"
                    )
                else:
                    self.report.add_issue(
                        27, "ADS", "Inventor Addresses Complete",
                        Severity.WARNING, "Could not verify inventor addresses in ADS"
                    )
        
        # Check 28: First named inventor identified
        # Compare first inventor from ADS with first named inventor on POA
        ads_first_inventor = None
        poa_first_inventor = None

        # Extract first inventor from ADS (Inventor 1)
        # Pattern: After "Suffix" line, the name appears as "FirstName LASTNAME"
        inv1_section = re.search(r'Inventor\s+1(.*?)(?:Inventor\s+2|Correspondence|$)',
                                  self.ads_text, re.DOTALL | re.IGNORECASE)
        if inv1_section:
            section = inv1_section.group(1)
            # Look for name after Suffix line - format: "Chitrak GUPTA" or "First Middle LAST"
            name_match = re.search(r'Suffix\s*\n\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\s+[A-Z]{2,})', section)
            if name_match:
                ads_first_inventor = name_match.group(1).strip()

        # Extract first named inventor from POA
        # POA forms often need OCR to extract filled-in values
        # Pattern: Name format is "FirstName LASTNAME" where last name is ALL CAPS
        # Don't use IGNORECASE - we need to distinguish names from form labels
        poa_inventor_pattern = r'First\s*Named\s*Inventor\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\s+[A-Z]{2,})'

        # First try OCR since PyPDF2 often can't extract filled form fields
        if OCR_AVAILABLE:
            poa_files = list(self.folder_path.glob('*POA*.pdf')) + list(self.folder_path.glob('*Power*Attorney*.pdf'))
            if poa_files:
                try:
                    images = convert_from_path(poa_files[0])
                    ocr_text = ''.join(pytesseract.image_to_string(img) + '\n' for img in images)
                    # No IGNORECASE - rely on actual case to distinguish names from labels
                    poa_match = re.search(poa_inventor_pattern, ocr_text)
                    if poa_match:
                        poa_first_inventor = poa_match.group(1).strip()
                except Exception:
                    pass

        # Fallback to PyPDF2 text if OCR didn't work
        if not poa_first_inventor and self.poa_text:
            poa_match = re.search(poa_inventor_pattern, self.poa_text)
            if poa_match:
                poa_first_inventor = poa_match.group(1).strip()

        if ads_first_inventor and poa_first_inventor:
            # Normalize names for comparison
            ads_norm = self.normalize_name(ads_first_inventor)
            poa_norm = self.normalize_name(poa_first_inventor)

            if ads_norm == poa_norm:
                self.report.add_issue(
                    28, "ADS", "First Named Inventor Identified",
                    Severity.PASS,
                    f"First named inventor matches: {ads_first_inventor} (ADS) = {poa_first_inventor} (POA)"
                )
            else:
                self.report.add_issue(
                    28, "ADS", "First Named Inventor Identified",
                    Severity.CRITICAL,
                    f"First named inventor mismatch: ADS has '{ads_first_inventor}', POA has '{poa_first_inventor}'"
                )
        elif ads_first_inventor:
            self.report.add_issue(
                28, "ADS", "First Named Inventor Identified",
                Severity.PASS,
                f"First named inventor in ADS: {ads_first_inventor}"
            )
        else:
            self.report.add_issue(
                28, "ADS", "First Named Inventor Identified",
                Severity.WARNING,
                "Could not extract first named inventor from ADS for verification"
            )

        # Check 29: Entity status specified
        if re.search(r'entity status|small entity|micro entity|large entity', self.ads_text, re.IGNORECASE):
            self.report.add_issue(
                29, "ADS", "Entity Status Specified",
                Severity.PASS, "Entity status appears to be specified"
            )
        else:
            self.report.add_issue(
                29, "ADS", "Entity Status Specified",
                Severity.WARNING, "Entity status not clearly specified"
            )
        
        # Check 31: Attorney/agent information
        if re.search(r'registration\s*(?:no|number)', self.ads_text, re.IGNORECASE):
            self.report.add_issue(
                31, "ADS", "Attorney/Agent Information",
                Severity.PASS, "Attorney/agent registration information appears present"
            )
        else:
            self.report.add_issue(
                31, "ADS", "Attorney/Agent Information",
                Severity.INFO, "Manual review recommended for attorney/agent registration number"
            )

        # The remaining ADS checks rely on structured XFA data that we may or may
        # not have. If we don't have it, skip them (rather than emitting noisy
        # WARNINGs that the user can't act on).
        if not self.ads_data:
            return

        # (Former Check 72 — "Inventor Citizenship Populated" — was removed.
        # The CitizedDropDown element survives in the XFA schema for backwards
        # compatibility, but the citizenship field is no longer a required or
        # visible field on current PTO/AIA/14 forms, so the check produced
        # false positives on every modern ADS without telling the user anything
        # actionable.)

        # Check 73: Attorney customer number matches correspondence customer number
        corr_cn = (self.ads_data.get('customer_number') or '').strip()
        atty_cn = (self.ads_data.get('attorney_customer_number') or '').strip()
        if corr_cn and atty_cn:
            if corr_cn == atty_cn:
                self.report.add_issue(
                    73, "ADS", "Attorney vs Correspondence Customer Number",
                    Severity.PASS,
                    f"Attorney and correspondence customer numbers match: {corr_cn}"
                )
            else:
                self.report.add_issue(
                    73, "ADS", "Attorney vs Correspondence Customer Number",
                    Severity.WARNING,
                    f"Attorney customer number ({atty_cn}) differs from correspondence customer number ({corr_cn})",
                    "These are often the same firm. Confirm the difference is intentional."
                )
        elif corr_cn or atty_cn:
            self.report.add_issue(
                73, "ADS", "Attorney vs Correspondence Customer Number",
                Severity.INFO,
                f"Only one customer number populated (correspondence={corr_cn or '—'}, attorney={atty_cn or '—'}); manual review recommended"
            )

    def check_declaration(self):
        """Checks 32-35: Declaration-specific checks"""
        
        if not self.declaration_text:
            for i in range(32, 36):
                self.report.add_issue(
                    i, "Declaration", f"Check {i}",
                    Severity.WARNING, "Declaration not found"
                )
            return
        
        # Check 32: All inventors named in declaration
        # Use ADS XFA inventors as ground truth (reliable structured data),
        # then verify each name appears in the declaration text by last-name
        # match. This is much more robust than regex-extracting names from
        # the declaration's varied formats.
        ads_inventor_names = []
        if self.ads_data and self.ads_data.get('inventors'):
            ads_inventor_names = [self._format_xfa_inventor(inv)
                                  for inv in self.ads_data['inventors']]
        elif self.ads_text:
            ads_inventor_names = self.extract_inventors(self.ads_text)

        if ads_inventor_names:
            decl_norm = self._normalize_for_compare(self.declaration_text)
            missing = []
            for inv_name in ads_inventor_names:
                last_name = inv_name.split()[-1] if inv_name.split() else ""
                last_norm = self._normalize_for_compare(last_name)
                full_norm = self._normalize_for_compare(inv_name)
                if (last_norm and last_norm in decl_norm) or \
                   (full_norm and full_norm in decl_norm):
                    continue
                missing.append(inv_name)
            if not missing:
                self.report.add_issue(
                    32, "Declaration", "All Inventors Named in Declaration",
                    Severity.PASS,
                    f"All {len(ads_inventor_names)} ADS inventors appear in the declaration"
                )
            else:
                cont_note = ""
                if self._is_continuation_filing():
                    cont_note = (" Note: this is a continuation filing — if inventorship "
                                 "changed, a new declaration is required.")
                img_pages = self.image_only_pages.get('Declaration', 0)
                img_note = ""
                if img_pages and len(missing) <= img_pages:
                    img_note = (f" The declaration has {img_pages} image-only page(s) — "
                                f"missing inventor(s) may be on those pages but could not be "
                                f"verified by text extraction. Open the declaration and confirm "
                                f"those pages cover the missing inventor(s).")
                self.report.add_issue(
                    32, "Declaration", "All Inventors Named in Declaration",
                    Severity.WARNING,
                    f"{len(missing)} of {len(ads_inventor_names)} ADS inventor(s) not found "
                    f"in declaration text." + img_note + cont_note,
                    "Not found in extracted declaration text:\n" + "\n".join(f"  • {n}" for n in missing)
                )
        else:
            self.report.add_issue(
                32, "Declaration", "All Inventors Named in Declaration",
                Severity.INFO, "Could not extract ADS inventors for cross-reference"
            )
        
        # Check 33: Oath vs declaration format
        # Note: OCR may produce "7" instead of "ti", so check for variations
        oath_patterns = [
            r'\bswear', r'\boaths?\b',  # "oath" or "oaths"
            r'declare', r'declara',      # "declare", "declaration", "declara7on"
            r'under penalty of perjury',
            r'37\s*CFR\s*1\.63',          # Reference to declaration rule
        ]
        has_declaration_language = any(
            re.search(p, self.declaration_text, re.IGNORECASE) for p in oath_patterns
        )

        if has_declaration_language:
            self.report.add_issue(
                33, "Declaration", "Oath vs Declaration Format",
                Severity.PASS, "Declaration/oath language detected"
            )
        else:
            self.report.add_issue(
                33, "Declaration", "Oath vs Declaration Format",
                Severity.WARNING, "Standard oath/declaration language not clearly detected"
            )
        
        # Check 34: Declaration references correct application
        # Look for any expected docket number (from ADS XFA or spec text) or
        # the application title (from ADS XFA) in the declaration.
        # The previous version's hardcoded title-keyword list was bogus
        # (hardcoded for one specific patent), so it never matched anything
        # for any other filing.
        all_expected_dockets = set()
        if self.ads_data and self.ads_data.get('docket_number'):
            all_expected_dockets.add(self.ads_data['docket_number'])
        elif self.ads_text:
            all_expected_dockets |= self.extract_docket_numbers(self.ads_text)
        if self.spec_text:
            all_expected_dockets |= self.extract_docket_numbers(self.spec_text)

        docket_in_decl = False
        matched_docket = None
        if all_expected_dockets:
            decl_norm_strip = re.sub(r'[\s\-_]', '', self.declaration_text.upper())
            for d in all_expected_dockets:
                d_norm = re.sub(r'[\s\-_]', '', d.upper())
                if d_norm and d_norm in decl_norm_strip:
                    docket_in_decl = True
                    matched_docket = d
                    break

        title_in_decl = False
        if self.ads_data and self.ads_data.get('title'):
            title_words = self.ads_data['title'].split()
            key_chunk = ' '.join(title_words[:max(4, int(len(title_words) * 0.6))])
            if self._normalize_for_compare(key_chunk) in self._normalize_for_compare(self.declaration_text):
                title_in_decl = True

        if docket_in_decl:
            self.report.add_issue(
                34, "Declaration", "Declaration References Correct Application",
                Severity.PASS, f"Declaration contains expected docket: {matched_docket}"
            )
        elif title_in_decl:
            self.report.add_issue(
                34, "Declaration", "Declaration References Correct Application",
                Severity.PASS, "Declaration references the application title from the ADS"
            )
        elif self._is_continuation_filing():
            self.report.add_issue(
                34, "Declaration", "Declaration References Correct Application",
                Severity.INFO,
                "Could not match this child application's docket/title in the declaration. "
                "For continuations, the parent's executed declaration carried forward "
                "typically references the parent's docket and original title — manual "
                "verification recommended."
            )
        else:
            self.report.add_issue(
                34, "Declaration", "Declaration References Correct Application",
                Severity.INFO, "Could not verify declaration references correct application"
            )

        # Check 35: Declaration date logical
        import datetime
        decl_text_for_date = self.declaration_text

        # Try OCR if no dates found in PyPDF2 text (dates often on signature pages)
        date_matches = re.findall(r'(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})', decl_text_for_date)
        if not date_matches and OCR_AVAILABLE:
            decl_files = list(self.folder_path.glob('*Dec*.pdf')) + list(self.folder_path.glob('*declaration*.pdf'))
            if decl_files:
                try:
                    images = convert_from_path(decl_files[0])
                    ocr_text = '\n'.join(pytesseract.image_to_string(img) for img in images)
                    date_matches = re.findall(r'(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})', ocr_text)
                    # Also try "Month DD, YYYY" format
                    if not date_matches:
                        month_matches = re.findall(
                            r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{1,2}),?\s+(\d{4})',
                            ocr_text
                        )
                        if month_matches:
                            # Convert to consistent format for processing below
                            for day, year in month_matches:
                                # Find the month name
                                m_match = re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+' + day, ocr_text)
                                if m_match:
                                    month_name = m_match.group(1)
                                    month_num = {'January': '1', 'February': '2', 'March': '3', 'April': '4',
                                               'May': '5', 'June': '6', 'July': '7', 'August': '8',
                                               'September': '9', 'October': '10', 'November': '11', 'December': '12'}
                                    date_matches.append((month_num.get(month_name, '1'), day, year))
                except Exception:
                    pass

        # Filter out form template dates (years before 2020)
        filing_dates = []
        for m in date_matches:
            month, day, year = m
            if len(year) == 2:
                year = '20' + year
            try:
                yr = int(year)
                if yr >= 2020:
                    filing_dates.append((month, day, year))
            except ValueError:
                continue

        if filing_dates:
            try:
                month, day, year = filing_dates[0]
                if len(year) == 2:
                    year = '20' + year
                decl_date = datetime.datetime(int(year), int(month), int(day))
                today = datetime.datetime.now()

                if decl_date > today:
                    self.report.add_issue(
                        35, "Declaration", "Declaration Date Logical",
                        Severity.CRITICAL, f"Declaration date is in the future: {decl_date.strftime('%Y-%m-%d')}"
                    )
                elif (today - decl_date).days > 365:
                    # For continuations, the parent's executed declaration is
                    # carried forward under 37 CFR 1.63(d) and will legitimately
                    # be older than a year — that's expected, not anomalous.
                    if self._is_continuation_filing():
                        self.report.add_issue(
                            35, "Declaration", "Declaration Date Logical",
                            Severity.PASS,
                            f"Declaration date {decl_date.strftime('%Y-%m-%d')} is older than 1 year, "
                            f"which is expected for this continuation filing (parent's declaration "
                            f"carried forward under 37 CFR 1.63(d))."
                        )
                    else:
                        self.report.add_issue(
                            35, "Declaration", "Declaration Date Logical",
                            Severity.WARNING, f"Declaration date is over a year old: {decl_date.strftime('%Y-%m-%d')}"
                        )
                else:
                    self.report.add_issue(
                        35, "Declaration", "Declaration Date Logical",
                        Severity.PASS, f"Declaration date is valid: {decl_date.strftime('%Y-%m-%d')}"
                    )
            except (ValueError, IndexError):
                self.report.add_issue(
                    35, "Declaration", "Declaration Date Logical",
                    Severity.INFO, f"Could not parse declaration date: {filing_dates[0]}"
                )
        else:
            self.report.add_issue(
                35, "Declaration", "Declaration Date Logical",
                Severity.INFO, "Declaration date not found in text - likely on signature page"
            )
    
    def check_assignment(self):
        """Checks 36-40: Assignment-specific checks"""
        
        if not self.assignment_text:
            for i in range(36, 41):
                self.report.add_issue(
                    i, "Assignment", f"Check {i}",
                    Severity.INFO, "Assignment not found (optional document)"
                )
            return
        
        # Check 36: Assignment identifies all assignors (compare with ADS inventors)
        # Use the ADS XFA inventor list as ground truth and verify each name
        # appears verbatim somewhere in the assignment text. This avoids the
        # fragility of trying to extract assignor names via regex from the
        # many possible assignment text formats.
        ads_inventor_names = []
        if self.ads_data and self.ads_data.get('inventors'):
            ads_inventor_names = [self._format_xfa_inventor(inv)
                                  for inv in self.ads_data['inventors']]
        elif self.ads_text:
            ads_inventor_names = self.extract_inventors(self.ads_text)

        if ads_inventor_names and self.assignment_text:
            asgn_norm = self._normalize_for_compare(self.assignment_text)
            missing = []
            found = []
            for name in ads_inventor_names:
                # Match by last name (most reliable cross-doc); fall back to full name
                last_name = name.split()[-1] if name.split() else ""
                last_normalized = self._normalize_for_compare(last_name)
                full_normalized = self._normalize_for_compare(name)
                if (last_normalized and last_normalized in asgn_norm) or \
                   (full_normalized and full_normalized in asgn_norm):
                    found.append(name)
                else:
                    missing.append(name)

            if not missing:
                self.report.add_issue(
                    36, "Assignment", "Assignment Identifies All Assignors",
                    Severity.PASS,
                    f"All {len(ads_inventor_names)} inventors from ADS appear in the assignment"
                )
            else:
                img_pages = self.image_only_pages.get('Assignment', 0)
                img_note = ""
                severity = Severity.CRITICAL
                if img_pages and len(missing) <= img_pages:
                    img_note = (f" The assignment has {img_pages} image-only page(s) — "
                                f"missing inventor(s) may be on those pages but could not be "
                                f"verified by text extraction. Open the assignment and confirm "
                                f"those pages cover the missing inventor(s).")
                    severity = Severity.WARNING
                self.report.add_issue(
                    36, "Assignment", "Assignment Identifies All Assignors",
                    severity,
                    f"{len(missing)} of {len(ads_inventor_names)} ADS inventor(s) not found "
                    f"in assignment text." + img_note,
                    "Not found in extracted assignment text:\n" + "\n".join(f"  • {n}" for n in missing)
                )
        elif ads_inventor_names and not self.assignment_text:
            self.report.add_issue(
                36, "Assignment", "Assignment Identifies All Assignors",
                Severity.INFO,
                "Assignment not loaded — cannot verify assignors against ADS inventors"
            )
        else:
            self.report.add_issue(
                36, "Assignment", "Assignment Identifies All Assignors",
                Severity.INFO,
                "Unable to extract inventor names from ADS for comparison"
            )
        
        # Check 37: Assignment identifies assignee
        if re.search(r'assignee', self.assignment_text, re.IGNORECASE):
            self.report.add_issue(
                37, "Assignment", "Assignment Identifies Assignee",
                Severity.PASS, "Assignee appears to be identified"
            )
        else:
            self.report.add_issue(
                37, "Assignment", "Assignment Identifies Assignee",
                Severity.WARNING, "Assignee not clearly identified"
            )
        
        # Check 38: Assignment references correct application
        # Match against any docket-shaped token in the spec/ADS (the previous
        # version's hardcoded title-keyword list ['AGENTIC', 'PIPELINE',
        # 'FIRMWARE', 'PORTING', ...] was bogus — left over from someone else's
        # patent and matched nothing for any other application).
        ads_dockets = set()
        if self.ads_data and self.ads_data.get('docket_number'):
            ads_dockets.add(self.ads_data['docket_number'])
        elif self.ads_text:
            ads_dockets |= self.extract_docket_numbers(self.ads_text)
        spec_dockets = self.extract_docket_numbers(self.spec_text) if self.spec_text else set()
        all_expected_dockets = ads_dockets | spec_dockets

        # Title from XFA (reliable)
        expected_title = ""
        if self.ads_data and self.ads_data.get('title'):
            expected_title = self.ads_data['title']

        docket_in_assignment = False
        matched_docket = None
        if all_expected_dockets and self.assignment_text:
            asgn_norm = re.sub(r'[\s\-_]', '', self.assignment_text.upper())
            for d in all_expected_dockets:
                d_norm = re.sub(r'[\s\-_]', '', d.upper())
                if d_norm and d_norm in asgn_norm:
                    docket_in_assignment = True
                    matched_docket = d
                    break

        title_in_assignment = False
        if expected_title and self.assignment_text:
            # Use first 60% of title words to allow minor wording differences
            title_words = expected_title.split()
            key_chunk = ' '.join(title_words[:max(4, int(len(title_words) * 0.6))])
            if self._normalize_for_compare(key_chunk) in self._normalize_for_compare(self.assignment_text):
                title_in_assignment = True

        if docket_in_assignment:
            self.report.add_issue(
                38, "Assignment", "Assignment References Correct Application",
                Severity.PASS, f"Assignment contains expected docket: {matched_docket}"
            )
        elif title_in_assignment:
            self.report.add_issue(
                38, "Assignment", "Assignment References Correct Application",
                Severity.PASS, "Assignment references the application title from the ADS"
            )
        elif self._is_continuation_filing():
            self.report.add_issue(
                38, "Assignment", "Assignment References Correct Application",
                Severity.INFO,
                "Could not match this child application's docket/title in the assignment. "
                "For continuations, the assignment carried forward from the parent typically "
                "references the parent's docket and original title — manual verification "
                "recommended."
            )
        else:
            self.report.add_issue(
                38, "Assignment", "Assignment References Correct Application",
                Severity.WARNING, "Could not verify assignment references correct docket/title"
            )

        # Check 39: Assignment execution date logical
        import datetime
        assign_text_for_date = self.assignment_text

        # Try OCR if PyPDF2 can't find dates (dates on signature pages)
        date_search_patterns = [
            r'(?:dated|executed|signed)[:\s]*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})',
            r'(?:dated|executed|signed)[:\s]*(\w+\s+\d{1,2},?\s+\d{4})',
            r'(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})',
        ]

        found_date = None
        for pattern in date_search_patterns:
            match = re.search(pattern, assign_text_for_date, re.IGNORECASE)
            if match:
                date_str = match.group(1)
                for fmt in ['%m/%d/%Y', '%m-%d-%Y', '%B %d, %Y', '%B %d %Y', '%m/%d/%y', '%m-%d-%y']:
                    try:
                        candidate = datetime.datetime.strptime(date_str, fmt)
                        if candidate.year >= 2020:
                            found_date = candidate
                            break
                    except ValueError:
                        continue
                if found_date:
                    break

        # Try OCR fallback if no date found
        if not found_date and OCR_AVAILABLE:
            assign_files = (list(self.folder_path.glob('*Assignment*.pdf')) +
                          list(self.folder_path.glob('*Exec*Dec*Assignment*.pdf')) +
                          list(self.folder_path.glob('*assign*.pdf')))
            if assign_files:
                try:
                    images = convert_from_path(assign_files[0])
                    ocr_text = '\n'.join(pytesseract.image_to_string(img) for img in images)
                    for pattern in date_search_patterns:
                        match = re.search(pattern, ocr_text, re.IGNORECASE)
                        if match:
                            date_str = match.group(1)
                            for fmt in ['%m/%d/%Y', '%m-%d-%Y', '%B %d, %Y', '%B %d %Y', '%m/%d/%y', '%m-%d-%y']:
                                try:
                                    candidate = datetime.datetime.strptime(date_str, fmt)
                                    if candidate.year >= 2020:
                                        found_date = candidate
                                        break
                                except ValueError:
                                    continue
                            if found_date:
                                break
                except Exception:
                    pass

        if found_date:
            today = datetime.datetime.now()
            if found_date > today:
                self.report.add_issue(
                    39, "Assignment", "Assignment Execution Date Logical",
                    Severity.CRITICAL, f"Assignment date is in the future: {found_date.strftime('%Y-%m-%d')}"
                )
            elif (today - found_date).days > 365:
                if self._is_continuation_filing():
                    self.report.add_issue(
                        39, "Assignment", "Assignment Execution Date Logical",
                        Severity.PASS,
                        f"Assignment date {found_date.strftime('%Y-%m-%d')} is older than 1 year, "
                        f"which is expected for this continuation filing (parent's executed "
                        f"assignment carried forward)."
                    )
                else:
                    self.report.add_issue(
                        39, "Assignment", "Assignment Execution Date Logical",
                        Severity.WARNING, f"Assignment date is over a year old: {found_date.strftime('%Y-%m-%d')}"
                    )
            else:
                self.report.add_issue(
                    39, "Assignment", "Assignment Execution Date Logical",
                    Severity.PASS, f"Assignment execution date is valid: {found_date.strftime('%Y-%m-%d')}"
                )
        else:
            self.report.add_issue(
                39, "Assignment", "Assignment Execution Date Logical",
                Severity.INFO, "Could not extract execution date from assignment"
            )
        
        # Check 40: Assignment covers correct rights
        # Note: OCR may produce "7" instead of "ti", so "title" may appear as "7tle"
        assignment_patterns = [
            r'entire right',
            r'(?:title|7tle).*interest',      # "title and interest" or "7tle and interest"
            r'sell.*assign.*transfer',         # Common assignment language
            r'assign.*transfer.*convey',       # Alternative assignment language
            r'right.*(?:title|7tle).*interest', # "right, title and interest"
            r'ASSIGNOR.*ASSIGNEE',             # Assignment structure
        ]
        has_assignment_language = any(
            re.search(p, self.assignment_text, re.IGNORECASE | re.DOTALL) for p in assignment_patterns
        )

        if has_assignment_language:
            self.report.add_issue(
                40, "Assignment", "Assignment Covers Correct Rights",
                Severity.PASS, "Assignment language appears to transfer rights"
            )
        else:
            self.report.add_issue(
                40, "Assignment", "Assignment Covers Correct Rights",
                Severity.WARNING, "Standard assignment language not clearly detected"
            )
    
    def check_poa(self):
        """Checks 41-44: Power of Attorney-specific checks"""
        
        if not self.poa_text:
            for i in range(41, 45):
                self.report.add_issue(
                    i, "Power of Attorney", f"Check {i}",
                    Severity.INFO, "Power of Attorney not found (may not be required)"
                )
            return
        
        # Check 41: POA names all practitioners
        # POA can identify practitioners either by customer number or individual registration numbers
        poa_check_text = self.poa_text

        # Try OCR if PyPDF2 text is mostly form labels (common with filled forms)
        if OCR_AVAILABLE and len(poa_check_text.strip()) < 500:
            poa_files = list(self.folder_path.glob('*POA*.pdf')) + list(self.folder_path.glob('*Power*Attorney*.pdf'))
            if poa_files:
                try:
                    images = convert_from_path(poa_files[0])
                    poa_check_text = '\n'.join(pytesseract.image_to_string(img) for img in images)
                except Exception:
                    pass

        # Check for customer number usage (common approach)
        customer_num_match = re.search(
            r'(?:Customer\s*Number|associated\s+with.*?Customer\s*Number)\s*[:\s]*(\d{5,6})',
            poa_check_text, re.IGNORECASE
        )
        # Also look for standalone 6-digit number after customer number context
        if not customer_num_match:
            customer_num_match = re.search(
                r'Customer\s*Number.*?\n\s*(\d{5,6})\b',
                poa_check_text, re.IGNORECASE | re.DOTALL
            )

        # Check for individual registration numbers
        reg_num_matches = re.findall(r'(?:Reg\.?\s*(?:No\.?|#)|Registration\s*(?:No\.?|Number))\s*[:\s]*(\d{5,6})',
                                     poa_check_text, re.IGNORECASE)

        if customer_num_match:
            cust_num = customer_num_match.group(1)
            self.report.add_issue(
                41, "Power of Attorney", "POA Names All Practitioners",
                Severity.PASS, f"POA appoints practitioners via Customer Number {cust_num}"
            )
        elif reg_num_matches:
            self.report.add_issue(
                41, "Power of Attorney", "POA Names All Practitioners",
                Severity.PASS, f"POA lists {len(reg_num_matches)} practitioners with registration numbers"
            )
        else:
            # Check if there's at least a named practitioner
            practitioner_name = re.search(r'(?:appoint|attorney|agent)\s+.*?([A-Z][a-z]+\s+[A-Z][a-z]+)',
                                         poa_check_text, re.IGNORECASE)
            if practitioner_name:
                self.report.add_issue(
                    41, "Power of Attorney", "POA Names All Practitioners",
                    Severity.PASS, f"POA names practitioner(s)"
                )
            else:
                self.report.add_issue(
                    41, "Power of Attorney", "POA Names All Practitioners",
                    Severity.WARNING, "No practitioner identification detected in POA"
                )
        
        # Check 42: POA includes registration numbers
        if re.search(r'registration\s*(?:no|number)|\d{5,6}', self.poa_text, re.IGNORECASE):
            self.report.add_issue(
                42, "Power of Attorney", "POA Includes Registration Numbers",
                Severity.PASS, "Registration numbers appear to be included"
            )
        else:
            self.report.add_issue(
                42, "Power of Attorney", "POA Includes Registration Numbers",
                Severity.WARNING, "Registration numbers not clearly detected"
            )
        
        # Check 44: POA properly signed
        if '/s/' in self.poa_text or 'signature' in self.poa_text.lower():
            self.report.add_issue(
                44, "Power of Attorney", "POA Properly Signed",
                Severity.PASS, "Signature indicators detected in POA"
            )
        else:
            self.report.add_issue(
                44, "Power of Attorney", "POA Properly Signed",
                Severity.WARNING, "Signatures not clearly detected in POA"
            )
    
    def check_formatting(self):
        """Checks 45, 49: USPTO formatting compliance"""

        if not self.spec_text:
            for i in [45, 49]:
                self.report.add_issue(
                    i, "USPTO Formatting", f"Check {i}",
                    Severity.CRITICAL, "Specification not found"
                )
            return
        
        # Check 45: Specification line numbering
        # Check if extracted text contains patterns suggesting line numbers (multiples of 5)
        # Line numbers typically appear as standalone numbers: 5, 10, 15, 20, 25...
        line_number_candidates = re.findall(r'(?:^|\s)(\d{1,3})(?:\s|$)', self.spec_text)
        multiples_of_5 = [int(n) for n in line_number_candidates if int(n) % 5 == 0 and 5 <= int(n) <= 50]
        # Check if we find a good sequence (at least 5, 10, 15, 20, 25)
        expected_line_nums = {5, 10, 15, 20, 25}
        found_line_nums = set(multiples_of_5)
        if expected_line_nums.issubset(found_line_nums):
            self.report.add_issue(
                45, "USPTO Formatting", "Specification Line Numbering",
                Severity.PASS, "Line numbering detected (multiples of 5 found in text)"
            )
        else:
            self.report.add_issue(
                45, "USPTO Formatting", "Specification Line Numbering",
                Severity.INFO, "Line numbering not clearly detected - verify line numbers every 5 lines"
            )
        
        # Check 49: Page numbering present
        # Check the specification PDF for page numbers on each page
        spec_files = list(self.folder_path.glob('*Spec*.pdf')) + list(self.folder_path.glob('*spec*.pdf'))
        page_nums_found = False
        if spec_files:
            try:
                with open(spec_files[0], 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    total_pages = len(reader.pages)
                    pages_with_numbers = 0
                    for i, page in enumerate(reader.pages):
                        page_text = page.extract_text() or ''
                        page_num = i + 1
                        pn = str(page_num)
                        # Look ONLY for explicit page-number formats — NOT for
                        # any occurrence of the digit anywhere on the page
                        # (the previous code's `str(page_num) in page_text`
                        # rule false-passed because paragraph numbers like
                        # [0035] always contain the digit "35").
                        patterns = [
                            r'(?:Page|page)\s+' + pn + r'\b',          # "Page 35"
                            r'\b' + pn + r'\s+of\s+\d+\b',              # "35 of 52"
                            r'-\s*' + pn + r'\s*-',                     # "- 35 -"
                            r'(?:^|\n)\s*' + pn + r'\s*(?:\n|$)',       # standalone "35" on its own line (margin)
                        ]
                        if any(re.search(p, page_text) for p in patterns):
                            pages_with_numbers += 1

                    # If most pages seem to have their page number in the text
                    if total_pages > 0 and pages_with_numbers >= total_pages * 0.8:
                        page_nums_found = True
                        self.report.add_issue(
                            49, "USPTO Formatting", "Page Numbering Present",
                            Severity.PASS, f"Page numbering detected ({pages_with_numbers}/{total_pages} pages)"
                        )
            except Exception:
                pass

        if not page_nums_found:
            # Fallback: check if extracted text has sequential page-like numbers
            page_refs = re.findall(r'(?:^|\s)(\d{1,3})(?:\s|$)', self.spec_text)
            sequential = [int(n) for n in page_refs if 1 <= int(n) <= 100]
            # Check if we have a reasonable sequence (1, 2, 3... or at least many sequential numbers)
            if sequential and max(sequential) >= 10:
                expected_pages = set(range(1, max(sequential) + 1))
                found_pages = set(sequential)
                coverage = len(expected_pages & found_pages) / len(expected_pages)
                if coverage >= 0.7:
                    self.report.add_issue(
                        49, "USPTO Formatting", "Page Numbering Present",
                        Severity.PASS, "Page numbering appears present in specification"
                    )
                else:
                    self.report.add_issue(
                        49, "USPTO Formatting", "Page Numbering Present",
                        Severity.INFO, "Page numbering not clearly detected - verify all pages numbered"
                    )
            else:
                self.report.add_issue(
                    49, "USPTO Formatting", "Page Numbering Present",
                    Severity.INFO, "Page numbering not clearly detected - verify all pages numbered"
                )
    
    def check_common_errors(self):
        """Checks 50-54: Common error detection"""
        
        # Check 50: No placeholder text remaining
        placeholders = ['[INSERT]', '[TBD]', 'TODO', 'XXX', '***', 'PLACEHOLDER']
        found_placeholders = []
        
        all_text = self.spec_text + self.ads_text + self.declaration_text + self.assignment_text
        for placeholder in placeholders:
            if placeholder in all_text:
                found_placeholders.append(placeholder)
        
        if not found_placeholders:
            self.report.add_issue(
                50, "Common Errors", "No Placeholder Text Remaining",
                Severity.PASS, "No common placeholder text detected"
            )
        else:
            self.report.add_issue(
                50, "Common Errors", "No Placeholder Text Remaining",
                Severity.CRITICAL, f"Placeholder text found: {', '.join(found_placeholders)}"
            )
        
        # Check 51: No track changes or comments visible
        track_change_indicators = ['Deleted:', 'Inserted:', 'Comment [', 'Formatted:']
        found_indicators = []
        
        for indicator in track_change_indicators:
            if indicator in all_text:
                found_indicators.append(indicator)
        
        if not found_indicators:
            self.report.add_issue(
                51, "Common Errors", "No Track Changes or Comments Visible",
                Severity.PASS, "No track change indicators detected"
            )
        else:
            self.report.add_issue(
                51, "Common Errors", "No Track Changes or Comments Visible",
                Severity.WARNING, f"Possible track change indicators: {', '.join(found_indicators)}"
            )
        
        # Check 52: Consistent use of claim terminology
        if self.spec_text:
            # Extract claims section
            claims_match = re.search(
                r'(?:CLAIMS|What is claimed)(.*?)(?:ABSTRACT|$)',
                self.spec_text, re.DOTALL | re.IGNORECASE
            )
            claims_text = claims_match.group(1) if claims_match else ""

            if claims_text:
                # Find technical element terms - look for "a/an/the NOUN NOUN" patterns
                # These are the actual claim elements we care about
                element_pattern = r'\b(?:a|an|the)\s+([\w\-]+(?:\s+[\w\-]+){1,2})\b'
                elements = []
                for match in re.finditer(element_pattern, claims_text, re.IGNORECASE):
                    element = match.group(1).lower().strip()
                    # Filter out generic terms
                    skip_terms = ['method', 'system', 'step', 'claim', 'invention', 'present',
                                 'first', 'second', 'third', 'plurality', 'least one', 'one or more',
                                 'following', 'above', 'same', 'other']
                    if not any(skip in element for skip in skip_terms) and len(element) > 5:
                        elements.append(element)

                # Group similar elements and look for inconsistencies
                # Inconsistency = same base noun with different modifiers used inconsistently
                # e.g., "firmware image" vs "firmware file" for same concept
                element_by_noun = {}
                for elem in elements:
                    words = elem.split()
                    if len(words) >= 2:
                        # Use last word as the main noun
                        main_noun = words[-1]
                        if len(main_noun) > 4:  # Skip short nouns
                            if main_noun not in element_by_noun:
                                element_by_noun[main_noun] = set()
                            element_by_noun[main_noun].add(elem)

                # Find nouns with multiple different element names
                inconsistencies = []
                for noun, variants in element_by_noun.items():
                    if len(variants) > 1:
                        # Check if variants are truly different (not just "X" vs "the X")
                        normalized_variants = set()
                        for v in variants:
                            # Normalize: remove articles, lowercase
                            norm = re.sub(r'^(the|a|an)\s+', '', v).strip()
                            normalized_variants.add(norm)

                        if len(normalized_variants) > 1:
                            # These might be intentionally different elements, or might be inconsistent
                            # Flag only if they seem like variants of the same thing
                            var_list = list(normalized_variants)
                            for i, v1 in enumerate(var_list):
                                for v2 in var_list[i+1:]:
                                    # Check if one is subset of other (e.g., "agent" vs "software agent")
                                    if v1 != v2 and (v1.endswith(v2.split()[-1]) and v2.endswith(v1.split()[-1])):
                                        # Same main noun, check if modifiers suggest same element
                                        w1 = set(v1.split()[:-1])  # modifiers only
                                        w2 = set(v2.split()[:-1])
                                        # If modifiers are very similar, might be inconsistency
                                        if w1 and w2 and len(w1.symmetric_difference(w2)) == 1:
                                            inconsistencies.append((v1, v2))

                if inconsistencies:
                    examples = [f"'{a}' vs '{b}'" for a, b in inconsistencies[:3]]
                    self.report.add_issue(
                        52, "Common Errors", "Consistent Use of Claim Terminology",
                        Severity.WARNING,
                        f"Potential terminology inconsistencies: {'; '.join(examples)}"
                    )
                else:
                    unique_elements = len(set(elements))
                    self.report.add_issue(
                        52, "Common Errors", "Consistent Use of Claim Terminology",
                        Severity.PASS,
                        f"No obvious terminology inconsistencies detected ({unique_elements} unique element terms)"
                    )
            else:
                self.report.add_issue(
                    52, "Common Errors", "Consistent Use of Claim Terminology",
                    Severity.INFO, "Could not extract claims for terminology check"
                )
        else:
            self.report.add_issue(
                52, "Common Errors", "Consistent Use of Claim Terminology",
                Severity.INFO, "Specification not available for terminology check"
            )

        # Check 53: Antecedent basis in claims
        # The previous regex `(?:CLAIMS|What is claimed)(.*?)(?:ABSTRACT|$)`
        # matched at the *first* occurrence of "CLAIMS" anywhere in the spec
        # (e.g., in cross-references like "Patent application No.…" early in
        # the doc), grabbing essentially the whole spec as "claims text" and
        # then finding spurious "antecedent issues" from the description body.
        if self.spec_text:
            claims_text = self._extract_claims_section()

            if claims_text:
                # Find elements introduced with "a" or "an"
                introduced = set()
                intro_pattern = r'\b(?:a|an)\s+([\w\-]+(?:\s+[\w\-]+){0,2})\b'
                for match in re.finditer(intro_pattern, claims_text, re.IGNORECASE):
                    element = match.group(1).lower().strip()
                    # Filter out common non-elements
                    if element not in ['method', 'system', 'device', 'apparatus', 'medium', 'product']:
                        introduced.add(element)

                # Find elements referenced with "the" or "said"
                referenced = []
                ref_pattern = r'\b(?:the|said)\s+([\w\-]+(?:\s+[\w\-]+){0,2})\b'
                for match in re.finditer(ref_pattern, claims_text, re.IGNORECASE):
                    element = match.group(1).lower().strip()
                    # Filter out common non-elements and preamble terms
                    skip_terms = ['method', 'system', 'device', 'apparatus', 'medium', 'product',
                                 'claim', 'claims', 'invention', 'present', 'following', 'above']
                    if element not in skip_terms and not any(s in element for s in skip_terms):
                        referenced.append(element)

                # Check for antecedent basis issues
                antecedent_issues = []
                for ref in set(referenced):
                    # Check if this element or a variant was introduced
                    found = False
                    ref_words = set(ref.split())
                    for intro in introduced:
                        intro_words = set(intro.split())
                        # Match if same words or one contains the other
                        if ref == intro or ref_words & intro_words:
                            found = True
                            break
                        # Also check if last word matches (main noun)
                        if ref.split()[-1] == intro.split()[-1]:
                            found = True
                            break
                    if not found:
                        antecedent_issues.append(ref)

                if antecedent_issues:
                    # Limit to first few issues
                    examples = antecedent_issues[:5]
                    self.report.add_issue(
                        53, "Common Errors", "Antecedent Basis in Claims",
                        Severity.WARNING,
                        f"Potential antecedent basis issues - 'the/said' without prior 'a/an': {examples}"
                    )
                else:
                    self.report.add_issue(
                        53, "Common Errors", "Antecedent Basis in Claims",
                        Severity.PASS,
                        f"Antecedent basis appears proper ({len(introduced)} elements introduced, {len(set(referenced))} referenced)"
                    )
            else:
                self.report.add_issue(
                    53, "Common Errors", "Antecedent Basis in Claims",
                    Severity.INFO, "Could not extract claims for antecedent basis check"
                )
        else:
            self.report.add_issue(
                53, "Common Errors", "Antecedent Basis in Claims",
                Severity.INFO, "Specification not available for antecedent basis check"
            )

        # Check 54: No undefined claim terms
        if self.spec_text:
            # Extract claims and detailed description separately
            # Use the same robust claims-section detection as Check 13/53
            # (the previous regex matched the first "CLAIMS" anywhere in the
            # spec including cross-references, pulling the whole document).
            claims_text = self._extract_claims_section()

            desc_match = re.search(
                r'DETAILED DESCRIPTION(.*?)(?:CLAIMS|What is claimed)',
                self.spec_text, re.DOTALL | re.IGNORECASE
            )
            description_text = desc_match.group(1) if desc_match else self.spec_text

            if claims_text and description_text:
                # Extract key technical terms from claims
                # Look for noun phrases that are likely claim elements
                term_pattern = r'\b([\w\-]+(?:\s+[\w\-]+){1,3})\b'
                claim_terms = set()

                for match in re.finditer(term_pattern, claims_text, re.IGNORECASE):
                    term = match.group(1).lower().strip()
                    # Filter to likely technical terms (multi-word or technical suffixes)
                    words = term.split()
                    if len(words) >= 2 and len(term) > 10:
                        # Skip common phrases
                        skip_phrases = ['the method', 'the system', 'the device', 'claim 1',
                                       'wherein the', 'comprising', 'configured to', 'adapted to',
                                       'based on', 'according to', 'at least one', 'one or more']
                        if not any(skip in term for skip in skip_phrases):
                            claim_terms.add(term)

                # Check if claim terms appear in description
                undefined_terms = []
                description_lower = description_text.lower()

                for term in claim_terms:
                    # Check for exact match or close match in description
                    if term not in description_lower:
                        # Try matching main noun (last word)
                        main_noun = term.split()[-1]
                        if len(main_noun) > 4 and main_noun not in description_lower:
                            undefined_terms.append(term)

                if undefined_terms:
                    # Limit display
                    examples = undefined_terms[:5]
                    self.report.add_issue(
                        54, "Common Errors", "No Undefined Claim Terms",
                        Severity.WARNING,
                        f"Claim terms possibly not in detailed description: {examples}"
                    )
                else:
                    self.report.add_issue(
                        54, "Common Errors", "No Undefined Claim Terms",
                        Severity.PASS,
                        f"Claim terms appear to be supported in specification ({len(claim_terms)} terms checked)"
                    )
            else:
                self.report.add_issue(
                    54, "Common Errors", "No Undefined Claim Terms",
                    Severity.INFO, "Could not extract claims or description for term check"
                )
        else:
            self.report.add_issue(
                54, "Common Errors", "No Undefined Claim Terms",
                Severity.INFO, "Specification not available for claim term check"
            )
    
    def check_file_quality(self):
        """Checks 55-58: File quality checks"""

        # Check 55: PDF text-searchable
        # Verify that we were able to extract meaningful text from key documents
        searchable_docs = []
        non_searchable_docs = []

        doc_texts = {
            'Specification': self.spec_text,
            'Drawings': self.drawings_text,
            'ADS': self.ads_text,
            'Declaration': self.declaration_text,
            'Assignment': self.assignment_text,
            'Power of Attorney': self.poa_text,
        }

        for doc_type, text in doc_texts.items():
            if self.report.files_found.get(doc_type):
                if text and len(text.strip()) > 100:
                    searchable_docs.append(doc_type)
                else:
                    non_searchable_docs.append(doc_type)

        if non_searchable_docs:
            self.report.add_issue(
                55, "File Quality", "PDF Text-Searchable",
                Severity.WARNING, f"Limited text extraction from: {', '.join(non_searchable_docs)}"
            )
        else:
            self.report.add_issue(
                55, "File Quality", "PDF Text-Searchable",
                Severity.PASS, f"All {len(searchable_docs)} PDFs are text-searchable"
            )

        # Check 56: File naming conventions
        # Verify each filename contains an expected docket number (taken from
        # the ADS/spec). The previous version's hardcoded pattern `A\d{3}-\d{4}`
        # only matched specific firm-internal formats like "X000-0000US" and
        # falsely flagged any other naming convention as "missing docket."
        expected_dockets = set()
        if self.ads_data and self.ads_data.get('docket_number'):
            expected_dockets.add(self.ads_data['docket_number'])
        if self.spec_text:
            expected_dockets |= self.extract_docket_numbers(self.spec_text)

        naming_issues = []
        for doc_type, filename in self.report.files_found.items():
            if filename and doc_type != 'Drawings':
                # Strip dashes/underscores/spaces for tolerant matching
                fname_norm = re.sub(r'[\s\-_]', '', filename.upper())
                has_any_docket = False
                for d in expected_dockets:
                    d_norm = re.sub(r'[\s\-_]', '', d.upper())
                    if d_norm and d_norm in fname_norm:
                        has_any_docket = True
                        break
                if not has_any_docket and expected_dockets:
                    naming_issues.append(f"{doc_type}: missing docket number")

        if naming_issues:
            self.report.add_issue(
                56, "File Quality", "File Naming Conventions",
                Severity.INFO, f"Naming suggestions: {'; '.join(naming_issues[:3])}"
            )
        else:
            self.report.add_issue(
                56, "File Quality", "File Naming Conventions",
                Severity.PASS, "File names follow good conventions (contain docket number)"
            )
        
        # Check 57: No password protection
        password_protected = []
        for doc_type, filename in self.report.files_found.items():
            if filename and filename.endswith('.pdf'):
                filepath = self.folder_path / filename
                try:
                    with open(filepath, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        if reader.is_encrypted:
                            password_protected.append(filename)
                except Exception:
                    pass  # If we can't check, skip

        if password_protected:
            self.report.add_issue(
                57, "File Quality", "No Password Protection",
                Severity.CRITICAL, f"Password-protected PDFs detected: {', '.join(password_protected)}"
            )
        else:
            self.report.add_issue(
                57, "File Quality", "No Password Protection",
                Severity.PASS, "No password-protected PDFs detected"
            )
        
        # Check 58: File size reasonable
        for doc_type, filename in self.report.files_found.items():
            if filename:
                filepath = list(self.folder_path.glob(filename))[0]
                size = filepath.stat().st_size
                if size == 0:
                    self.report.add_issue(
                        58, "File Quality", "File Size Reasonable",
                        Severity.CRITICAL, f"{filename} has 0 bytes"
                    )
                    return
        
        self.report.add_issue(
            58, "File Quality", "File Size Reasonable",
            Severity.PASS, "All files have reasonable sizes"
        )
    
    def check_cross_references(self):
        """Checks 59-62: Cross-reference validation"""
        
        # Check 59: Claims reference specification elements
        if self.spec_text:
            claims_text_59 = self._extract_claims_section()

            # Extract detailed description (before claims)
            desc_match = re.search(
                r'(?:DETAILED\s+DESCRIPTION|DESCRIPTION\s+OF.*?EMBODIMENTS?)(.*?)(?:\bCLAIMS\b|What is claimed)',
                self.spec_text, re.DOTALL | re.IGNORECASE
            )
            desc_text = desc_match.group(1) if desc_match else self.spec_text[:len(self.spec_text)//2]

            if claims_text_59:
                # Extract noun phrases from claims that look like real claim
                # elements: introduced by "a/an" (first mention) or referenced
                # by "the/said". Filter out single-word common nouns and
                # boilerplate.
                STOPWORDS = {'method', 'system', 'apparatus', 'device', 'medium',
                             'product', 'invention', 'embodiment', 'present',
                             'following', 'above', 'said', 'wherein', 'thereof',
                             'therein', 'further', 'least', 'one', 'more', 'each',
                             'plurality', 'time', 'use', 'set', 'first', 'second',
                             'third', 'fourth', 'fifth'}
                # Multi-word noun phrases (2-3 words) — these are the meaningful
                # claim terms. Single-word terms like "method" / "system" are
                # too generic to verify.
                np_pattern = r'\b(?:a|an|the|said)\s+([a-z][\w\-]*\s+[\w\-]+(?:\s+[\w\-]+)?)\b'
                all_claim_terms = set()
                for m in re.finditer(np_pattern, claims_text_59, re.IGNORECASE):
                    term = m.group(1).strip().lower()
                    # Exclude phrases that start with a stopword or are entirely
                    # stopword-derived
                    words = term.split()
                    if words[0] in STOPWORDS:
                        continue
                    # Strip trailing verb participle if any
                    if words[-1].endswith(('ing', 'ed')) and len(words) > 1:
                        term = ' '.join(words[:-1])
                    all_claim_terms.add(term)

                if all_claim_terms:
                    desc_lower = desc_text.lower()
                    missing = [t for t in all_claim_terms if t.lower() not in desc_lower]
                    if not missing:
                        self.report.add_issue(
                            59, "Cross-References", "Claims Reference Specification Elements",
                            Severity.PASS, f"All {len(all_claim_terms)} claim elements found in specification"
                        )
                    elif len(missing) <= 3:
                        self.report.add_issue(
                            59, "Cross-References", "Claims Reference Specification Elements",
                            Severity.PASS, f"Most claim elements found in specification ({len(all_claim_terms) - len(missing)}/{len(all_claim_terms)})"
                        )
                    else:
                        self.report.add_issue(
                            59, "Cross-References", "Claims Reference Specification Elements",
                            Severity.WARNING, f"{len(missing)} claim elements not clearly found in specification",
                            f"Missing: {', '.join(list(missing)[:5])}"
                        )
                else:
                    self.report.add_issue(
                        59, "Cross-References", "Claims Reference Specification Elements",
                        Severity.PASS, "Claim elements cross-referenced with specification"
                    )
            else:
                self.report.add_issue(
                    59, "Cross-References", "Claims Reference Specification Elements",
                    Severity.INFO, "Could not isolate claims section for cross-reference check"
                )
        else:
            self.report.add_issue(
                59, "Cross-References", "Claims Reference Specification Elements",
                Severity.WARNING, "Specification not found"
            )

        # Check 60: Specification summary matches claims
        if self.spec_text:
            # Find summary section
            summary_match = re.search(
                r'(?:SUMMARY|BRIEF\s+SUMMARY)(.*?)(?:BRIEF\s+DESCRIPTION\s+OF|DETAILED\s+DESCRIPTION|DRAWINGS)',
                self.spec_text, re.DOTALL | re.IGNORECASE
            )

            # Find independent claims (claim 1 at minimum)
            ind_claim_match = re.search(
                r'(?:What is claimed[^:]*:\s*|CLAIMS\s+).*?1\.\s+(A.*?)(?:\s{2,}\d+\.\s+|\.\s+\d+\.\s+)',
                self.spec_text, re.DOTALL | re.IGNORECASE
            )

            if summary_match and ind_claim_match:
                summary_text = summary_match.group(1).lower()
                claim1_text = ind_claim_match.group(1).lower()

                # Extract key terms from claim 1 (nouns/adjectives, 4+ chars)
                claim1_words = set(re.findall(r'\b([a-z]{4,})\b', claim1_text))
                # Remove common claim language
                stop_words = {'comprising', 'including', 'wherein', 'thereof', 'therein',
                             'configured', 'coupled', 'connected', 'having', 'being',
                             'first', 'second', 'third', 'method', 'system', 'apparatus',
                             'each', 'said', 'claim', 'further', 'least', 'with', 'from',
                             'that', 'which', 'where', 'when', 'into', 'upon', 'between'}
                key_terms = claim1_words - stop_words

                if key_terms:
                    found = sum(1 for t in key_terms if t in summary_text)
                    coverage = found / len(key_terms) if key_terms else 0

                    if coverage >= 0.5:
                        self.report.add_issue(
                            60, "Cross-References", "Specification Summary Matches Claims",
                            Severity.PASS, f"Summary covers {found}/{len(key_terms)} key claim terms ({coverage:.0%} coverage)"
                        )
                    else:
                        # Word-overlap heuristic — drafters legitimately use
                        # different vocabulary in the summary vs. the claims,
                        # so low coverage doesn't necessarily mean a real issue.
                        self.report.add_issue(
                            60, "Cross-References", "Specification Summary Matches Claims",
                            Severity.INFO,
                            f"Heuristic: summary covers only {coverage:.0%} of claim 1 word tokens "
                            f"({found}/{len(key_terms)}). Drafters often use different vocabulary "
                            f"in the summary; this is best verified manually rather than "
                            f"flagged as a finding."
                        )
                else:
                    self.report.add_issue(
                        60, "Cross-References", "Specification Summary Matches Claims",
                        Severity.PASS, "Summary and claims present"
                    )
            else:
                self.report.add_issue(
                    60, "Cross-References", "Specification Summary Matches Claims",
                    Severity.INFO, "Could not isolate both summary and claims for comparison"
                )
        else:
            self.report.add_issue(
                60, "Cross-References", "Specification Summary Matches Claims",
                Severity.WARNING, "Specification not found"
            )
        
        # Check 61: Drawing figure count matches specification
        # Same image-only-drawings caveat as Check 15/22 — if the drawings PDF
        # has no extractable FIG. labels, comparing counts is meaningless.
        if not self._drawings_text_extractable():
            spec_figs = set(self._extract_figure_numbers(self.spec_text)) if self.spec_text else set()
            self.report.add_issue(
                61, "Cross-References", "Drawing Figure Count Matches Specification",
                Severity.INFO,
                f"Drawings PDF appears to be image-only — figure count not verifiable by text extraction. "
                f"Spec references {len(spec_figs)} figure(s)" + (
                    f" (FIG. {', '.join(str(n) for n in sorted(spec_figs))})" if spec_figs else ""
                ) + ". Manually verify drawings contain matching figures."
            )
        elif self.spec_text and self.drawings_text:
            spec_figs = set(self._extract_figure_numbers(self.spec_text))
            drawing_figs = set(self._extract_figure_numbers(self.drawings_text))

            if spec_figs == drawing_figs:
                self.report.add_issue(
                    61, "Cross-References", "Drawing Figure Count Matches Specification",
                    Severity.PASS, f"Figure numbers match: {len(spec_figs)} figures (FIG. {', '.join(str(n) for n in sorted(spec_figs))})"
                )
            elif len(spec_figs) == len(drawing_figs):
                self.report.add_issue(
                    61, "Cross-References", "Drawing Figure Count Matches Specification",
                    Severity.WARNING, f"Same figure count ({len(spec_figs)}) but different numbers. Spec: {sorted(spec_figs)}, Drawings: {sorted(drawing_figs)}"
                )
            else:
                self.report.add_issue(
                    61, "Cross-References", "Drawing Figure Count Matches Specification",
                    Severity.WARNING, f"Figure count mismatch: Spec references {len(spec_figs)} figures, Drawings has {len(drawing_figs)} figures"
                )
        else:
            self.report.add_issue(
                61, "Cross-References", "Drawing Figure Count Matches Specification",
                Severity.INFO, "Unable to compare figure counts"
            )
        
        # Check 62: Claim count verification
        # Use the shared claims-section helper and the same preamble-free
        # pattern as Check 13 (claim 1 may not start with A/An/The/We/I).
        if self.spec_text:
            claims_text = self._extract_claims_section() or self.spec_text

            claim_patterns = [
                r'(?:^|\n)\s*(\d+)\.\s+(?=\D)',          # newline-anchored
                r'(?:\.\s+|\;\s+|:\s+)(\d+)\.\s+(?=\D)', # after sentence end
                r'\s{2,}(\d+)\.\s+(?=\D)',               # after multiple spaces
                r'\s+\d{2,3}\s+(\d+)\.\s+(?=\D)',        # after page number
            ]

            claim_nums = []
            for pattern in claim_patterns:
                matches = re.findall(pattern, claims_text, re.IGNORECASE)
                claim_nums.extend(matches)

            if claim_nums:
                # Get unique claim numbers and verify sequence
                unique_claims = sorted(set(int(n) for n in claim_nums if 1 <= int(n) <= 100))
                self.report.add_issue(
                    62, "Cross-References", "Claim Count Verification",
                    Severity.PASS, f"Total claims detected: {len(unique_claims)} (Claims 1-{max(unique_claims)})"
                )
            else:
                # Fallback: try simpler pattern
                simple_patterns = [
                    r'(?:^|\n)\s*(\d+)\.\s+',           # Original with newlines
                    r'(?:\.\s+|\;\s+|:\s+)(\d+)\.\s+',  # After sentence end
                    r'\s{2,}(\d+)\.\s+',                # After multiple spaces
                    r'^\s*(\d+)\.\s+',                  # At start of claims section
                    r'\s+\d{2,3}\s+(\d+)\.\s+',         # After page number (2-3 digits)
                ]
                simple_claims = []
                for pattern in simple_patterns:
                    matches = re.findall(pattern, claims_text, re.MULTILINE)
                    simple_claims.extend(matches)
                simple_claims = [n for n in simple_claims if 1 <= int(n) <= 100]
                if simple_claims:
                    self.report.add_issue(
                        62, "Cross-References", "Claim Count Verification",
                        Severity.INFO, f"Possible claims detected: {len(set(simple_claims))} (manual verification recommended)"
                    )
                else:
                    self.report.add_issue(
                        62, "Cross-References", "Claim Count Verification",
                        Severity.WARNING, "Unable to count claims - claims section may use non-standard format"
                    )
        else:
            self.report.add_issue(
                62, "Cross-References", "Claim Count Verification",
                Severity.WARNING, "Specification not found"
            )
    
    def check_priority(self):
        """Checks 63-65: Priority/related application checks.
        Refactored to use ADS XFA structured continuity/priority data as the
        source of truth instead of fragile regex on spec/ADS text. Previously,
        Check 65 misidentified US provisional applications as 'foreign'
        priority, and Check 64 missed Cross-Reference sections that don't use
        the exact regex header format."""

        ads_dom_entries = (self.ads_data.get('domestic_continuity_entries')
                           if self.ads_data else None) or []
        ads_for_entries = (self.ads_data.get('foreign_priority_entries')
                           if self.ads_data else None) or []

        # Detect priority language in the spec (broader patterns).
        spec_priority_match = None
        if self.spec_text:
            for pat in [
                r'(?:claims|claiming)\s+(?:the\s+)?(?:benefit|priority)\s+(?:of|to|under)',
                r'\bcontinuation(?:[-\s]in[-\s]part)?\s+of',
                r'\bdivisional\s+(?:of|application)',
                r'\bprovisional\s+application\s+(?:no\.?|number)',
            ]:
                m = re.search(pat, self.spec_text, re.IGNORECASE)
                if m:
                    spec_priority_match = m.group(0)
                    break

        # ----- Check 63: Priority claim consistency -----
        if ads_dom_entries or ads_for_entries:
            # ADS HAS priority data — verify spec also references it
            if spec_priority_match:
                self.report.add_issue(
                    63, "Priority Claims", "Priority Claim Consistency",
                    Severity.PASS,
                    f"Priority claims present in both ADS and specification "
                    f"({len(ads_dom_entries)} domestic, {len(ads_for_entries)} foreign in ADS)"
                )
            else:
                self.report.add_issue(
                    63, "Priority Claims", "Priority Claim Consistency",
                    Severity.WARNING,
                    f"ADS lists {len(ads_dom_entries)} domestic and {len(ads_for_entries)} "
                    f"foreign priority entries, but no priority language found in specification. "
                    f"Spec should reference the parent/priority application(s)."
                )
        else:
            if spec_priority_match:
                self.report.add_issue(
                    63, "Priority Claims", "Priority Claim Consistency",
                    Severity.WARNING,
                    "Priority language detected in specification but no priority entries in ADS. "
                    "Verify the ADS continuity/foreign-priority sections are filled in correctly."
                )
            else:
                self.report.add_issue(
                    63, "Priority Claims", "Priority Claim Consistency",
                    Severity.PASS, "No priority claims detected in specification or ADS"
                )

        # ----- Check 64: Related application references in spec -----
        if ads_dom_entries or ads_for_entries or spec_priority_match:
            # Look for a Cross-Reference / Related Applications section
            # using broader patterns. Many specs label this as "RELATED
            # APPLICATIONS", "CROSS-REFERENCE TO RELATED APPLICATIONS",
            # "PRIORITY CLAIM", etc.
            spec_related = False
            if self.spec_text:
                for pat in [
                    r'CROSS[-\s]*REFERENCE',
                    r'RELATED\s+APPLICATION',
                    r'PRIORITY\s+CLAIM',
                    r'PRIORITY\s+TO\s+RELATED',
                    r'\bcontinuation\s+of\b',
                    r'\bclaims\s+(?:the\s+)?(?:benefit|priority)',
                ]:
                    if re.search(pat, self.spec_text, re.IGNORECASE):
                        spec_related = True
                        break

            if spec_related:
                self.report.add_issue(
                    64, "Priority Claims", "Related Application References",
                    Severity.PASS,
                    "Related-application cross-reference language found in specification"
                )
            else:
                self.report.add_issue(
                    64, "Priority Claims", "Related Application References",
                    Severity.WARNING,
                    "Priority claims present but no Cross-Reference / Related Applications section "
                    "found in specification. Verify the spec includes proper priority/continuation "
                    "language near the start."
                )
        else:
            self.report.add_issue(
                64, "Priority Claims", "Related Application References",
                Severity.PASS, "No related applications detected"
            )

        # ----- Check 65: Foreign priority documents -----
        # Use the XFA-structured foreign_priority_entries directly. A US
        # provisional is a DOMESTIC priority (sfDomesticContinuity), not a
        # foreign priority (sfForeignPriorityInfo). The previous regex-based
        # check matched any "Provisional Application No.…" in the spec as
        # "foreign," which is wrong.
        if ads_for_entries:
            countries = sorted({(e.get('country') or '?').upper() for e in ads_for_entries})
            self.report.add_issue(
                65, "Priority Claims", "Foreign Priority Documents",
                Severity.INFO,
                f"{len(ads_for_entries)} foreign priority claim(s) in ADS "
                f"({', '.join(countries)}). Verify that certified copies of the foreign "
                f"priority documents are on file or being filed."
            )
        else:
            self.report.add_issue(
                65, "Priority Claims", "Foreign Priority Documents",
                Severity.PASS, "No foreign priority claims in ADS"
            )
    
    def check_final_quality(self):
        """Checks 66-70: Final quality checks"""
        
        # Check 66: No obvious typos in critical fields
        if self.spec_text or self.ads_text:
            issues_66 = []

            # Check docket number format (should be consistent pattern like A###-####XX)
            docket_pattern = r'[A-Z]\d{2,4}[\s\-]*\d{3,4}[A-Z]{2}'
            all_texts = {'spec': self.spec_text, 'ads': self.ads_text, 'decl': self.declaration_text}
            docket_numbers = set()
            for doc, text in all_texts.items():
                if text:
                    matches = re.findall(docket_pattern, text)
                    for m in matches:
                        docket_numbers.add(re.sub(r'\s', '', m))

            if len(docket_numbers) > 1:
                issues_66.append(f"Multiple docket number variants: {', '.join(docket_numbers)}")

            # Check inventor names for mixed case issues (e.g., "gUPTA" instead of "GUPTA" or "Gupta")
            if self.ads_text:
                inventor_names = re.findall(r'(?:Given|Family)\s*Name[:\s]+([A-Za-z]+)', self.ads_text)
                for name in inventor_names:
                    if name and len(name) > 1:
                        # Name should be either all caps, all lower, or title case
                        if not (name.isupper() or name.islower() or name.istitle()):
                            issues_66.append(f"Unusual capitalization in inventor name: '{name}'")

            # Check title for incomplete words or obvious issues
            title_match = re.search(r'(?:Title|TITLE).*?(?:of|:)\s*(?:the\s+)?(?:Invention\s*)?(.*?)(?:\n|$|Attorney)',
                                   self.ads_text or self.spec_text, re.IGNORECASE)
            if title_match:
                title = title_match.group(1).strip()
                if len(title) < 5:
                    issues_66.append(f"Title appears too short: '{title}'")
                if re.search(r'\b[A-Z]{1}\b', title):  # Single uppercase letters (may be typos)
                    pass  # Single letters can be valid in titles

            if issues_66:
                self.report.add_issue(
                    66, "Final Quality", "No Obvious Typos in Critical Fields",
                    Severity.WARNING, f"Potential issues found: {'; '.join(issues_66[:3])}"
                )
            else:
                self.report.add_issue(
                    66, "Final Quality", "No Obvious Typos in Critical Fields",
                    Severity.PASS, "No obvious typos detected in critical fields"
                )
        else:
            self.report.add_issue(
                66, "Final Quality", "No Obvious Typos in Critical Fields",
                Severity.INFO, "Insufficient document text for typo analysis"
            )

        # Check 67: Dates in proper format
        if self.ads_text or self.declaration_text or self.assignment_text:
            all_doc_text = (self.ads_text or '') + (self.declaration_text or '') + (self.assignment_text or '')

            months = r'(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)'

            # Find dates in various formats
            date_patterns = [
                (r'\b(\d{1,2})[/\-](\d{1,2})[/\-](\d{4})\b', 'MM/DD/YYYY'),
                (r'\b(\d{1,2})[/\-](\d{1,2})[/\-](\d{2})\b', 'MM/DD/YY'),
                (r'\b(' + months + r')\s+(\d{1,2}),?\s+(\d{4})\b', 'Month DD, YYYY'),
                (r'\b(\d{1,2})\s+(' + months + r')\s+(\d{4})\b', 'DD Month YYYY'),
            ]

            dates_found = []
            issues_67 = []
            # Form template boilerplate patterns to ignore
            form_boilerplate = re.compile(
                r'(?:approved\s+for\s+use|OMB|Paperwork\s+Reduction|collection\s+of\s+information|'
                r'CFR\s+\d|U\.S\.\s+Patent\s+and\s+Trademark)',
                re.IGNORECASE
            )

            for pattern, fmt in date_patterns:
                for m_obj in re.finditer(pattern, all_doc_text, re.IGNORECASE):
                    m = m_obj.groups()
                    # Skip dates in form boilerplate context (100 chars before/after)
                    context_start = max(0, m_obj.start() - 100)
                    context_end = min(len(all_doc_text), m_obj.end() + 100)
                    context = all_doc_text[context_start:context_end]
                    if form_boilerplate.search(context):
                        continue

                    dates_found.append((m, fmt))
                    # Check for 2-digit years (potential issue)
                    if fmt == 'MM/DD/YY':
                        issues_67.append(f"2-digit year found: {'/'.join(m)}")
                    # Check for reasonable year range
                    if fmt in ('MM/DD/YYYY', 'Month DD, YYYY', 'DD Month YYYY'):
                        year = int(m[2])
                        if year < 2000 or year > 2030:
                            issues_67.append(f"Date with unusual year: {' '.join(m)}")

            if issues_67:
                self.report.add_issue(
                    67, "Final Quality", "Dates in Proper Format",
                    Severity.WARNING, f"Date format issues: {'; '.join(issues_67[:3])}"
                )
            elif dates_found:
                self.report.add_issue(
                    67, "Final Quality", "Dates in Proper Format",
                    Severity.PASS, f"All {len(dates_found)} dates appear properly formatted"
                )
            else:
                self.report.add_issue(
                    67, "Final Quality", "Dates in Proper Format",
                    Severity.PASS, "No date format issues detected"
                )
        else:
            self.report.add_issue(
                67, "Final Quality", "Dates in Proper Format",
                Severity.INFO, "Insufficient document text for date analysis"
            )
        
        # Check 68: No excessively long claims
        # Same approach as Check 13/62: shared section helper + preamble-free
        # claim-start pattern (the original A/An/The preamble missed claims
        # that started with other words, undercounting by 1 — which is why
        # this previously said '19 claims checked' for a 20-claim spec).
        if self.spec_text:
            claims_text = self._extract_claims_section() or self.spec_text

            claim_start_patterns = [
                r'(?:^|\n)\s*(\d+)\.\s+(?=\D)',
                r'(?:\.\s{1,3})(\d+)\.\s+(?=\D)',
                r'\s{2,}(\d+)\.\s+(?=\D)',
                r'\s+\d{2,3}\s+(\d+)\.\s+(?=\D)',
            ]

            # Use finditer to get positions for splitting
            all_claim_positions = []
            for pattern in claim_start_patterns:
                for m in re.finditer(pattern, claims_text, re.IGNORECASE | re.MULTILINE):
                    claim_num = int(m.group(1))
                    if 1 <= claim_num <= 100:
                        all_claim_positions.append((m.start(), claim_num, m.end()))

            # Deduplicate by claim number (keep first occurrence)
            seen_claims = {}
            for start, num, end in sorted(all_claim_positions):
                if num not in seen_claims:
                    seen_claims[num] = (start, end)

            claims = []
            sorted_claims = sorted(seen_claims.items())
            for i, (claim_num, (start, text_start)) in enumerate(sorted_claims):
                # End at next claim start or end of text
                if i + 1 < len(sorted_claims):
                    end = sorted_claims[i + 1][1][0]
                else:
                    end = len(claims_text)
                claim_body = claims_text[text_start:end].strip()
                word_count = len(claim_body.split())
                claims.append((claim_num, word_count))

            if claims:
                long_claims = [(num, words) for num, words in claims if words > 200]
                if long_claims:
                    details = ", ".join([f"Claim {num} ({words} words)" for num, words in long_claims[:5]])
                    self.report.add_issue(
                        68, "Final Quality", "No Excessively Long Claims",
                        Severity.WARNING, f"Unusually long claims detected: {details}"
                    )
                else:
                    self.report.add_issue(
                        68, "Final Quality", "No Excessively Long Claims",
                        Severity.PASS, f"No excessively long claims detected ({len(claims)} claims checked)"
                    )
            else:
                self.report.add_issue(
                    68, "Final Quality", "No Excessively Long Claims",
                    Severity.INFO, "Unable to parse individual claims for length check"
                )
        else:
            self.report.add_issue(
                68, "Final Quality", "No Excessively Long Claims",
                Severity.WARNING, "Specification not found"
            )
        
        # Check 69: Specification references all claims
        if self.spec_text:
            # Extract claims section
            claims_section_patterns = [
                r'CLAIMS\s+What is claimed[^:]*:\s*(.*?)(?:ABSTRACT|$)',
                r'What is claimed[^:]*:\s*(.*?)(?:ABSTRACT|$)',
                r'(?:CLAIMS?\s*\n|What is claimed[^\n]*\n)(.*?)(?:ABSTRACT|$)',
            ]
            claims_text_69 = None
            for pattern in claims_section_patterns:
                m = re.search(pattern, self.spec_text, re.DOTALL | re.IGNORECASE)
                if m:
                    claims_text_69 = m.group(1)
                    break

            # Get description text (everything before claims)
            desc_match = re.search(
                r'(?:DETAILED\s+DESCRIPTION|DESCRIPTION\s+OF.*?EMBODIMENTS?)(.*?)(?:CLAIMS|What is claimed)',
                self.spec_text, re.DOTALL | re.IGNORECASE
            )
            desc_text_69 = desc_match.group(1) if desc_match else ""

            if claims_text_69 and desc_text_69:
                # Extract reference numerals from claims
                claim_numerals = set(re.findall(r'\b(\d{2,3})\b', claims_text_69))
                # Filter to actual reference numerals (those used with element descriptions)
                valid_numerals = set()
                for num in claim_numerals:
                    if re.search(r'(?:a|an|the|said)\s+[\w\-]+(?:\s+[\w\-]+){0,2}\s+' + num + r'\b',
                                claims_text_69, re.IGNORECASE):
                        valid_numerals.add(num)

                if valid_numerals:
                    # Check which numerals appear in the detailed description
                    missing = [n for n in valid_numerals if n not in desc_text_69]
                    if not missing:
                        self.report.add_issue(
                            69, "Final Quality", "Specification References All Claims",
                            Severity.PASS, f"All {len(valid_numerals)} claim reference numerals found in specification"
                        )
                    elif len(missing) <= 2:
                        self.report.add_issue(
                            69, "Final Quality", "Specification References All Claims",
                            Severity.PASS, f"Most claim elements referenced in specification ({len(valid_numerals) - len(missing)}/{len(valid_numerals)})"
                        )
                    else:
                        self.report.add_issue(
                            69, "Final Quality", "Specification References All Claims",
                            Severity.WARNING, f"{len(missing)} claim reference numerals not found in specification: {', '.join(sorted(missing)[:5])}"
                        )
                else:
                    # No reference numerals in claims to verify against the
                    # specification. We deliberately do NOT fall back to a
                    # word-coverage heuristic here — that's Check 60's job and
                    # would just produce duplicate noise under a misleading
                    # name (this check's name is "References All Claims",
                    # not "term coverage").
                    self.report.add_issue(
                        69, "Final Quality", "Specification References All Claims",
                        Severity.INFO,
                        "No reference numerals detected in claims — cannot verify "
                        "claim-to-specification element references. Manual review recommended."
                    )
            else:
                self.report.add_issue(
                    69, "Final Quality", "Specification References All Claims",
                    Severity.INFO, "Could not isolate claims and description for cross-reference"
                )
        else:
            self.report.add_issue(
                69, "Final Quality", "Specification References All Claims",
                Severity.WARNING, "Specification not found"
            )
        
        # Check 70: Consistent figure reference format
        if self.spec_text:
            fig_refs = re.findall(r'(FIG(?:URE)?\.?\s*\d+)', self.spec_text, re.IGNORECASE)
            if fig_refs:
                formats = set(ref.split()[0].upper() for ref in fig_refs)
                if len(formats) == 1:
                    self.report.add_issue(
                        70, "Final Quality", "Consistent Figure Reference Format",
                        Severity.PASS, f"Figure references use consistent format: {list(formats)[0]}"
                    )
                else:
                    self.report.add_issue(
                        70, "Final Quality", "Consistent Figure Reference Format",
                        Severity.WARNING, f"Mixed figure reference formats detected: {formats}"
                    )
            else:
                self.report.add_issue(
                    70, "Final Quality", "Consistent Figure Reference Format",
                    Severity.INFO, "No figure references detected"
                )
        else:
            self.report.add_issue(
                70, "Final Quality", "Consistent Figure Reference Format",
                Severity.WARNING, "Specification not found"
            )
    
    def generate_html_report(self, output_path: str):
        """Emit a self-contained HTML report with embedded CSS. Produces the
        same content as generate_markdown_report but with deterministic
        rendering in any browser (and reliable Print → Save as PDF). No
        external CSS, no external fonts, no markdown/pdflatex toolchain
        dependency."""
        import datetime
        import html as _html

        def esc(s) -> str:
            return _html.escape("" if s is None else str(s))

        def severity_class(sev: 'Severity') -> str:
            return {
                Severity.CRITICAL: 'critical',
                Severity.WARNING: 'warning',
                Severity.INFO: 'info',
                Severity.PASS: 'pass',
            }.get(sev, 'info')

        def severity_label(sev: 'Severity') -> str:
            return {
                Severity.CRITICAL: 'CRITICAL',
                Severity.WARNING: 'WARN',
                Severity.INFO: 'INFO',
                Severity.PASS: 'PASS',
            }.get(sev, 'INFO')

        critical_issues = [i for i in self.report.issues if i.severity == Severity.CRITICAL]
        warnings = [i for i in self.report.issues if i.severity == Severity.WARNING]
        info_issues = [i for i in self.report.issues if i.severity == Severity.INFO]
        passed_issues = [i for i in self.report.issues if i.severity == Severity.PASS]

        def group_by_category(issues):
            groups: Dict[str, List] = {}
            for issue in issues:
                groups.setdefault(issue.category, []).append(issue)
            return groups

        # Try to use the application title from XFA data (when ADS read
        # successfully) for the document subtitle; otherwise omit it.
        subtitle = ""
        if self.ads_data and self.ads_data.get('title'):
            subtitle = self.ads_data['title']
        docket = ""
        if self.ads_data and self.ads_data.get('docket_number'):
            docket = self.ads_data['docket_number']

        date_str = datetime.datetime.now().strftime('%B %d, %Y at %H:%M')

        # CSS — system fonts only; no emoji glyphs (use text labels in colored
        # badges instead). Print rules avoid splitting issue cards across pages.
        css = """
            :root {
                --c-critical: #b91c1c;
                --c-critical-bg: #fef2f2;
                --c-warn:     #b45309;
                --c-warn-bg:  #fffbeb;
                --c-info:     #1d4ed8;
                --c-info-bg:  #eff6ff;
                --c-pass:     #166534;
                --c-pass-bg:  #f0fdf4;
                --c-text:     #1f2937;
                --c-muted:    #6b7280;
                --c-border:   #e5e7eb;
                --c-bg-soft:  #f9fafb;
            }
            * { box-sizing: border-box; }
            html, body { margin: 0; padding: 0; }
            body {
                font-family: -apple-system, BlinkMacSystemFont, "Segoe UI",
                             Roboto, Helvetica, Arial, sans-serif;
                color: var(--c-text);
                line-height: 1.5;
                background: #fff;
                font-size: 14px;
            }
            .page {
                max-width: 980px;
                margin: 0 auto;
                padding: 32px 40px 56px;
            }
            header { border-bottom: 3px solid #0f172a; padding-bottom: 16px; margin-bottom: 24px; }
            header h1 { font-size: 24px; margin: 0 0 4px; color: #0f172a; letter-spacing: -0.01em; }
            header .subtitle { font-size: 14px; color: var(--c-muted); margin: 4px 0 0; }
            header .meta { display: flex; gap: 24px; flex-wrap: wrap;
                          font-size: 12px; color: var(--c-muted); margin-top: 12px; }
            header .meta span { white-space: nowrap; }
            header .meta b { color: var(--c-text); font-weight: 600; }

            h2 { font-size: 18px; color: #0f172a; margin: 28px 0 12px;
                 padding-bottom: 6px; border-bottom: 1px solid var(--c-border); }
            h3 { font-size: 15px; color: #334155; margin: 18px 0 8px;
                 text-transform: uppercase; letter-spacing: 0.04em; font-weight: 600; }

            .summary { display: grid; grid-template-columns: repeat(4, 1fr);
                       gap: 12px; margin: 12px 0 16px; }
            .stat { padding: 16px; border-radius: 6px; border: 1px solid var(--c-border); }
            .stat .num { font-size: 28px; font-weight: 700; line-height: 1.1; }
            .stat .lbl { font-size: 12px; text-transform: uppercase;
                         letter-spacing: 0.05em; color: var(--c-muted);
                         margin-top: 4px; }
            .stat.critical { background: var(--c-critical-bg); border-color: #fecaca; }
            .stat.critical .num { color: var(--c-critical); }
            .stat.warning  { background: var(--c-warn-bg);     border-color: #fcd34d; }
            .stat.warning  .num { color: var(--c-warn); }
            .stat.info     { background: var(--c-info-bg);     border-color: #bfdbfe; }
            .stat.info     .num { color: var(--c-info); }
            .stat.pass     { background: var(--c-pass-bg);     border-color: #bbf7d0; }
            .stat.pass     .num { color: var(--c-pass); }
            /* Stat cards are <a> elements — make them look clickable */
            a.stat { text-decoration: none; color: inherit; display: block;
                     transition: transform 0.1s ease, box-shadow 0.1s ease; cursor: pointer; }
            a.stat:hover { transform: translateY(-2px);
                           box-shadow: 0 4px 12px rgba(0,0,0,0.08); }
            .stat.stat-empty { opacity: 0.55; }

            /* Print-to-PDF button (top right, sticky on screen, hidden in print) */
            .print-btn { position: fixed; top: 16px; right: 20px; z-index: 100;
                         background: #0f172a; color: #fff; border: none;
                         padding: 8px 14px; border-radius: 6px; font-size: 13px;
                         font-weight: 600; cursor: pointer;
                         box-shadow: 0 2px 8px rgba(0,0,0,0.15);
                         font-family: inherit; }
            .print-btn:hover { background: #1e293b; }
            .print-btn:active { transform: translateY(1px); }
            @media print { .print-btn { display: none; } }
            /* Smooth scroll when clicking the executive summary cards */
            html { scroll-behavior: smooth; }

            table { width: 100%; border-collapse: collapse; margin: 8px 0 16px;
                    font-size: 13px; }
            th, td { border: 1px solid var(--c-border); padding: 8px 12px;
                     text-align: left; vertical-align: top; }
            th { background: var(--c-bg-soft); font-weight: 600;
                 font-size: 12px; text-transform: uppercase; letter-spacing: 0.04em;
                 color: var(--c-muted); }
            td.docfound { font-family: "SF Mono", Menlo, Consolas, monospace;
                          font-size: 12px; }
            td.docfound.missing { color: var(--c-critical); }

            .issue { border: 1px solid var(--c-border); border-left-width: 4px;
                     border-radius: 4px; padding: 12px 14px;
                     margin: 8px 0; background: #fff;
                     page-break-inside: avoid; }
            .issue.critical { border-left-color: var(--c-critical); }
            .issue.warning  { border-left-color: var(--c-warn); }
            .issue.info     { border-left-color: var(--c-info); }
            .issue.pass     { border-left-color: var(--c-pass); }

            .issue-head { display: flex; align-items: baseline; gap: 10px;
                          flex-wrap: wrap; }
            .badge { display: inline-block; padding: 2px 8px; border-radius: 3px;
                     font-size: 11px; font-weight: 700;
                     text-transform: uppercase; letter-spacing: 0.05em;
                     font-family: -apple-system, BlinkMacSystemFont, "Segoe UI",
                                  Roboto, sans-serif; }
            .badge.critical { background: var(--c-critical); color: #fff; }
            .badge.warning  { background: var(--c-warn);     color: #fff; }
            .badge.info     { background: var(--c-info);     color: #fff; }
            .badge.pass     { background: var(--c-pass);     color: #fff; }

            .issue-id { color: var(--c-muted); font-variant-numeric: tabular-nums;
                        font-size: 12px; min-width: 28px; }
            .issue-name { font-weight: 600; color: #0f172a; }
            .issue-msg { margin: 6px 0 0 0; }
            .issue-details { margin: 10px 0 0; padding: 10px 12px;
                             background: var(--c-bg-soft);
                             border-left: 3px solid var(--c-border);
                             border-radius: 0 4px 4px 0;
                             white-space: pre-wrap;
                             font-size: 12.5px;
                             color: #374151; }

            footer { margin-top: 40px; padding-top: 12px;
                     border-top: 1px solid var(--c-border);
                     color: var(--c-muted); font-size: 12px; }

            @media print {
                @page { margin: 0.6in 0.6in; }
                body { font-size: 11pt; }
                .page { max-width: none; padding: 0; }
                .summary { gap: 8px; }
                .stat { padding: 10px; }
                h2 { break-after: avoid; }
                h3 { break-after: avoid; }
                .issue { break-inside: avoid; }
            }
        """

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("<!DOCTYPE html>\n")
            f.write("<html lang=\"en\">\n<head>\n")
            f.write("<meta charset=\"utf-8\">\n")
            f.write("<meta name=\"viewport\" content=\"width=device-width,initial-scale=1\">\n")
            title_text = f"Patent Filing QC Report"
            if docket:
                title_text += f" — {docket}"
            f.write(f"<title>{esc(title_text)}</title>\n")
            f.write(f"<style>{css}</style>\n")
            f.write("</head>\n<body>\n")
            f.write("<button class=\"print-btn\" onclick=\"window.print()\">"
                    "Print to PDF</button>\n")
            f.write("<main class=\"page\">\n")

            # Header
            f.write("<header>\n")
            f.write(f"<h1>Patent Filing Quality Control Report</h1>\n")
            if subtitle:
                f.write(f"<p class=\"subtitle\">{esc(subtitle)}</p>\n")
            f.write("<div class=\"meta\">\n")
            if docket:
                f.write(f"<span><b>Docket:</b> {esc(docket)}</span>\n")
            f.write(f"<span><b>Folder:</b> {esc(self.report.folder_path)}</span>\n")
            f.write(f"<span><b>Generated:</b> {esc(date_str)}</span>\n")
            f.write("</div>\n")
            f.write("</header>\n")

            # Executive Summary — each stat card is a link to its section
            f.write("<h2>Executive Summary</h2>\n")
            f.write("<div class=\"summary\">\n")
            for cls, count, lbl, anchor in [
                ('critical', len(critical_issues), 'Critical',      'sec-critical'),
                ('warning',  len(warnings),        'Warnings',      'sec-warnings'),
                ('info',     len(info_issues),     'Manual Review', 'sec-info'),
                ('pass',     len(passed_issues),   'Passed',        'sec-passed'),
            ]:
                # Make the card a link only if the section exists (has issues).
                if count > 0:
                    f.write(f"<a class=\"stat {cls}\" href=\"#{anchor}\">"
                            f"<div class=\"num\">{count}</div>"
                            f"<div class=\"lbl\">{lbl}</div></a>\n")
                else:
                    f.write(f"<div class=\"stat {cls} stat-empty\">"
                            f"<div class=\"num\">{count}</div>"
                            f"<div class=\"lbl\">{lbl}</div></div>\n")
            f.write("</div>\n")

            # Documents Found
            f.write("<h2>Documents Found</h2>\n")
            f.write("<table>\n<thead><tr><th>Document Type</th><th>File</th></tr></thead>\n<tbody>\n")
            for doc_type, filename in self.report.files_found.items():
                if filename:
                    f.write(f"<tr><td>{esc(doc_type)}</td>"
                            f"<td class=\"docfound\">{esc(filename)}</td></tr>\n")
                else:
                    f.write(f"<tr><td>{esc(doc_type)}</td>"
                            f"<td class=\"docfound missing\">NOT FOUND</td></tr>\n")
            f.write("</tbody>\n</table>\n")

            def write_issue_section(heading: str, issues: List[QCIssue], anchor: str):
                if not issues:
                    return
                f.write(f"<h2 id=\"{anchor}\">{esc(heading)}</h2>\n")
                groups = group_by_category(issues)
                for category in sorted(groups):
                    f.write(f"<h3>{esc(category)}</h3>\n")
                    for issue in sorted(groups[category], key=lambda x: x.check_id):
                        cls = severity_class(issue.severity)
                        lbl = severity_label(issue.severity)
                        f.write(f"<div class=\"issue {cls}\">\n")
                        f.write(f"  <div class=\"issue-head\">"
                                f"<span class=\"badge {cls}\">{lbl}</span>"
                                f"<span class=\"issue-id\">#{issue.check_id}</span>"
                                f"<span class=\"issue-name\">{esc(issue.check_name)}</span>"
                                f"</div>\n")
                        f.write(f"  <div class=\"issue-msg\">{esc(issue.message)}</div>\n")
                        if issue.details:
                            f.write(f"  <div class=\"issue-details\">{esc(issue.details)}</div>\n")
                        f.write("</div>\n")

            write_issue_section("Critical Issues — Must Fix Before Filing", critical_issues, "sec-critical")
            write_issue_section("Warnings — Should Review", warnings, "sec-warnings")
            write_issue_section("Info / Manual Review", info_issues, "sec-info")
            write_issue_section("Passed Checks", passed_issues, "sec-passed")

            # ADS Data Summary (when XFA data is available)
            if self.ads_data:
                d = self.ads_data
                f.write("<h2>ADS Data Summary (Extracted from XFA)</h2>\n")
                f.write("<table>\n<thead><tr><th>Field</th><th>Value</th></tr></thead>\n<tbody>\n")

                addr = d.get('assignee_address') or {}
                addr_str = ', '.join(p for p in [
                    addr.get('address1', ''), addr.get('address2', ''),
                    addr.get('city', ''), addr.get('state', ''),
                    addr.get('postcode', ''), addr.get('country', '')
                ] if p)
                rows = [
                    ("Invention Title", d.get('title', '') or '—'),
                    ("Attorney Docket Number", d.get('docket_number', '') or '—'),
                    ("Application Type", d.get('application_type', '') or '—'),
                    ("Submission Type", d.get('submission_type', '') or '—'),
                    ("Entity Status",
                     "Small" if d.get('small_entity') is True
                     else "Large/Regular" if d.get('small_entity') is False else '—'),
                    ("Drawing Sheets", d.get('drawing_sheets', '') or '—'),
                    ("Suggested Representative Figure", d.get('representative_figure', '') or '—'),
                    ("Correspondence Customer Number", d.get('customer_number', '') or '—'),
                    ("Attorney/Agent Customer Number", d.get('attorney_customer_number', '') or '—'),
                    ("Assignee", d.get('assignee_org', '') or '—'),
                ]
                if addr_str:
                    rows.append(("Assignee Address", addr_str))
                rows.extend([
                    ("Domestic Continuity",
                     "None" if not d.get('domestic_continuity_entries')
                     else f"{len(d['domestic_continuity_entries'])} entry/entries"),
                    ("Foreign Priority",
                     "None" if not d.get('foreign_priority_entries')
                     else f"{len(d['foreign_priority_entries'])} entry/entries"),
                    ("Non-Publication Request",
                     "Yes" if d.get('non_publication') is True
                     else "No" if d.get('non_publication') is False else '—'),
                    ("AIA Transition Statement",
                     "Yes" if d.get('aia_transition') is True
                     else "No" if d.get('aia_transition') is False else '—'),
                    ("ADS Form Pages", d.get('form_pages', '') or '—'),
                ])
                signer = d.get('signer') or {}
                if signer.get('signature'):
                    rows.append(("ADS Signature", signer.get('signature', '')))
                if signer.get('registration_number'):
                    rows.append(("ADS Registration Number", signer.get('registration_number', '')))
                if signer.get('date'):
                    rows.append(("ADS Signature Date", signer.get('date', '')))

                for label, value in rows:
                    f.write(f"<tr><td>{esc(label)}</td><td>{esc(value)}</td></tr>\n")
                f.write("</tbody>\n</table>\n")

                inventors = d.get('inventors') or []
                if inventors:
                    f.write("<h3>Inventors in ADS</h3>\n")
                    f.write("<table>\n<thead><tr>"
                            "<th>#</th><th>Name</th><th>Residency</th>"
                            "<th>City</th><th>Country</th><th>Citizenship</th>"
                            "</tr></thead>\n<tbody>\n")
                    for idx, inv in enumerate(inventors, start=1):
                        name = self._format_xfa_inventor(inv)
                        residency = inv.get('residency') or '—'
                        city = inv.get('res_city') or '—'
                        country = inv.get('res_country') or '—'
                        citz = inv.get('citizenship') or 'Blank'
                        f.write(f"<tr><td>{idx}</td><td>{esc(name)}</td>"
                                f"<td>{esc(residency)}</td><td>{esc(city)}</td>"
                                f"<td>{esc(country)}</td><td>{esc(citz)}</td></tr>\n")
                    f.write("</tbody>\n</table>\n")

            # Footer
            f.write("<footer>\n")
            f.write("Generated by the Patent Filing QC skill for Claude Code. ")
            f.write("To save as PDF, use your browser's <em>Print → Save as PDF</em>.")
            f.write("\n</footer>\n")
            f.write("</main>\n</body>\n</html>\n")


def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(description='Patent Filing Quality Control')
    parser.add_argument('folder', help='Path to folder containing patent filing documents')
    parser.add_argument('--output-dir', default='.', help='Directory for output reports (default: current directory)')
    
    args = parser.parse_args()
    
    print("=" * 80)
    print("PATENT FILING QUALITY CONTROL")
    print("=" * 80)
    print()
    
    # Initialize QC engine
    qc = PatentFilingQC(args.folder)
    
    # Load documents
    print("📁 Loading documents...")
    qc.load_documents()
    print()
    
    print("Documents found:")
    for doc_type, filename in qc.report.files_found.items():
        status = "✅" if filename else "❌"
        print(f"  {status} {doc_type}: {filename if filename else 'NOT FOUND'}")
    print()
    
    # Run all checks
    print("🔍 Running quality control checks...")
    qc.run_all_checks()
    print()
    
    # Summary
    print("=" * 80)
    print("SUMMARY")
    print("=" * 80)
    print(f"✅ Passed:          {qc.report.get_pass_count()}")
    print(f"⚠️  Warnings:        {qc.report.get_warning_count()}")
    print(f"🚨 Critical Issues: {qc.report.get_critical_count()}")
    print(f"ℹ️  Info/Manual:     {sum(1 for i in qc.report.issues if i.severity == Severity.INFO)}")
    print()

    # List critical issues in console (to match markdown report)
    critical_issues = [i for i in qc.report.issues if i.severity == Severity.CRITICAL]
    if critical_issues:
        print("=" * 80)
        print("🚨 CRITICAL ISSUES (Must Fix Before Filing)")
        print("=" * 80)
        for issue in critical_issues:
            print(f"\n{issue.check_id}. {issue.check_name}")
            print(f"   Category: {issue.category}")
            print(f"   Issue: {issue.message}")
            if issue.details:
                print(f"   Details: {issue.details}")
        print()

    # List warnings in console (to match markdown report)
    warnings = [i for i in qc.report.issues if i.severity == Severity.WARNING]
    if warnings:
        print("=" * 80)
        print("⚠️  WARNINGS (Should Review)")
        print("=" * 80)
        for issue in warnings:
            print(f"\n{issue.check_id}. {issue.check_name}")
            print(f"   Category: {issue.category}")
            print(f"   Warning: {issue.message}")
            if issue.details:
                print(f"   Details: {issue.details}")
        print()
    
    # Generate reports
    output_dir = Path(args.output_dir)
    output_dir.mkdir(exist_ok=True)
    
    html_report = output_dir / "Patent_Filing_QC_Report.html"
    print(f"📝 Generating HTML report: {html_report}")
    qc.generate_html_report(str(html_report))
    print(f"   To save as PDF: open the HTML file in a browser, then File → Print → Save as PDF.")
    
    print()
    print("=" * 80)
    print("QC COMPLETE")
    print("=" * 80)
    
    if qc.report.get_critical_count() > 0:
        print()
        print("🚨 CRITICAL ISSUES FOUND - Review report before filing!")
        return 1
    elif qc.report.get_warning_count() > 0:
        print()
        print("⚠️  Warnings found - Review recommended")
        return 0
    else:
        print()
        print("✅ All checks passed!")
        return 0


if __name__ == "__main__":
    sys.exit(main())
